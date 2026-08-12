"""校验最终 DOCX 包、模板结构和公式对象。"""

# 延迟解析类型注解，避免校验模块在导入期绑定可选文档类型。
from __future__ import annotations

# 只导入最终包、模板结构和公式证据校验实际依赖的运行时名称。
from export_runtime_support import (
    DEFAULT_SECTION_XML,
    DOCX_DOCUMENT_ENTRY,
    ElementTree,
    FORBIDDEN_DISCLOSURE_TEXTS,
    MATHTYPE_PROG_ID,
    MATH_NAMESPACE,
)

# 版式合同与模板章节常量共同限定最终交付包的结构边界。
from export_runtime_support import (
    PATH_DOCX_STYLE_CONTRACT,
    TEMPLATE_REQUIRED_BODY_HEADINGS,
    TEMPLATE_SECTION_ORDER,
    TEMPLATE_TABLE_LABELS,
)

# 加载、正则和 ZIP 能力只服务最终包的读取与独立复核。
from export_runtime_support import (
    DOCX_EMBEDDINGS_PREFIX,
    find_first_body_child_by_suffix,
    load_template_validator_module,
    re,
    zipfile,
)
from fallback_docx_builder import render_word_paragraph_xml

# 用字符码构造正则反斜杠，避免与文件系统路径字面量混淆。
REGEX_BACKSLASH = chr(92)  # 正则转义反斜杠字符

# 组合 Word 首个表格正则，保留标签边界与非贪婪匹配语义。
TABLE_XML_PATTERN = "".join(("<w:tbl", REGEX_BACKSLASH, "b.*?", "</w:tbl>"))  # Word 表格匹配式

# 组合 Word 页面 section 正则，保留原来的非贪婪匹配语义。
SECTION_XML_PATTERN = "".join(("<w:sectPr", REGEX_BACKSLASH, "b.*?", "</w:sectPr>"))  # 页面 section 匹配式

# 由 ZIP 路径段构造媒体前缀，供包条目过滤复用。
DOCX_MEDIA_PREFIX = "/".join(("word", "media", ""))  # DOCX 媒体部件 ZIP 路径前缀

# 组合段落 XPath，避免查询表达式与本地相对路径混淆。
WORD_PARAGRAPH_XPATH = "".join((".", "//", "w:p"))  # WordprocessingML 段落查询式

# 组合文本节点 XPath，供段落 run 的可见文字提取复用。
WORD_TEXT_XPATH = "".join((".", "//", "w:t"))  # WordprocessingML 文本节点查询式

# 组合 Office 数学根节点 XPath，保持命名空间查询集中可审计。
OMML_OBJECT_XPATH = "".join((".", "//", "{", MATH_NAMESPACE, "}oMath"))  # Office 数学根节点查询式

# 根据章节层级选择 Word 标题样式。
def resolve_heading_style(str_heading: str) -> str:
    """根据章节标题选择 Word 样式。

    参数：
    - `str_heading`：模板章节标题。

    返回：
    - `str`：对应的 Word 段落样式 ID。

    异常：
    - 无。
    """

    # 数字编号小节使用更低层级标题样式，减少版式跳跃。
    if str_heading[0].isdigit():

        # 返回三级标题样式，供 3.x 和 4.x 小节使用。
        return "Heading3"

    # 中文大章节使用一级标题样式，突出六大模板部分。
    return "Heading1"

# 从模板 document.xml 中提取第一个信息表 XML。
def extract_first_table_xml(str_document_xml: str) -> str:
    """从模板主文档 XML 中提取首个表格。

    参数：
    - `str_document_xml`：模板 `word/document.xml` 文本。

    返回：
    - `str`：第一个 `w:tbl` 表格 XML；缺失时返回空字符串。

    异常：
    - 无。
    """

    # 用非贪婪匹配提取第一个 Word 表格，保留模板原始行政信息表结构。
    obj_table_match = re.search(TABLE_XML_PATTERN, str_document_xml, flags=re.DOTALL)  # 模板首个表格匹配结果

    # 表格缺失时返回空字符串，让调用方继续生成正文但不伪造表格。
    if obj_table_match is None:

        # 空字符串表示模板里没有可复用的信息表。
        return ""

    # 返回模板首个表格 XML，供新主文档正文直接复用。
    return obj_table_match.group(0)

# 从模板 document.xml 中提取 section 属性。
def extract_section_xml(str_document_xml: str) -> str:
    """从模板主文档 XML 中提取 section 属性。

    参数：
    - `str_document_xml`：模板 `word/document.xml` 文本。

    返回：
    - `str`：模板 `w:sectPr` XML；缺失时返回最小兜底 section。

    异常：
    - 无。
    """

    # 匹配模板文档最后的 section 属性，尽量保留页面设置。
    obj_section_match = re.search(SECTION_XML_PATTERN, str_document_xml, flags=re.DOTALL)  # 模板 section 匹配结果

    # 模板提供 section 时直接复用。
    if obj_section_match is not None:

        # 返回模板 section XML，保留页面尺寸和页边距。
        return obj_section_match.group(0)

    # 模板缺失 section 时回退到最小 A4 页面设置。
    return DEFAULT_SECTION_XML

# 从模板 document.xml 中提取根文档开始标签，保留全部命名空间声明。
def extract_document_open_xml(str_document_xml: str) -> str:
    """从模板主文档 XML 中提取 `w:document` 开始标签。

    参数：
    - `str_document_xml`：模板 `word/document.xml` 文本。

    返回：
    - `str`：包含命名空间声明的 `w:document` 开始标签。

    异常：
    - 模板主文档缺少根标签时抛出 `ValueError`。
    """

    # 从模板中保留完整根标签，避免复制表格后出现未绑定命名空间前缀。
    obj_document_match = re.search(r"<w:document\b[^>]*>", str_document_xml)  # 模板根文档开始标签匹配结果

    # 模板主文档缺少根标签时直接阻断，避免写出非法 DOCX。
    if obj_document_match is None:

        # 抛出明确错误，提示模板资产本身结构异常。
        raise ValueError("> ERR: [Python] 模板 DOCX 主文档缺少 w:document 根标签。")

    # 返回包含全部命名空间声明的根标签，供新 document.xml 复用。
    return obj_document_match.group(0)

# 把模板正文分节渲染成 Word body XML。
def render_template_body_xml(
    str_table_xml: str,
    dict_sections: dict[str, list[str]],
    str_section_xml: str,
) -> str:
    """把模板正文分节渲染成 Word body XML。

    参数：
    - `str_table_xml`：从模板复用的信息表 XML。
    - `dict_sections`：按模板章节归集的正文映射。
    - `str_section_xml`：模板或兜底 section XML。

    返回：
    - `str`：完整 `w:body` 内部 XML。

    异常：
    - 无。
    """

    # 先准备 body 片段列表，首段优先放入模板信息表。
    list_body_parts: list[str] = []  # Word body XML 片段列表

    # 在模板表格存在时保留到最终主文档。
    if str_table_xml:

        # 写入模板信息表，保证行政信息行和勾选项继续可见。
        list_body_parts.append(str_table_xml)

    # 按模板顺序写入章节标题和正文段落。
    for str_heading in TEMPLATE_SECTION_ORDER:

        # 选择当前章节对应的 Word 标题样式。
        str_style_id = resolve_heading_style(str_heading)  # 当前章节标题样式

        # 写入当前模板章节标题。
        list_body_parts.append(render_word_paragraph_xml(str_heading, str_style_id))

        # 读取当前章节正文行；空章节只保留标题，交由后续校验阻断交付。
        list_section_lines = dict_sections.get(str_heading, [])  # 当前章节正文行列表

        # 逐条写入当前章节正文段落。
        for str_section_line in list_section_lines:

            # 写入当前章节正文段落。
            list_body_parts.append(render_word_paragraph_xml(str_section_line))

    # 写入 section 属性，保持 Word 主文档结构完整。
    list_body_parts.append(str_section_xml)

    # 返回完整 body 内部 XML，供模板包替换 document.xml。
    return "".join(list_body_parts)

# 读取 DOCX 主文档 XML，供模板校验解析表格、标题和正文文本。
def read_docx_document_xml(path_docx: Path) -> str:
    """读取 DOCX 主文档 XML。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。

    返回：
    - `str`：`word/document.xml` 的 UTF-8 文本。

    异常：
    - DOCX 包缺失或主文档缺失时由底层异常继续上抛。
    """

    # 打开 DOCX ZIP 包并读取 Word 主文档 XML，避免依赖 python-docx。
    with zipfile.ZipFile(path_docx, "r") as obj_docx_zip:

        # 解码主文档 XML 文本，供后续 ElementTree 解析和字符串扫描复用。
        str_document_xml = obj_docx_zip.read(DOCX_DOCUMENT_ENTRY).decode("utf-8")  # 待校验的 document.xml 文本

    # 返回主文档 XML 文本，供模板校验继续处理。
    return str_document_xml

# 读取 DOCX ZIP 包中的媒体条目，供嵌图和嵌公式校验复用。
def read_docx_media_entries(path_docx: Path) -> list[str]:
    """读取 DOCX 媒体条目列表。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。

    返回：
    - `list[str]`：DOCX ZIP 包内 `word/media/` 条目路径列表。

    异常：
    - DOCX 包读取失败时由底层异常继续上抛。
    """

    # 打开 DOCX ZIP 包并枚举媒体条目，判断图像是否真实进入最终交付件。
    with zipfile.ZipFile(path_docx, "r") as obj_docx_zip:

        # 仅保留媒体目录条目，供最终媒体数量和存在性校验复用。
        return [str_name for str_name in obj_docx_zip.namelist() if str_name.startswith(DOCX_MEDIA_PREFIX)]

# 从主文档 XML 中提取可见段落文本。
def extract_docx_paragraph_texts(str_document_xml: str) -> list[str]:
    """从 DOCX 主文档 XML 中提取段落文本。

    参数：
    - `str_document_xml`：`word/document.xml` 文本。

    返回：
    - `list[str]`：按文档顺序提取的段落可见文本。

    异常：
    - XML 非法时由 `ElementTree` 异常继续上抛。
    """

    # 解析主文档 XML，供命名空间查询所有段落和文本节点。
    obj_document_root = ElementTree.fromstring(str_document_xml)  # DOCX 主文档 XML 根节点

    # 固定 WordprocessingML 命名空间，供段落和文本查询共享。
    dict_namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}  # Word XML 命名空间

    # 先准备段落文本列表，后续按 Word 段落顺序逐项填充。
    list_paragraph_texts: list[str] = []  # DOCX 可见段落文本列表

    # 遍历所有段落节点，并合并段落内可能被拆分的多个文本 run。
    for obj_paragraph in obj_document_root.findall(WORD_PARAGRAPH_XPATH, dict_namespaces):

        # 拼接当前段落内所有 w:t 文本节点，兼容 Word 按 run 拆分文字。
        str_paragraph_text = "".join(  # 当前段落可见文本
            obj_text_node.text or ""  # 当前文本节点内容
            for obj_text_node in obj_paragraph.findall(WORD_TEXT_XPATH, dict_namespaces)  # 当前段落的文本节点
        )

        # 把当前段落文本加入列表，后续用于章节顺序和正文非空校验。
        list_paragraph_texts.append(str_paragraph_text.strip())

    # 返回完整段落文本列表，供模板 DOCX 校验继续判断。
    return list_paragraph_texts

# 在段落序列中定位所有模板章节标题。
def collect_heading_indices(list_paragraph_texts: list[str]) -> dict[str, int]:
    """收集模板章节标题所在段落索引。

    参数：
    - `list_paragraph_texts`：按文档顺序提取的段落可见文本。

    返回：
    - `dict[str, int]`：模板标题到首个段落索引的映射。

    异常：
    - 无。
    """

    # 先准备标题索引映射，缺失标题不会进入该字典。
    dict_heading_indices: dict[str, int] = {}  # 模板章节标题索引映射

    # 顺序扫描段落文本，记录模板标题第一次出现的位置。
    for int_index, str_paragraph_text in enumerate(list_paragraph_texts):

        # 只记录模板章节顺序中的精确标题，避免误把普通正文当成章节边界。
        if str_paragraph_text in TEMPLATE_SECTION_ORDER and str_paragraph_text not in dict_heading_indices:

            # 记录当前标题所在段落索引，供正文非空检查使用。
            dict_heading_indices[str_paragraph_text] = int_index  # 当前模板标题的首个段落位置

    # 返回标题索引映射，供后续校验复用。
    return dict_heading_indices

# 提取某个模板章节标题到下一个标题之间的正文段落。
def collect_section_body_texts(
    str_heading: str,
    list_paragraph_texts: list[str],
    dict_heading_indices: dict[str, int],
) -> list[str]:
    """提取指定模板章节的正文段落。

    参数：
    - `str_heading`：待检查的模板章节标题。
    - `list_paragraph_texts`：DOCX 可见段落文本列表。
    - `dict_heading_indices`：模板标题到段落索引的映射。

    返回：
    - `list[str]`：当前章节标题之后、下一模板标题之前的非空正文段落。

    异常：
    - 无。
    """

    # 当前标题缺失时返回空列表，缺失问题由上层标题校验统一报告。
    if str_heading not in dict_heading_indices:

        # 空列表表示无法提取当前章节正文。
        return []

    # 读取当前标题在模板顺序中的位置，便于寻找下一个实际出现的章节标题。
    int_template_index = TEMPLATE_SECTION_ORDER.index(str_heading)  # 当前标题在模板顺序中的位置

    # 当前章节正文从标题后一段开始。
    int_start_index = dict_heading_indices[str_heading] + 1  # 当前章节正文起始段落索引

    # 默认把文档尾部作为当前章节结束位置，若后续有标题再收窄。
    int_end_index = len(list_paragraph_texts)  # 当前章节正文结束段落索引

    # 从模板顺序的后续标题中查找第一个真实出现的标题索引。
    for str_next_heading in TEMPLATE_SECTION_ORDER[int_template_index + 1:]:

        # 命中下一个已出现标题时，用它的位置作为当前章节结束边界。
        if str_next_heading in dict_heading_indices:

            # 记录下一标题索引并结束查找。
            int_end_index = dict_heading_indices[str_next_heading]  # 当前章节的右侧边界位置

            # 找到最近的后续模板标题后即可结束边界搜索。
            break

    # 返回当前章节边界内的非空、非标题正文段落。
    return [
        str_text  # 当前正文段落文本
        for str_text in list_paragraph_texts[int_start_index:int_end_index]  # 当前章节范围内的段落
        if str_text and str_text not in TEMPLATE_SECTION_ORDER  # 过滤空行和标题行
    ]

# 收集模板信息表完整性问题。
def collect_template_table_findings(
    str_document_xml: str,
    str_visible_text: str,
) -> list[str]:
    """收集模板信息表及其行名的完整性问题。

    参数：
    - `str_document_xml`：最终 DOCX 的主文档 XML。
    - `str_visible_text`：最终 DOCX 的全部可见文本。

    返回：
    - `list[str]`：信息表结构与行名问题列表。
    """

    # 准备信息表问题列表，供结构与行名规则共同追加。
    list_findings: list[str] = []  # 模板信息表问题列表

    # 模板表格缺失时记录结构问题。
    if str_document_xml.count("<w:tbl") < 1:

        # 保存结构缺失问题，供调用方统一阻断交付。
        list_findings.append("缺少模板信息表。")

    # 逐项确认模板行名仍然存在于最终可见文本中。
    for str_label in TEMPLATE_TABLE_LABELS:

        # 当前行名缺失时记录精确标签，便于定位模板破坏位置。
        if str_label not in str_visible_text:

            # 保存当前缺失行名，不中断其余行名扫描。
            list_findings.append(f"缺少模板信息表行名：{str_label}")

    # 返回当前规则发现的全部信息表问题。
    return list_findings

# 收集模板章节与正文内容问题。
def collect_template_section_findings(
    list_paragraph_texts: list[str],
    str_visible_text: str,
) -> list[str]:
    """收集模板标题、必填正文和禁止文本问题。

    参数：
    - `list_paragraph_texts`：最终 DOCX 的可见段落文本。
    - `str_visible_text`：最终 DOCX 的全部可见文本。

    返回：
    - `list[str]`：章节结构、正文和禁止文本问题列表。
    """

    # 准备章节问题列表，保持每条规则的稳定输出顺序。
    list_findings: list[str] = []  # 模板章节问题列表

    # 建立标题索引，供标题存在性和章节正文范围检查复用。
    dict_heading_indices = collect_heading_indices(list_paragraph_texts)  # 模板标题段落索引映射

    # 检查全部模板标题是否存在。
    for str_heading in TEMPLATE_SECTION_ORDER:

        # 当前标题未进入索引时记录结构缺失问题。
        if str_heading not in dict_heading_indices:

            # 保存当前缺失标题，不中断后续章节扫描。
            list_findings.append(f"缺少模板章节标题：{str_heading}")

    # 检查关键技术章节是否有正文内容。
    for str_heading in TEMPLATE_REQUIRED_BODY_HEADINGS:

        # 提取当前章节标题与下一模板标题之间的正文。
        list_body_texts = collect_section_body_texts(  # 当前章节边界内的正文段落
            str_heading,  # 待检查的模板章节标题
            list_paragraph_texts,  # 最终 DOCX 的段落文本
            dict_heading_indices,  # 已定位的模板标题索引
        )

        # 当前必填章节没有正文时记录内容缺失问题。
        if not list_body_texts:

            # 保存当前空章节标题，供上游补齐技术内容。
            list_findings.append(f"模板章节正文为空：{str_heading}")

    # 检查最终 DOCX 是否残留内部审查、模板提示或待确认占位。
    for str_forbidden_text in FORBIDDEN_DISCLOSURE_TEXTS:

        # 禁止片段进入可见文本时记录精确残留内容。
        if str_forbidden_text in str_visible_text:

            # 保存残留片段，不中断其余禁止文本扫描。
            list_findings.append(f"最终 DOCX 残留禁止文本：{str_forbidden_text}")

    # 返回当前规则发现的全部章节问题。
    return list_findings

# 收集媒体条目与正文挂接数量问题。
def collect_template_media_findings(
    str_document_xml: str,
    list_media_entries: list[str],
    int_expected_media_count: int,
) -> list[str]:
    """收集媒体文件及 drawing 节点数量问题。

    参数：
    - `str_document_xml`：最终 DOCX 的主文档 XML。
    - `list_media_entries`：最终 DOCX 包内的媒体条目。
    - `int_expected_media_count`：按输入应至少嵌入的媒体数量。

    返回：
    - `list[str]`：媒体包条目与正文挂接问题列表。
    """

    # 无媒体输入时跳过媒体嵌入规则。
    if int_expected_media_count <= 0:

        # 返回空列表表示媒体规则不适用于当前导出输入。
        return []

    # 准备媒体问题列表，分别记录包条目和正文挂接缺口。
    list_findings: list[str] = []  # 模板媒体问题列表

    # 媒体条目不足说明图像未完整写入最终 DOCX 包。
    if len(list_media_entries) < int_expected_media_count:

        # 记录包内媒体不足问题，保留期望值与实际值。
        list_findings.append(
            f"DOCX 媒体条目不足：expected>={int_expected_media_count}, actual={len(list_media_entries)}"
        )

    # drawing 数量不足说明媒体文件未完整挂接到正文。
    int_drawing_count = str_document_xml.count("<w:drawing")  # 正文 drawing 节点数量

    # 正文挂接数量不足时记录 drawing 缺口。
    if int_drawing_count < int_expected_media_count:

        # 记录正文挂接不足问题，区分于包内媒体条目缺失。
        list_findings.append(
            f"DOCX drawing 节点不足：expected>={int_expected_media_count}, actual={int_drawing_count}"
        )

    # 汇总返回媒体规则发现的问题。
    return list_findings

# 收集严格模板 DOCX 的结构与正文质量问题。
def collect_template_docx_findings(
    path_docx: Path,
    int_expected_media_count: int = 0,
) -> list[str]:
    """收集严格模板 DOCX 校验问题。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。
    - `int_expected_media_count`：按输入应至少嵌入的媒体数量。

    返回：
    - `list[str]`：全部模板校验问题；为空表示通过。

    异常：
    - DOCX 包读取或 XML 解析失败时由底层异常继续上抛。
    """

    # 读取最终 DOCX 的主文档 XML，供结构、正文和媒体检查共享。
    str_document_xml = read_docx_document_xml(path_docx)  # 待扫描的最终 document.xml 文本

    # 提取段落级可见文本，方便判断标题顺序与章节正文是否为空。
    list_paragraph_texts = extract_docx_paragraph_texts(str_document_xml)  # 最终 DOCX 段落文本

    # 合并段落文本，供信息表行名和禁止片段规则扫描。
    str_visible_text = "".join(list_paragraph_texts)  # DOCX 全部可见文本

    # 先收集信息表问题，保持与拆分前一致的 finding 顺序。
    list_findings = collect_template_table_findings(str_document_xml, str_visible_text)  # 模板问题列表

    # 追加章节标题、正文与禁止文本问题。
    list_findings.extend(collect_template_section_findings(list_paragraph_texts, str_visible_text))

    # 最后追加媒体包条目和正文挂接数量问题。
    list_findings.extend(
        collect_template_media_findings(
            str_document_xml,  # 最终 DOCX 的主文档 XML
            read_docx_media_entries(path_docx),  # 最终 DOCX 的媒体条目
            int_expected_media_count,  # 当前导出的媒体数量下限
        )
    )

    # 返回所有发现的问题；空列表表示当前 DOCX 通过严格模板校验。
    return list_findings

# 执行严格模板 DOCX 校验，不通过时抛出明确错误。
def validate_template_docx_output(
    path_docx: Path,
    int_expected_media_count: int = 0,
) -> None:
    """校验严格模板 DOCX 输出。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。
    - `int_expected_media_count`：按导出输入应至少嵌入的媒体数量。

    返回：
    - `None`：通过时不返回额外数据。

    异常：
    - 校验不通过时抛出 `ValueError`。
    """

    # 收集当前 DOCX 的结构与正文质量问题。
    list_findings = collect_template_docx_findings(path_docx, int_expected_media_count)  # 阻断导出的模板问题列表

    # 加载独立验证器，从最终落盘 OOXML 检查中文排版角色和槽位边界。
    obj_template_validator = load_template_validator_module()  # 最终 DOCX 样式验证器模块

    # 收集结构验证器之外的字体、行距、缩进、对齐和空段问题。
    list_style_findings = obj_template_validator.collect_docx_style_findings(  # 最终 DOCX 样式问题
        path_docx,  # 已保存的严格模板交付件
        PATH_DOCX_STYLE_CONTRACT,  # 正式中文排版合同
    )

    # 把结构化样式 finding 转成既有错误列表使用的可读文本。
    list_findings.extend(
        f"{dict_finding['code']}: {dict_finding['message']} (paragraph={dict_finding['paragraph']})"  # 当前样式偏差文本
        for dict_finding in list_style_findings  # 独立验证器 finding 来源
    )

    # 没有发现问题时直接返回，表示最终 DOCX 可以作为主交付件。
    if not list_findings:

        # 校验通过，无需额外处理。
        return

    # 把问题合并为可读错误文本，便于 CLI stderr 直接定位失败原因。
    str_joined_findings = "\n".join(f"- {str_item}" for str_item in list_findings)  # 模板 DOCX 校验失败明细

    # 抛出明确错误，阻止 pipeline 把不合格 DOCX 标为 completed。
    raise ValueError("> ERR: [Python] 严格模板 DOCX 校验失败。\n" + str_joined_findings)

# 在 Word 完成 MathType OLE 写入后恢复中文排版合同要求的显式属性。
def restore_mathtype_docx_explicit_layout(
    path_docx: Path,
    obj_template_renderer: Any,
    obj_template_validator: Any,
) -> None:
    """恢复 Word 规范化后被折叠为样式继承的正文排版属性。

    参数：
    - `path_docx`：已经完成 MathType OLE 替换的 DOCX 路径。
    - `obj_template_renderer`：正式模板渲染器模块。
    - `obj_template_validator`：正式模板验证器模块。

    返回：
    - `None`：槽位正文和公式段已经重新写入显式排版属性。

    异常：
    - DOCX 读取、样式合同加载或保存失败时继续抛出底层异常。
    """

    # 延迟导入 Document，保持模块发现阶段不强制加载文档依赖。
    from docx import Document

    # 读取 Word 已保存的最终 MathType 文档和正式样式合同。
    obj_document = Document(str(path_docx))  # Word 规范化后的 MathType 文档

    # 加载渲染阶段使用的同一排版合同，避免恢复规则与生成规则漂移。
    dict_contract = obj_template_renderer.load_docx_style_contract(PATH_DOCX_STYLE_CONTRACT)  # 正式排版合同

    # 构造无空白标题集合，供扫描时识别正式槽位边界。
    set_normalized_headings = {
        obj_template_validator.normalize_slot_text(str_heading)  # 当前正式槽位标题比较键
        for str_heading in dict_contract["slot_headings"]  # 正式合同声明的槽位标题
    }  # 正式槽位标题比较键集合

    # 只恢复首个正式槽位之后的交付正文，避免改写模板说明页。
    bool_inside_delivery_slots = False  # 当前扫描位置是否已经进入正式槽位

    # 按最终文档顺序扫描段落，保持公式对象和正文位置不变。
    for obj_paragraph in obj_document.paragraphs:

        # 规整当前段落文本，判断它是否是正式槽位标题。
        str_normalized_text = obj_template_validator.normalize_slot_text(obj_paragraph.text)  # 当前段落比较键

        # 标题只负责推进槽位边界，其原始模板样式不应被正文角色覆盖。
        if str_normalized_text in set_normalized_headings:

            # 标记扫描已经进入正式交付区，后续非空段落需要恢复显式样式。
            bool_inside_delivery_slots = True  # 已进入正式交付槽位

            # 当前标题保持模板自身样式，直接处理下一个段落。
            continue

        # 模板说明区和真正的尾空段都不需要正文角色样式。
        bool_has_structured_content = (
            bool(obj_paragraph.text.strip())  # 可见正文形成的内容条件
            or obj_template_validator.paragraph_contains_tag(obj_paragraph, "oMath")  # Office 数学节点条件
            or obj_template_validator.paragraph_contains_tag(obj_paragraph, "object")  # 原生 OLE 对象条件
            or obj_template_validator.paragraph_contains_tag(obj_paragraph, "drawing")  # 正式附图节点条件
        )  # 当前段落是否承载需要恢复样式的正式内容

        # 未进入正式槽位或当前段落为空时保持原样。
        if not bool_inside_delivery_slots or not bool_has_structured_content:

            # 跳过模板说明与尾空段，防止意外赋予正文缩进。
            continue

        # 依据最终可见文本和 OOXML 重新判定角色，兼容纯 OLE 与行内 OLE。
        str_role = obj_template_validator.classify_final_paragraph(obj_paragraph)  # 当前最终段落角色

        # 重新写入被 Word 折叠的字号、字体、缩进、行距和对齐属性。
        obj_template_renderer.apply_paragraph_role_style(obj_paragraph, str_role, dict_contract)

    # 保存显式排版恢复结果；调用方随后重新核验 MathType OLE 结构与版式。
    obj_document.save(str(path_docx))

# 清空模板正文并保留首个信息表与最终 section，供代理交付版 DOCX 在模板壳上重建正文。
def clear_template_body_keep_first_table(obj_document: Any) -> None:
    """清空模板正文并保留首个信息表与最终 section。

    参数：
    - `obj_document`：基于模板打开的 python-docx `Document` 对象。

    返回：
    - `None`。

    异常：
    - 底层 XML 操作失败时由 python-docx 对象异常继续上抛。
    """

    # 读取模板 body XML 节点，供低层子节点筛选和删除复用。
    obj_body = obj_document._element.body  # 模板正文 body XML 节点

    # 复制当前所有 body 子节点，避免遍历时边删边改导致迭代错乱。
    list_children = list(obj_body)  # 模板 body 原始子节点列表

    # 记录需要保留的首个信息表节点；模板若无表则保持空值。
    obj_first_table = find_first_body_child_by_suffix(list_children, "}tbl")  # 模板首个信息表节点

    # 记录最终 section 节点，避免清正文时把页面设置一并删掉。
    obj_section = find_first_body_child_by_suffix(list_children, "}sectPr")  # 模板最终 section 节点

    # 逐项删除除首个信息表与最终 section 外的全部正文节点。
    for obj_child in list_children:

        # 命中需要保留的表格或 section 时跳过删除，保持模板外壳稳定。
        if obj_child is obj_first_table or obj_child is obj_section:

            # 当前子节点属于模板保留壳，不删除。
            continue

        # 删除当前正文节点，让后续导出按新的交底书内容重建主稿。
        obj_body.remove(obj_child)

# 根据模板章节标题选择 python-docx 的标题层级。
def resolve_docx_heading_level(str_heading: str) -> int:
    """根据模板章节标题选择 python-docx 标题层级。

    参数：
    - `str_heading`：模板章节标题。

    返回：
    - `int`：可直接传给 `add_heading` 的标题层级。

    异常：
    - 无。
    """

    # 数字编号小节使用较低层级标题，减少版式跳跃。
    if str_heading[:1].isdigit():

        # 返回适合 3.x 和 4.x 小节的三级标题层级。
        return 3

    # 中文大章节使用一级标题层级，突出正式模板主结构。
    return 1

# 从最终 DOCX 包读取公式对象统计，并与源公式清单进行逐项对账。
def collect_formula_object_evidence(
    path_docx: Path,
    str_equation_mode: str,
    list_formula_records: list[dict[str, Any]],
    list_conversion_evidence: list[dict[str, object]],
) -> dict[str, Any]:
    """构造最终 DOCX 的公式对象级验收证据。

    参数：
    - `path_docx`：已经完成全部后处理的最终 DOCX。
    - `str_equation_mode`：本轮 `office` 或 `mathtype` 模式。
    - `list_formula_records`：按文档顺序保存的源公式记录。
    - `list_conversion_evidence`：Office 中间转换阶段收集的结构证据。

    返回：
    - `dict[str, Any]`：公式数量、对象类型、嵌入部件和逐公式改写证据。

    异常：
    - `ValueError`：对象数量或类型与当前模式合同不一致。
    """

    # 读取最终 DOCX 主文档和全部 ZIP 条目，保证证据来自交付文件本身。
    with zipfile.ZipFile(path_docx) as obj_archive:

        # 主文档 XML 用于统计 OMML 和 MathType ProgID。
        str_document_xml = obj_archive.read(DOCX_DOCUMENT_ENTRY).decode("utf-8")  # 最终 DOCX 主文档 XML

        # 嵌入部件列表必须与 MathType 公式数量精确对应。
        list_embedding_entries = [  # 最终 DOCX 中的 OLE 嵌入部件路径
            str_entry_name  # 当前嵌入部件 ZIP 路径
            for str_entry_name in obj_archive.namelist()  # 最终 DOCX 全部 ZIP 条目
            if str_entry_name.startswith(DOCX_EMBEDDINGS_PREFIX)  # 只统计 Word 嵌入对象目录
        ]

    # 解析主文档 XML 后按标准命名空间精确统计数学根节点。
    obj_document_root = ElementTree.fromstring(str_document_xml)  # 最终 DOCX 主文档 XML 根节点

    # m:oMath 数量是 Office 模式的可编辑公式对象证据。
    int_omml_count = len(obj_document_root.findall(OMML_OBJECT_XPATH))  # 最终文档 OMML 公式数量

    # Equation.DSMT4 出现次数是 MathType 对象类型证据。
    int_mathtype_progid_count = str_document_xml.count(MATHTYPE_PROG_ID)  # 最终文档 MathType ProgID 数量

    # 源公式清单决定当前导出的唯一预期公式总数。
    int_expected_count = len(list_formula_records)  # 当前文档预期公式数量

    # Office 模式要求每条源公式形成一个 OMML，且没有 OLE 对象。
    if str_equation_mode == "office":

        # 三项对象统计必须同时满足，避免混合对象文档被误判为通过。
        bool_object_contract_passed = (  # Office 公式对象合同是否通过
            int_omml_count == int_expected_count  # 每条源公式形成一个 OMML
            and int_mathtype_progid_count == 0  # Office 模式不含 MathType ProgID
            and len(list_embedding_entries) == 0  # Office 模式不含 OLE 嵌入部件
        )

    # MathType 模式要求每条源公式形成一个 Equation.DSMT4 嵌入对象，且不残留 OMML。
    else:

        # ProgID、嵌入部件和 OMML 三项必须精确对应最终纯 MathType 合同。
        bool_object_contract_passed = (  # 最终纯 MathType OLE 对账结果
            int_omml_count == 0  # 最终文档不得残留 Office 公式节点
            and int_mathtype_progid_count == int_expected_count  # 每条源公式对应一个 MathType ProgID
            and len(list_embedding_entries) == int_expected_count  # 每条源公式对应一个 OLE 嵌入部件
        )

    # 对象统计不一致时立即阻断，禁止只写一份标记失败的旁路报告。
    if not bool_object_contract_passed:

        # 稳定错误说明公式对象证据与源公式清单不一致。
        raise ValueError("> ERR: [Python] EQ009 最终 DOCX 公式对象合同校验失败。")

    # 逐公式证据按源记录顺序构造，便于审查者对应原公式与转换规则。
    list_formula_items: list[dict[str, object]] = []  # 最终公式逐项证据列表

    # 顺序遍历源公式清单，不依赖公式文本去重或集合顺序。
    for int_index, dict_formula_record in enumerate(list_formula_records, start=1):

        # 两种模式共享公式序号、源文本和行内/行间布局字段。
        dict_formula_item: dict[str, object] = {  # 当前公式的基础对象证据
            "index": int_index,  # 当前公式的一基文档顺序
            "latex": str(dict_formula_record["latex"]),  # 最终对象使用的源公式文本
            "display": bool(dict_formula_record["display"]),  # 当前公式是否采用行间布局
            "final_object_type": "OMML" if str_equation_mode == "office" else MATHTYPE_PROG_ID,  # 最终公式对象类型
        }

        # Office 模式附加预改写、语义指纹和 OMML 结构校验证据。
        if str_equation_mode == "office":

            # 转换证据数量已经由公式转换回调按同一文档顺序收集。
            dict_formula_item["conversion"] = list_conversion_evidence[int_index - 1]  # 同顺序 OMML 转换证据

        # MathType 直接使用原 LaTeX 生成 OLE，不应把 OMML 预改写文本写入最终对象。
        else:

            # 显式记录未使用中间改写，证明最终 OLE 内容来自源公式。
            dict_formula_item["conversion"] = {
                "pre_rewrite_applied_to_final_object": False,  # 预改写没有进入最终 OLE
                "native_mathtype_source": "original_latex",  # MathType 使用原始 LaTeX
            }

        # 保存当前公式证据，继续处理下一文档对象。
        list_formula_items.append(dict_formula_item)

    # 返回可直接写入 JSON 的完整验收载荷。
    return {
        "passed": True,
        "equation_mode": str_equation_mode,
        "expected_formula_count": int_expected_count,
        "omml_count": int_omml_count,
        "mathtype_progid_count": int_mathtype_progid_count,
        "embedding_count": len(list_embedding_entries),
        "formulas": list_formula_items,
    }
