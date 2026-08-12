#!/usr/bin/env python3
"""把正式交底书 Markdown 导出为 DOCX。

参数：
- 无。

返回：
- 无。

异常：
- 无。
"""
# 启用未来版本注解行为，保证类型标注在当前解释器下稳定可用。
from __future__ import annotations
# 引入参数解析和按路径加载模块能力，供导出入口解析参数并加载共享支持模块。
import argparse
import importlib.util

# 引入环境变量、正则和子进程能力，供运行时切换与正文解析逻辑复用。
import os
import re
import subprocess

# 引入标准输出、临时目录和 ZIP 打包能力，供导出流程写回结果并处理 DOCX 包。
import sys
import tempfile
import zipfile

# 引入路径、类型和 XML 处理能力，供正式 DOCX 渲染与校验逻辑复用。
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

# 引入 XML 转义能力，供标准库回退导出路径安全写入正文文本。
from xml.sax.saxutils import escape

# 固定共享运行时支持模块位置，避免通过改写 sys.path 查找公共工具。
PATH_RUNTIME_SUPPORT = Path(__file__).resolve().parents[1] / "support" / "runtime_support.py"  # 共享运行时支持模块路径

# 固定模板槽位渲染器路径，避免脚本直执行与测试按路径导入时依赖 sys.path 副作用。
PATH_TEMPLATE_RENDERER = Path(__file__).resolve().with_name("template_docx_renderer.py")  # 模板槽位渲染器路径

# 固定默认模板路径，供 python-docx 增强路径按需读取页面版式。
DEFAULT_TEMPLATE = Path(__file__).resolve().parents[3] / "assets" / "cn_technical_disclosure_template.docx"  # 默认模板 DOCX 路径

# 用环境变量标记当前进程是否已经切到文档运行时，避免模板导出递归重启。
ENV_TEMPLATE_RUNTIME_REEXEC = "READABLE_PATENT_EXPORT_DOCX_REEXEC"  # 模板运行时重启标记环境变量

# 固定 Pillow 公式回退图的探针画布尺寸，避免测量文本边界时使用过小底图。
FORMULA_FALLBACK_PROBE_SIZE = 32  # Pillow 探针画布尺寸

# 固定 Pillow 公式回退图的最小内容宽度，避免很短公式生成过窄图片。
FORMULA_FALLBACK_TEXT_WIDTH_DEFAULT = 160  # 回退公式图的默认文本宽度

# 固定 Pillow 公式回退图的默认单行高度，作为文本测量异常时的兜底值。
FORMULA_FALLBACK_LINE_HEIGHT_DEFAULT = 16  # 回退公式图的默认单行高度

# 固定 Pillow 公式回退图的水平留白，保证公式文本不会紧贴图片边缘。
FORMULA_FALLBACK_HORIZONTAL_PADDING = 24  # 回退公式图的左右留白

# 固定 Pillow 公式回退图的垂直留白，保证公式文本不会紧贴上下边缘。
FORMULA_FALLBACK_VERTICAL_PADDING = 20  # 回退公式图的上下留白

# 固定 Pillow 公式回退图的行间距，避免多行文本粘连。
FORMULA_FALLBACK_LINE_GAP = 8  # 回退公式图的多行文本间距

# 固定 Pillow 公式回退图的最小宽度，避免单短式图片过窄影响阅读。
FORMULA_FALLBACK_MIN_WIDTH = 280  # 回退公式图的最小宽度

# 固定 Pillow 公式回退图的最小高度，避免单行公式图片过矮。
FORMULA_FALLBACK_MIN_HEIGHT = 96  # 回退公式图的最小高度

# 基于系统环境变量和当前用户主目录盘符推导 Windows 字体目录，避免在源码里写死本机盘符。
def build_windows_fonts_dir() -> Path:
    """推导当前环境下的 Windows 字体目录。

    参数：
    - 无。

    返回：
    - `Path`：Windows 字体目录路径。

    异常：
    - 无。
    """

    # 先按环境变量与当前用户盘符顺序推导系统根目录，兼容本机与受限运行环境。
    str_windows_root = os.environ.get("WINDIR")  # 优先尝试当前进程显式暴露的 Windows 根目录

    # WINDIR 缺失时，继续尝试兼容旧式环境变量名。
    if not str_windows_root:

        # 再读取 SystemRoot，兼容只暴露旧式系统根目录变量的运行环境。
        str_windows_root = os.environ.get("SystemRoot")  # 兼容只提供旧式变量名的 Windows 根目录

    # 系统环境变量仍缺失时，再回退到当前用户主目录盘符。
    if not str_windows_root:

        # 主目录盘符通常仍能反映当前 Windows 安装所在盘符。
        str_windows_root = Path.home().anchor  # 从当前用户主目录反推出本机系统盘符

    # 如果以上候选都为空，再用平台根路径保底，避免得到空路径。
    if not str_windows_root:

        # 平台根路径是最后的兜底值，保证后续字体目录拼装总有基准。
        str_windows_root = os.sep  # 在非 Windows 受限环境下保留根路径兜底

    # 再在系统根目录下拼出 Fonts 子目录，供字体探测逻辑稳定复用。
    path_windows_fonts_dir = Path(str_windows_root) / "Fonts"  # Windows 字体目录路径

    # 返回推导出的 Windows 字体目录。
    return path_windows_fonts_dir

# 固定 Windows 字体目录路径，供公式 Pillow fallback 优先解析支持运算符的 TrueType 字体。
WINDOWS_FONTS_DIR = build_windows_fonts_dir()  # 公式回退使用的字体目录

# 固定公式文本回退优先尝试的字体文件名，优先覆盖乘号与常见运算符字形。
FORMULA_FALLBACK_FONT_FILENAMES = (  # 公式回退优先字体文件名
    "cambria.ttc",  # Cambria Math 常见可用字体
    "calibri.ttf",  # Calibri 常见办公字体
    "arial.ttf",  # Arial 对常见数学运算符支持稳定
    "msyh.ttc",  # 微软雅黑兼容中文与基础数学符号
    "simhei.ttf",  # 黑体兼容中文与基础数学符号
)  # 公式回退优先字体文件名序列

# 固定 Pillow 公式回退字体字号，避免默认位图字体过小且缺少符号字形。
FORMULA_FALLBACK_FONT_SIZE = 28  # 公式回退 TrueType 字号

# 构造 Codex 主运行时 Python 路径，供当前解释器缺少 python-docx 时复用文档导出能力。
def build_codex_bundled_python_path() -> Path:
    """构造 Codex 主运行时 Python 路径。

    参数：
    - 无。

    返回：
    - `Path`：Codex 主运行时 Python 可执行文件路径。

    异常：
    - 无。
    """

    # 先登记 Codex 主运行时依赖目录段，避免完整路径拼接落成过长单行。
    tuple_runtime_segments = (".cache", "codex-runtimes", "codex-primary-runtime", "dependencies")  # Codex 主运行时依赖目录段

    # 再用目录段列表拼出依赖根目录，后续统一在此基础上定位外部 Python。
    path_runtime_root = Path.home().joinpath(*tuple_runtime_segments)  # Codex 主运行时依赖根目录

    # 再在依赖根目录下拼出 Python 可执行文件路径，供模板导出回退逻辑复用。
    path_bundled_python = path_runtime_root / "python" / "python.exe"  # Codex 主运行时 Python 路径

    # 返回构造好的 Codex 文档运行时 Python 路径。
    return path_bundled_python

# 固定 Codex 主运行时 Python 路径，供模板导出受控切换到外部运行时时复用。
PATH_CODEX_BUNDLED_PYTHON = build_codex_bundled_python_path()  # 模板导出回退使用的外部 Python 路径

# 固定常见 LaTeX 命令到可读文本的映射，供纯文本公式 fallback 做语义不变改写。
FORMULA_FALLBACK_LATEX_REPLACEMENTS = (  # 公式 fallback 的 LaTeX 命令替换映射
    (r"\times", "×"),  # 乘号命令改写为直接可读乘号
    (r"\cdot", "·"),  # 点乘命令改写为中点运算符
    (r"\leq", "≤"),  # 小于等于命令改写为 Unicode 运算符
    (r"\geq", "≥"),  # 大于等于命令改写为 Unicode 运算符
    (r"\neq", "≠"),  # 不等于命令改写为 Unicode 运算符
    (r"\approx", "≈"),  # 近似命令改写为 Unicode 运算符
    (r"\pm", "±"),  # 正负号命令改写为 Unicode 运算符
    (r"\left", ""),  # 左定界尺寸命令直接去掉，保留原括号正文
    (r"\right", ""),  # 右定界尺寸命令直接去掉，保留原括号正文
    (r"\_", "_"),  # 转义下划线改回正文下划线字符
    (r"\,", " "),  # 紧缩空白命令改成普通空格，保证纯文本可读性
    (r"\;", " "),  # 中等空白命令改成普通空格，保证纯文本可读性
)

# 预编译多空白压缩规则，供纯文本公式 fallback 规整输出间距。
RE_FORMULA_MULTISPACE = re.compile(r"\s+")  # 公式 fallback 多空白压缩规则

# 规整一组文本行，只保留去首尾空白后的非空结果。
def collect_nonempty_stripped_lines(list_lines: list[str]) -> list[str]:
    """收集去空白后的非空文本行。

    参数：
    - `list_lines`：待规整的原始文本行列表。

    返回：
    - `list[str]`：去首尾空白并过滤空白行后的文本行列表。

    异常：
    - 无。
    """

    # 逐行去首尾空白并过滤空结果，供正文与公式规整逻辑共享。
    list_clean_lines = [str_line.strip() for str_line in list_lines if str_line.strip()]  # 去空白后的非空文本行列表

    # 返回规整后的非空文本行列表，供调用方继续拼接正文或公式块。
    return list_clean_lines

# 把 LaTeX 风格公式改写成适合纯文本渲染的可读表达，避免 fallback 泄露源码命令。
def normalize_formula_text_for_fallback(str_formula_text: str) -> str:
    """把 LaTeX 风格公式改写成纯文本可读表达。

    参数：
    - `str_formula_text`：待规整的原始公式正文文本。

    返回：
    - `str`：去掉主要 LaTeX 命令后的可读公式文本。

    异常：
    - 无。
    """

    # 先去掉首尾空白，避免替换前后的偶发换行影响纯文本公式可读性。
    str_readable_formula = str_formula_text.strip()  # 待改写的纯文本公式正文

    # 逐项替换常见 LaTeX 命令，把代理侧能直接阅读的等价符号写回正文。
    for str_source_token, str_target_token in FORMULA_FALLBACK_LATEX_REPLACEMENTS:

        # 把当前 LaTeX 命令改写成目标可读字符或普通空格。
        str_readable_formula = str_readable_formula.replace(str_source_token, str_target_token)  # 当前替换后的公式正文

    # 去掉成组控制用的大括号，保留组内运算内容本身。
    str_readable_formula = str_readable_formula.replace("{", "").replace("}", "")  # 去掉分组包裹符后的公式正文

    # 把多余空白压缩成单空格，避免命令替换后出现难读的空隙堆叠。
    str_readable_formula = RE_FORMULA_MULTISPACE.sub(" ", str_readable_formula).strip()  # 压缩空白后的可读公式正文

    # 返回规整后的可读公式文本，供纯文本 fallback 图片直接渲染。
    return str_readable_formula

# 解析首个可用于公式纯文本回退的 TrueType 字体路径。
def find_formula_fallback_font_path() -> Path | None:
    """解析公式纯文本回退优先字体路径。

    参数：
    - 无。

    返回：
    - `Path | None`：命中时返回首个可用字体路径，否则返回 `None`。

    异常：
    - 无。
    """

    # 按优先顺序逐项检查字体文件是否真实存在，保证符号字形尽量稳定。
    for str_filename in FORMULA_FALLBACK_FONT_FILENAMES:

        # 拼出当前候选字体路径，供存在性检查复用。
        path_candidate_font = WINDOWS_FONTS_DIR / str_filename  # 当前候选字体路径

        # 找到首个存在的候选字体后立即返回，减少重复扫描。
        if path_candidate_font.exists():

            # 返回首个真实存在的候选字体路径，供 Pillow TrueType 回退复用。
            return path_candidate_font

    # 当前环境没有命中任何候选字体时返回空值，让上层继续走默认字体回退。
    return None

# 为公式 Pillow fallback 构造优先支持运算符的字体对象。
def build_formula_fallback_font(class_image_font: Any) -> Any:
    """为公式 Pillow fallback 构造字体对象。

    参数：
    - `class_image_font`：Pillow 的字体模块对象；shape=单个模块句柄，dtype=runtime object，unit=none。

    返回：
    - `Any`：优先返回 TrueType 字体对象；shape=单个字体句柄，dtype=Pillow font object，unit=font handle。

    异常：
    - 无；TrueType 字体加载失败时直接回退默认字体。
    """

    # 先解析当前环境可用的公式回退字体路径，供 TrueType 加载复用。
    path_font = find_formula_fallback_font_path()  # 当前环境可用公式回退字体路径

    # 在当前环境命中了真实可用字体时优先构造 TrueType 字体对象。
    if path_font is not None:

        # 尝试按真实字体路径加载 TrueType 字体，优先保证乘号等符号能正确绘制。
        try:

            # 返回命中字形更完整的 TrueType 字体对象，供公式文本回退图复用。
            return class_image_font.truetype(str(path_font), FORMULA_FALLBACK_FONT_SIZE)

        # 字体文件虽然存在但无法加载时立即改走默认字体，避免当前公式图因为单个字库异常而失败。
        except Exception:

            # 直接回退默认字体，保证极端环境下仍能完成受控导出。
            return class_image_font.load_default()

    # 回退到 Pillow 默认字体，保证极端环境下至少仍能完成导出链路。
    return class_image_font.load_default()

# 构造文档运行时最小依赖探针命令，供子进程能力检查复用。
def build_template_runtime_probe_command(path_python: Path) -> list[str]:
    """构造文档运行时最小依赖探针命令。

    参数：
    - `path_python`：待探测的 Python 可执行文件路径。

    返回：
    - `list[str]`：可直接传给 `subprocess.run` 的探针命令参数列表。

    异常：
    - 无。
    """

    # 组织最小导入探针命令，只验证 `docx` 和 `PIL` 是否在目标运行时可用。
    list_probe_command = [str(path_python), "-c", "import docx, PIL"]  # 文档运行时最小依赖探针命令

    # 返回最小依赖探针命令，供运行时探测逻辑直接执行。
    return list_probe_command

# 构造文档运行时重启命令，供模板导出缺包时复用。
def build_template_runtime_reexec_command(path_python: Path) -> list[str]:
    """构造文档运行时重启命令。

    参数：
    - `path_python`：待执行的外部 Python 可执行文件路径。

    返回：
    - `list[str]`：可直接传给 `subprocess.run` 的重启命令参数列表。

    异常：
    - 无。
    """

    # 组织重启命令，保持脚本路径和原始命令行参数在外部运行时中原样复用。
    list_reexec_command = [str(path_python), str(Path(__file__).resolve()), *sys.argv[1:]]  # 文档运行时重启命令

    # 返回受控重启命令，供模板导出回退逻辑直接执行。
    return list_reexec_command

# 在模板 body 节点列表中按后缀查找第一个匹配节点。
def find_first_body_child_by_suffix(list_children: list[Any], str_suffix: str) -> Any | None:
    """查找首个命中后缀的 body 子节点。

    参数：
    - `list_children`：模板 body 子节点列表。
    - `str_suffix`：目标 XML 标签后缀。

    返回：
    - `Any | None`：命中时返回首个匹配节点，否则返回 `None`。

    异常：
    - 无。
    """

    # 从模板 body 子节点中按顺序定位首个命中后缀的候选节点。
    obj_matched_child = next((obj_child for obj_child in list_children if obj_child.tag.endswith(str_suffix)), None)  # 首个命中后缀的 body 子节点

    # 返回首个命中的 body 子节点，供模板表格和 section 提取逻辑复用。
    return obj_matched_child

# 向当前画布中央写入公式文本，供数学模式和纯文本回退模式复用。
def write_centered_formula_text(
    obj_axes: Any,
    str_formula_text: str,
    int_fontsize: int,
    bool_use_math_mode: bool,
) -> Any:
    """向当前画布中央写入公式文本。

    参数：
    - `obj_axes`：当前公式图画布对象。
    - `str_formula_text`：待写入的公式正文文本。
    - `int_fontsize`：当前文本字号。
    - `bool_use_math_mode`：是否使用数学公式模式包装文本。

    返回：
    - `Any`：matplotlib 返回的文本对象。

    异常：
    - 底层文本写入失败时由 matplotlib 异常继续上抛。
    """

    # 根据当前渲染模式决定要写入画布的真实文本内容。
    str_render_text = f"${str_formula_text}$" if bool_use_math_mode else str_formula_text  # 当前要写入画布的文本内容

    # 把文本居中写到当前公式画布，供后续边界框估算和图片导出复用。
    obj_text = obj_axes.text(0.5, 0.5, str_render_text, ha="center", va="center", fontsize=int_fontsize)  # 当前画布的公式文本对象

    # 返回当前公式文本对象，供上层继续测量边界框和导出图片。
    return obj_text

# 测量多行公式文本块的整体边界框，供 Pillow 回退图尺寸计算复用。
def measure_formula_block_bbox(
    obj_probe_draw: Any,
    str_formula_block: str,
    obj_font: Any,
) -> Any:
    """测量多行公式文本块的整体边界框。

    参数：
    - `obj_probe_draw`：探针画布绘图对象。
    - `str_formula_block`：待测量的多行公式文本块。
    - `obj_font`：当前使用的 Pillow 字体对象。

    返回：
    - `Any`：Pillow 返回的多行文本块边界框对象。

    异常：
    - 底层边界框测量失败时由 Pillow 异常继续上抛。
    """

    # 直接返回多行公式文本块的整体边界框，减少调用方逐行统计的噪声。
    return obj_probe_draw.multiline_textbbox(
        (0, 0),
        str_formula_block,
        font=obj_font,
        spacing=FORMULA_FALLBACK_LINE_GAP,
        align="center",
    )

# 预编译 Markdown 标题匹配规则，统一提取标题层级和标题正文。
RE_HEADING = re.compile(r"^(#{1,6})\s+(.+)$")  # Markdown 标题匹配规则

# 预编译 Markdown 无序列表匹配规则，统一识别项目符号行。
RE_BULLET = re.compile(r"^\s*[-*]\s+(.+)$")  # Markdown 无序列表匹配规则

# 预编译 Markdown 有序列表匹配规则，统一识别编号条目行。
RE_ORDERED = re.compile(r"^\s*\d+[.)、]\s+(.+)$")  # Markdown 有序列表匹配规则

# 预编译 Markdown 表格分隔行规则，避免把 `---` 样式行写入正文。
RE_TABLE_SEPARATOR = re.compile(r"^:?-{3,}:?$")  # Markdown 表格分隔行匹配规则

# 预编译 Markdown 行内公式规则，供回退路径去掉 `$...$` 源语法标记。
RE_INLINE_FORMULA = re.compile(r"\$([^$\n]+)\$")  # Markdown 行内公式匹配规则

# 固定 Word 标题层级上限，避免正文标题样式超出最小样式集。
WORD_HEADING_LEVEL_LIMIT = 4  # Word 标题层级上限

# 固定标准库 DOCX 回退导出的 A4 页面宽度。
WORD_PAGE_WIDTH = 11906  # A4 页面宽度

# 固定标准库 DOCX 回退导出的 A4 页面高度。
WORD_PAGE_HEIGHT = 16838  # A4 页面高度

# 固定标准库 DOCX 回退导出的统一页边距。
WORD_PAGE_MARGIN = 1440  # DOCX 页面边距

# 固定标准库 DOCX 回退导出的页眉页脚边距。
WORD_HEADER_FOOTER_MARGIN = 720  # 页眉页脚边距

# 固定正文中的代码块占位文本，提醒评审人在提交前补正式附件内容。
TEXT_ATTACHMENT_PLACEHOLDER = "[代码块或图表示意已移入附件，请在提交前替换为正式内容。]"  # 正文中的附件占位文本

# 固定附件章节标题，集中收纳正文中摘出的代码块和图表示意。
TEXT_ATTACHMENT_TITLE = "附件：待人工转写的代码块与图表示意"  # 附件章节标题

# 固定导出说明标题文本，保证 sidecar 说明文件格式稳定。
TEXT_EXPORT_NOTE_TITLE = "# Export Note"  # 导出说明标题文本

# 固定提交说明标题文本，用来承接不进入主 DOCX 的行政空白和内部审查内容。
TEXT_SUBMISSION_NOTE_TITLE = "# Submission Note"  # 提交说明标题文本

# 固定模板正文章节顺序，确保导出件始终呈现专利技术交底书主结构。
TEMPLATE_SECTION_ORDER = [  # 模板正文章节顺序
    "一、发明名称",  # 发明名称章节
    "二、所属技术领域",  # 所属技术领域章节
    "三、现有技术（背景技术）",  # 背景技术总章
    "3.1相关技术背景以及最接近的现有技术",  # 背景与最接近现有技术小节
    "3.2与本发明最相似的现有技术实现方案",  # 最相似现有技术实现方案小节
    "3.3现有技术的缺点",  # 现有技术缺点小节
    "四、发明内容：",  # 发明内容总章
    "4.1 发明目的",  # 发明目的小节
    "4.2 技术解决方案",  # 技术解决方案小节
    "4.2.1 装置、结构类",  # 装置结构类小节
    "4.2.2 方法类",  # 方法类小节
    "4.3、技术效果",  # 技术效果小节
    "五、附图及附图的简单说明",  # 附图说明章节
    "六、具体实施方式",  # 具体实施方式章节
]

# 固定 Markdown 章节到模板章节的映射，兼容旧草稿标题和严格模板标题。
TEMPLATE_HEADING_ALIASES = {  # Markdown 标题归一化映射
    "一、发明名称": "一、发明名称",  # 保留已合规发明名称标题
    "二、所属技术领域": "二、所属技术领域",  # 保留已合规技术领域标题
    "三、现有技术": "三、现有技术（背景技术）",  # 旧版总章补全括号说明
    "三、现有技术（背景技术）": "三、现有技术（背景技术）",  # 接收模板原文背景标题
    "3.1 相关技术背景": "3.1相关技术背景以及最接近的现有技术",  # 旧版背景标题扩展为模板长标题
    "3.1相关技术背景以及最接近的现有技术": "3.1相关技术背景以及最接近的现有技术",  # 接收模板原文 3.1
    "3.2 最接近现有技术": "3.2与本发明最相似的现有技术实现方案",  # 旧版 3.2 标题扩展为模板表述
    "3.2与本发明最相似的现有技术实现方案": "3.2与本发明最相似的现有技术实现方案",  # 保持相似方案锚点原样输出
    "3.3 现有技术缺点": "3.3现有技术的缺点",  # 旧版缺点标题去掉多余空格
    "3.3现有技术的缺点": "3.3现有技术的缺点",  # 保持缺点小节锚点不改名
    "四、发明内容": "四、发明内容：",  # 旧版发明内容补模板冒号
    "四、发明内容：": "四、发明内容：",  # 接收模板原文发明内容标题
    "4.1 发明目的": "4.1 发明目的",  # 目的小节无需改写
    "4.2 技术解决方案": "4.2 技术解决方案",  # 解决方案小节无需改写
    "4.2.1 方法流程": "4.2.2 方法类",  # 旧方法小节迁到模板方法类
    "4.2.1 装置、结构类": "4.2.1 装置、结构类",  # 接收模板装置结构小节
    "4.2.2 系统/装置方案": "4.2.1 装置、结构类",  # 旧装置小节迁到模板装置类
    "4.2.2 方法类": "4.2.2 方法类",  # 接收模板方法类小节
    "4.3 技术效果": "4.3、技术效果",  # 旧技术效果标题补顿号
    "4.3、技术效果": "4.3、技术效果",  # 接收模板原文效果标题
    "五、附图及附图说明": "五、附图及附图的简单说明",  # 旧附图标题补全模板措辞
    "五、附图及附图的简单说明": "五、附图及附图的简单说明",  # 接收模板原文附图标题
    "六、具体实施方式": "六、具体实施方式",  # 实施方式章节无需改写
}

# 固定内部审查章节前缀，导出主 DOCX 时必须转移到 sidecar。
INTERNAL_SECTION_PREFIXES = ("七、", "八、", "九、")  # 内部审查章节前缀

# 固定模板表格行名，行政信息允许空白但必须在提交说明中显式列出。
ADMIN_LABELS = [  # 进入 sidecar 的可空行政字段顺序
    "申请人名称",  # 机构主体字段允许后补
    "申请人地址",  # 地址字段通常由代理确认
    "发明人排名",  # 发明人顺序需人工核对
    "第一发明人身份证号码",  # 证件号码不得自动推断
    "交底书撰写人",  # 撰写人信息可由提交方补录
    "手机号码",  # 手机联系方式不从技术材料猜测
    "办公电话",  # 办公电话缺失不阻塞技术正文
    "E-mail",  # 邮箱字段保留给人工填写
    "所属项目",  # 项目归属不由生成链路臆造
]

# 固定最终 DOCX 信息表必须可见的行名，防止模板表格被导出器误删。
TEMPLATE_TABLE_LABELS = ["发明名称", "发明创造类型", *ADMIN_LABELS]  # 模板信息表完整行名

# 固定必须具备正文内容的技术章节，父级总章只要求标题存在。
TEMPLATE_REQUIRED_BODY_HEADINGS = [  # 需要在最终主 DOCX 中有正文的叶子章节
    "一、发明名称",  # 名称不能只停留在信息表
    "二、所属技术领域",  # 技术领域必须给代理判断分类
    "3.1相关技术背景以及最接近的现有技术",  # 背景小节需要事实上下文
    "3.2与本发明最相似的现有技术实现方案",  # 相似方案用于区别特征定位
    "3.3现有技术的缺点",  # 缺点小节承接发明目的
    "4.1 发明目的",  # 目的小节解释解决方向
    "4.2.1 装置、结构类",  # 装置小节需明确适用或不适用
    "4.2.2 方法类",  # 方法小节承载步骤方案
    "4.3、技术效果",  # 效果小节支撑代理撰写有益效果
    "五、附图及附图的简单说明",  # 附图小节说明图号含义
    "六、具体实施方式",  # 实施方式小节提供可实施样例
]

# 固定最终 DOCX 正文禁止残留的内部文本，避免把审查材料交给代理。
FORBIDDEN_DISCLOSURE_TEXTS = [  # 校验最终主 DOCX 时直接阻断的文本片段
    "```",  # 代码围栏不得进入代理提交件
    "【",  # 左提示括号说明模板残留
    "】",  # 右提示括号说明模板残留
    "待确认",  # 待确认项只能进 sidecar
    "TODO",  # 英文大写待办属于内部痕迹
    "todo",  # 英文小写待办也按占位处理
    TEXT_ATTACHMENT_PLACEHOLDER,  # 代码块占位不是正式交底书正文
    "七、术语说明",  # 术语说明改为内部审查材料
    "八、来源证据摘要",  # 证据摘要不提交给代理
    "九、待确认事项",  # 待确认事项只进提交说明
]

# 固定最小可用 section 属性，供模板包缺失 section 时兜底。
DEFAULT_SECTION_XML = (  # 最小 Word section XML
    f'<w:sectPr><w:pgSz w:w="{WORD_PAGE_WIDTH}" w:h="{WORD_PAGE_HEIGHT}"/>'
    f'<w:pgMar w:top="{WORD_PAGE_MARGIN}" w:right="{WORD_PAGE_MARGIN}" '
    f'w:bottom="{WORD_PAGE_MARGIN}" w:left="{WORD_PAGE_MARGIN}" '
    f'w:header="{WORD_HEADER_FOOTER_MARGIN}" '
    f'w:footer="{WORD_HEADER_FOOTER_MARGIN}" '
    'w:gutter="0"/></w:sectPr>'
)

# 按文件路径加载共享运行时支持模块，避免导入期改写解释器搜索路径。
def load_runtime_support_module() -> Any:
    """按路径加载共享运行时支持模块。

    参数：
    - 无。

    返回：
    - `Any`：已经执行源码的共享运行时支持模块对象。

    异常：
    - 支持模块缺失或加载规格不完整时抛出 `ImportError`。
    """

    # 先根据共享模块路径创建加载规格，供后续执行源码。
    obj_runtime_spec = importlib.util.spec_from_file_location("readable_patent_runtime_support", PATH_RUNTIME_SUPPORT)  # 共享运行时支持模块加载规格

    # 在加载规格或加载器缺失时立即终止，避免主流程继续走到空模块对象。
    if obj_runtime_spec is None or obj_runtime_spec.loader is None:

        # 抛出明确阻断原因，提醒调用方先修复公共运行时支持文件。
        raise ImportError("> ERR: [Python] 无法加载 support/runtime_support.py。")

    # 根据有效加载规格创建临时模块对象，供 exec_module 写入工具函数。
    obj_runtime_module = importlib.util.module_from_spec(obj_runtime_spec)  # 已创建但尚未执行源码的运行时支持模块

    # 执行共享运行时支持模块源码，把统一路径工具装入模块对象。
    obj_runtime_spec.loader.exec_module(obj_runtime_module)

    # 把已完成加载的共享模块对象交回导出流程继续复用。
    return obj_runtime_module

# 按文件路径加载模板槽位渲染器，保证 CLI 与单元测试使用同一份正式实现。
def load_template_renderer_module() -> Any:
    """加载模板槽位渲染器模块。

    参数：
    - 无。

    返回：
    - `Any`：已执行的模板槽位渲染器模块对象。

    异常：
    - 渲染器文件缺失或无法加载时抛出 `RuntimeError`。
    """

    # 从受管导出目录定位模板渲染器，避免依赖调用方的模块搜索路径。
    obj_renderer_spec = importlib.util.spec_from_file_location(  # 模板渲染器加载规格
        "readable_patent_template_renderer",  # 模板渲染器的内部模块名
        PATH_TEMPLATE_RENDERER,  # 模板渲染器源码路径
    )

    # 加载规格或加载器缺失时无法安全执行独立渲染模块。
    if obj_renderer_spec is None or obj_renderer_spec.loader is None:

        # 用稳定错误信息阻断不完整的模板渲染链。
        raise RuntimeError("> ERR: [Python] 无法加载模板槽位渲染器。")

    # 根据已验证的加载规格创建模块对象，等待下一步执行源码。
    obj_renderer_module = importlib.util.module_from_spec(obj_renderer_spec)  # 已创建但尚未执行的模板渲染器模块

    # 执行模块定义以暴露正式模板渲染入口。
    obj_renderer_spec.loader.exec_module(obj_renderer_module)

    # 返回完成初始化的渲染器模块供导出协调层调用。
    return obj_renderer_module

# 构造导出入口的参数解析器，统一声明案件目录、输入、输出和模板参数。
def build_parser() -> argparse.ArgumentParser:
    """构造导出入口的命令行参数解析器。

    参数：
    - 无。

    返回：
    - `argparse.ArgumentParser`：已注册导出参数的解析器对象。

    异常：
    - 无。
    """

    # 先准备解析器说明文本，避免初始化行过长影响阅读。
    str_description = "Export governed disclosure markdown to DOCX."  # 导出入口命令行说明

    # 初始化导出入口解析器，后续逐项注册所有受控参数。
    obj_parser = argparse.ArgumentParser(description=str_description)  # 导出入口参数解析器

    # 注册案件目录参数，允许调用方按正式案件目录自动定位正文草稿。
    obj_parser.add_argument("--case-dir", help="Case directory containing the disclosure draft.")

    # 注册显式输入参数，允许调用方直接指定待导出的 Markdown 文件。
    obj_parser.add_argument("--input", help="Optional markdown path.")

    # 注册显式输出参数，允许调用方覆盖默认导出目录与文件名。
    obj_parser.add_argument("--output", help="Optional DOCX output path.")

    # 注册模板参数，允许调用方覆盖默认模板路径。
    obj_parser.add_argument("--template", default=str(DEFAULT_TEMPLATE), help="Optional DOCX template path.")

    # 将已经装配完成的解析器对象交给主流程继续解析参数。
    return obj_parser

# 检查当前环境是否安装 python-docx，供导出后端选择逻辑复用。
def is_python_docx_available() -> bool:
    """检查 python-docx 是否可用。

    参数：
    - 无。

    返回：
    - `bool`：模块规格可用时返回 `True`，否则返回 `False`。

    异常：
    - 无。
    """

    # 通过真实导入验证 python-docx 是否可用，避免仅有模块名却缺失二进制依赖时误判。
    try:

        # 直接导入 `Document`，确保模板导出所需核心对象真正可用。
        from docx import Document as _Document

    # 任意导入失败都视为当前解释器不具备模板 DOCX 能力。
    except Exception:

        # 返回不可用状态，交由上层决定是否切换解释器或走其他后端。
        return False

    # 返回已成功加载的结果，表示当前解释器具备模板 DOCX 导出能力。
    return _Document is not None

# 定位可选的 Codex 文档运行时 Python，供模板导出在当前解释器缺包时受控切换。
def find_codex_bundled_python() -> Path | None:
    """定位 Codex 主运行时 Python。

    参数：
    - 无。

    返回：
    - `Path | None`：存在可执行文件时返回其路径，否则返回 `None`。

    异常：
    - 无。
    """

    # 允许调用方通过环境变量覆盖默认文档运行时 Python 路径，便于受控调试。
    str_override_python = os.environ.get("CODEX_BUNDLED_PYTHON_EXE", "").strip()  # 可选覆盖的运行时 Python 路径文本

    # 在显式覆盖存在时优先使用覆盖路径。
    if str_override_python:

        # 解析覆盖路径并在文件存在时直接返回，保持人工指定优先级最高。
        path_override_python = Path(str_override_python).expanduser().resolve()  # 覆盖的运行时 Python 路径

        # 覆盖路径真实存在时直接返回。
        if path_override_python.exists():

            # 返回显式覆盖的运行时 Python 路径。
            return path_override_python

    # 默认使用 Codex 主运行时 Python 路径；缺失时返回空值交由上层降级。
    if PATH_CODEX_BUNDLED_PYTHON.exists():

        # 返回默认 Codex 文档运行时 Python。
        return PATH_CODEX_BUNDLED_PYTHON

    # 当前环境不存在可复用的 Codex 文档运行时 Python。
    return None

# 检查给定 Python 是否具备模板 DOCX 导出所需的最小依赖。
def is_bundled_template_runtime_usable(path_python: Path) -> bool:
    """检查外部 Python 是否可用于模板 DOCX 导出。

    参数：
    - `path_python`：待检查的 Python 可执行文件路径。

    返回：
    - `bool`：同时具备 `docx` 和 `PIL` 时返回 `True`。

    异常：
    - 子进程启动失败时由底层异常继续上抛。
    """

    # 先构造最小导入探针命令，只验证候选运行时能否导入关键依赖。
    list_probe_command = build_template_runtime_probe_command(path_python)  # 候选运行时依赖探针命令

    # 固定探针子进程参数，只关心最小导入是否成功，不消费额外标准输出。
    dict_probe_run_kwargs = {"capture_output": True, "text": True, "check": False}  # 依赖探针固定运行参数

    # 启动一次最小依赖探针子进程，专门判断候选运行时能否承担模板导出。
    completed_process_probe = subprocess.run(list_probe_command, **dict_probe_run_kwargs)  # 文档运行时最小依赖探针结果

    # 仅在探针命令成功时返回可用，避免把半可用环境误判成有效运行时。
    return completed_process_probe.returncode == 0

# 在当前解释器不具备模板导出能力时，尝试切换到 Codex 文档运行时继续执行当前脚本。
def maybe_reexec_with_bundled_template_runtime(path_template: Path | None) -> int | None:
    """按需切换到 Codex 文档运行时重新执行当前脚本。

    参数：
    - `path_template`：当前导出请求使用的模板路径。

    返回：
    - `int | None`：发生重启时返回子进程退出码；不需要或无法重启时返回 `None`。

    异常：
    - 子进程启动失败时由底层异常继续上抛。
    """

    # 只有严格模板导出才需要考虑切换到文档运行时。
    if path_template is None or not path_template.exists():

        # 当前请求不依赖严格模板导出，无需切换解释器。
        return None

    # 当前解释器已经能导入 python-docx 时直接继续本地执行。
    if is_python_docx_available():

        # 不需要切换到外部运行时，继续当前主流程。
        return None

    # 已经处于重启后的子进程时不再重复切换，避免递归重启。
    if os.environ.get(ENV_TEMPLATE_RUNTIME_REEXEC) == "1":

        # 返回空值，让上层抛出更明确的当前解释器能力错误。
        return None

    # 定位可复用的 Codex 文档运行时 Python；缺失时只能交由上层报错。
    path_bundled_python = find_codex_bundled_python()  # 候选文档运行时 Python 路径

    # 未找到可复用运行时时直接返回，让上层给出明确阻断信息。
    if path_bundled_python is None:

        # 返回空值表示当前环境没有可切换的文档运行时。
        return None

    # 当前解释器和候选运行时完全相同时无需重启，避免重复自调用。
    if Path(sys.executable).resolve() == path_bundled_python.resolve():

        # 返回空值，让上层继续按当前解释器直接执行。
        return None

    # 在候选运行时不具备模板导出最小依赖时直接放弃切换。
    if not is_bundled_template_runtime_usable(path_bundled_python):

        # 返回空值，让上层继续给出清晰错误。
        return None

    # 复制当前环境并写入单次重启标记，避免新进程再次递归切换。
    dict_env = os.environ.copy()  # 当前脚本重启时透传的环境变量字典

    # 标记后续子进程已经进入文档运行时，阻止再次自重启。
    dict_env[ENV_TEMPLATE_RUNTIME_REEXEC] = "1"  # 标记当前子进程已经切到文档运行时

    # 先拼出重启命令，确保文档运行时沿用当前脚本参数继续执行。
    list_reexec_command = build_template_runtime_reexec_command(path_bundled_python)  # 切换文档运行时的重启命令

    # 固定重启子进程参数，确保模板导出切换运行时后仍保留原始标准流和退出码语义。
    dict_reexec_run_kwargs = {"capture_output": True, "text": True, "check": False, "env": dict_env}  # 文档运行时重启固定参数

    # 让当前脚本在文档运行时里重跑一遍，保持 CLI 契约不变但补足 python-docx 能力。
    completed_process_reexec = subprocess.run(list_reexec_command, **dict_reexec_run_kwargs)  # 文档运行时重启结果

    # 把子进程标准输出原样回放给当前调用方，保持路径型输出契约不变。
    if completed_process_reexec.stdout:

        # 写回子进程标准输出，供 pipeline 和测试继续消费。
        sys.stdout.write(completed_process_reexec.stdout)

    # 把子进程标准错误原样回放给当前调用方，保留真实失败原因。
    if completed_process_reexec.stderr:

        # 先按受控错误头包裹外部运行时 stderr，再统一回放给当前调用方。
        sys.stderr.write(f"> ERR: [Python] 文档运行时子进程输出如下。\n{completed_process_reexec.stderr}")

    # 返回重启后子进程的退出码，让当前进程对调用方保持同态行为。
    return completed_process_reexec.returncode

# 清洗 Markdown 行内标记，避免导出到 Word 后残留源语法噪声。
def strip_markdown_inline_text(str_text: str) -> str:
    """清洗 Markdown 行内标记并返回纯文本。

    参数：
    - `str_text`：待清洗的原始 Markdown 行文本。

    返回：
    - `str`：移除图片、链接和强调标记后的纯文本。

    异常：
    - 无。
    """

    # 先把 Markdown 图片标记替换成可读占位文本，避免图示信息完全丢失。
    str_clean_text = re.sub(r"!\[([^\]]*)\]\(([^)]+)\)", r"[图片：\1]", str_text)  # 替换图片语法后的文本

    # 再把 Markdown 链接降级成文字标签，避免把链接语法原样带进 DOCX。
    str_clean_text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1", str_clean_text)  # 替换链接语法后的文本

    # 随后去掉行内代码反引号，只保留其中的实际文本内容。
    str_clean_text = re.sub(r"`([^`]+)`", r"\1", str_clean_text)  # 去掉行内代码标记后的文本

    # 再把行内公式标记降级成纯文本，避免 `$...$` 源语法直接进入 DOCX 可见正文。
    str_clean_text = RE_INLINE_FORMULA.sub(r"\1", str_clean_text)  # 去掉行内公式标记后的文本

    # 最后去掉常见强调标记，让导出正文保留可阅读文本而非 Markdown 语法。
    str_clean_text = re.sub(r"(\*\*|__)(.+?)\1", r"\2", str_clean_text)  # 去掉强调语法后的文本

    # 返回去除首尾空白后的纯文本，供后续段落构造逻辑继续使用。
    return str_clean_text.strip()

# 把单条 Markdown 表格行转换成线性文本，供最小导出路径稳定保留表格信息。
def normalize_table_row(str_line: str) -> str:
    """把单条 Markdown 表格行转换成线性文本。

    参数：
    - `str_line`：当前 Markdown 表格行原文。

    返回：
    - `str`：当前表格行的线性化文本；分隔行会返回空字符串。

    异常：
    - 无。
    """

    # 先准备单元格纯文本列表，后续逐项记录当前表格行的可读内容。
    list_cells: list[str] = []  # 当前表格行的单元格纯文本列表

    # 顺序遍历当前表格行的原始单元格文本，逐项完成行内 Markdown 清洗。
    for str_cell in str_line.strip().strip("|").split("|"):

        # 把当前单元格清洗成纯文本后收进列表，供后续判断分隔行和拼接正文。
        list_cells.append(strip_markdown_inline_text(str_cell.strip()))

    # 先记录每个单元格是否命中分隔符规则，供整行判断是否为表头分隔线复用。
    list_separator_matches = [bool(RE_TABLE_SEPARATOR.fullmatch(str_cell or "")) for str_cell in list_cells]  # 各单元格的分隔符匹配结果

    # 判断当前行是否只是 `---` 分隔行，避免把分隔符误写入正文。
    bool_is_separator = bool(list_cells) and all(list_separator_matches)  # 当前表格行是否只是分隔符

    # 在当前行仅承担表头分隔职责时直接返回空串，交给上游跳过。
    if bool_is_separator:

        # 用空串通知上游忽略当前分隔行，避免导出无意义横线。
        return ""

    # 把当前表格行拼成 `|` 分隔的线性文本，保留最小可读结构。
    str_table_text = " | ".join(str_cell for str_cell in list_cells if str_cell)  # 当前表格行的线性化文本

    # 将线性化后的表格行文本交回正文解析逻辑继续登记。
    return str_table_text

# 把单行 Markdown 正文登记成线性 block，统一处理标题、列表和普通段落。
def append_text_block(
    list_blocks: list[dict[str, Any]],
    str_line: str,
    str_stripped_line: str,
) -> None:
    """把单行 Markdown 正文登记成线性 block。

    参数：
    - `list_blocks`：当前累计的线性 block 列表。
    - `str_line`：保留原始空白的 Markdown 原始文本行。
    - `str_stripped_line`：去掉首尾空白后的 Markdown 文本行。

    返回：
    - `None`。

    异常：
    - 无。
    """

    # 先尝试按 Markdown 标题规则识别当前文本行。
    obj_heading_match = RE_HEADING.match(str_stripped_line)  # 当前文本行的标题匹配结果

    # 在命中标题时直接登记受控层级的标题 block。
    if obj_heading_match:

        # 计算当前标题可映射到 Word 的实际层级，避免超过最小样式集上限。
        int_heading_level = min(len(obj_heading_match.group(1)), WORD_HEADING_LEVEL_LIMIT)  # 当前标题映射后的 Word 层级

        # 清洗标题正文文本，避免行内 Markdown 标记残留到导出结果中。
        str_heading_text = strip_markdown_inline_text(obj_heading_match.group(2))  # 当前标题的纯文本内容

        # 把当前标题加入线性输出序列，供两个导出后端复用同一结构。
        list_blocks.append({"kind": "heading", "text": str_heading_text, "level": int_heading_level})

        # 结束当前函数，避免后续列表和段落分支重复处理标题行。
        return

    # 再按项目符号规则识别当前文本行，覆盖 `-` 和 `*` 样式列表。
    obj_bullet_match = RE_BULLET.match(str_line)  # 当前文本行的无序列表匹配结果

    # 在命中无序列表时登记带项目符号前缀的普通段落 block。
    if obj_bullet_match:

        # 清洗无序列表正文，并补上统一项目符号前缀。
        str_bullet_text = "• " + strip_markdown_inline_text(obj_bullet_match.group(1))  # 当前无序列表的线性化正文

        # 把当前项目符号条目加入线性输出序列。
        list_blocks.append({"kind": "paragraph", "text": str_bullet_text})

        # 结束当前函数，避免后续分支重复处理同一条列表项。
        return

    # 若前面都未命中，这里再判断是否属于带编号的步骤条目。
    obj_ordered_match = RE_ORDERED.match(str_line)  # “1.” 样式步骤条目的匹配结果

    # 在命中有序列表时登记保留序号正文的普通段落 block。
    if obj_ordered_match:

        # 清洗编号列表正文，并保留 `1.` 风格前缀以维持阅读顺序。
        str_ordered_text = strip_markdown_inline_text(str_stripped_line)  # 当前编号列表的线性化正文

        # 把当前编号条目加入线性输出序列。
        list_blocks.append({"kind": "paragraph", "text": str_ordered_text})

        # 结束当前函数，避免后续普通段落逻辑再次写入同一行。
        return

    # 将剩余普通文本行清洗成纯文本段落，供导出正文直接写入。
    str_paragraph_text = strip_markdown_inline_text(str_stripped_line)  # 当前普通正文行的纯文本内容

    # 把当前普通段落加入线性输出序列，形成最终 DOCX 正文内容。
    list_blocks.append({"kind": "paragraph", "text": str_paragraph_text})

# 把已闭合的 fenced code block 转成正文占位和附件正文，避免正文直接嵌入大段源码。
def flush_code_block(
    list_blocks: list[dict[str, Any]],
    list_attachments: list[str],
    list_code_lines: list[str],
) -> None:
    """把已闭合的 fenced code block 转成正文占位和附件正文。

    参数：
    - `list_blocks`：当前累计的线性 block 列表。
    - `list_attachments`：当前累计的附件正文列表。
    - `list_code_lines`：当前代码块收集到的原始正文行列表。

    返回：
    - `None`。

    异常：
    - 无。
    """

    # 把代码块正文拼成单段附件文本，供文末附件章节集中收纳。
    str_attachment_text = "\n".join(list_code_lines).strip()  # 当前代码块的附件正文文本

    # 仅在代码块正文非空时才把它记入附件列表。
    if str_attachment_text:

        # 将当前代码块正文收进附件列表，保持与正文出现顺序一致。
        list_attachments.append(str_attachment_text)

    # 在正文当前位置写入附件占位提示，提醒评审人后续人工替换。
    list_blocks.append({"kind": "paragraph", "text": TEXT_ATTACHMENT_PLACEHOLDER})

# 把附件正文统一展开成文末章节，确保两个导出后端看到相同的附件结构。
def append_attachment_blocks(
    list_blocks: list[dict[str, Any]],
    list_attachments: list[str],
) -> None:
    """把附件正文统一展开成文末章节。

    参数：
    - `list_blocks`：当前累计的线性 block 列表。
    - `list_attachments`：正文中摘出的附件正文列表。

    返回：
    - `None`。

    异常：
    - 无。
    """

    # 只有在存在附件正文时才需要追加分页和附件章节。
    if not list_attachments:

        # 在没有附件内容时提前返回，避免正文尾部平白新增附件标题。
        return

    # 先加入分页 block，把附件章节与正文主体视觉分隔开。
    list_blocks.append({"kind": "page_break"})

    # 再写入附件章节标题，帮助评审人快速定位需要人工处理的内容。
    list_blocks.append({"kind": "heading", "text": TEXT_ATTACHMENT_TITLE, "level": 1})

    # 按原正文出现顺序逐项展开附件正文，保持回看路径稳定。
    for int_attachment_index, str_attachment_text in enumerate(list_attachments, start=1):

        # 先登记当前附件的小标题，便于文末逐条核对代码块来源。
        list_blocks.append({"kind": "paragraph", "text": f"附件 {int_attachment_index}"})

        # 再按行展开当前附件正文，避免超长整段影响 Word 可读性。
        for str_attachment_line in str_attachment_text.splitlines():

            # 只在当前附件行有实际内容时才把它写入文末章节。
            if str_attachment_line.strip():

                # 截断过长附件行，避免单行源码把页面版式横向撑开。
                str_trimmed_attachment_line = str_attachment_line[:600]  # 当前附件行的受控文本

                # 把当前附件行作为普通段落加入附件章节。
                list_blocks.append({"kind": "paragraph", "text": str_trimmed_attachment_line})

# 统一把 Markdown 解析成线性 block 列表，供两个导出后端共用同一份正文结构。
def collect_markdown_blocks(str_markdown: str) -> list[dict[str, Any]]:
    """把 Markdown 解析成线性 block 列表。

    参数：
    - `str_markdown`：待导出的 Markdown 全文。

    返回：
    - `list[MarkdownBlock]`：按正文顺序整理好的线性 block 列表。

    异常：
    - 无。
    """

    # 先准备线性 block 列表，后续会按正文顺序逐项登记内容。
    list_blocks: list[dict[str, Any]] = []  # Markdown 线性 block 列表

    # 把 Markdown 全文按行拆开，便于顺序扫描标题、表格和代码块。
    list_lines = str_markdown.splitlines()  # Markdown 原始行列表

    # 用行游标控制顺序扫描位置，避免复杂递归解析。
    int_index = 0  # 当前顺序扫描到的 Markdown 行号

    # 用状态位记录当前是否位于 fenced code block 内部。
    bool_in_code_block = False  # 当前是否位于代码块内部

    # 用列表暂存当前代码块正文行，待闭合后统一转入附件章节。
    list_code_lines: list[str] = []  # 当前代码块正文暂存列表

    # 用列表暂存所有附件正文，供扫描结束后集中追加到文末。
    list_attachments: list[str] = []  # 正文中摘出的附件正文列表

    # 顺序扫描 Markdown 各行，持续构造受控的线性正文结构。
    while int_index < len(list_lines):

        # 读取当前原始文本行，供后续判断标题、表格和代码块边界。
        str_line = list_lines[int_index]  # 当前原始 Markdown 行

        # 读取当前去首尾空白后的文本行，便于统一判断空行和 fenced code 标记。
        str_stripped_line = str_line.strip()  # 当前去首尾空白后的 Markdown 行

        # 在命中 fenced code block 边界时切换代码块状态。
        if str_stripped_line.startswith("```"):

            # 在代码块闭合点把已收集正文转入附件并写入正文占位。
            if bool_in_code_block:

                # 把当前已闭合代码块刷新成正文占位和附件正文。
                flush_code_block(list_blocks, list_attachments, list_code_lines)

            # 在遇到新的代码块起始边界时清空代码块正文暂存列表。
            else:

                # 为新的代码块正文重新准备暂存列表，避免污染上一段附件内容。
                list_code_lines = []  # 新代码块的正文暂存列表

            # 翻转代码块状态，让下一轮扫描按新状态继续解析正文。
            bool_in_code_block = not bool_in_code_block  # 切换后的代码块状态

            # 跳过当前 fenced code 边界行，继续扫描下一行正文。
            int_index += 1  # 跳过代码块边界后的下一行游标位置

            # 直接进入下一轮扫描，避免边界行继续落入普通正文分支。
            continue

        # 在代码块内部时仅累计原始文本行，不做标题和列表识别。
        if bool_in_code_block:

            # 把当前代码行加入暂存列表，等待代码块闭合后统一写入附件。
            list_code_lines.append(str_line)

            # 推进行游标到下一行，继续读取代码块正文。
            int_index += 1  # 代码块正文继续扫描后的下一行游标位置

            # 直接进入下一轮扫描，避免代码行被误判为普通正文。
            continue

        # 在普通空白行场景下仅跳过当前行，不生成额外段落。
        if not str_stripped_line:

            # 推进行游标到下一行，保持原始空白只承担段落分隔职责。
            int_index += 1  # 空白行跳过后的下一行游标位置

            # 直接进入下一轮扫描，避免空白行进入正文列表。
            continue

        # 在命中 Markdown 表格时收集连续表格行，并按最小线性格式写入正文。
        if str_stripped_line.startswith("|"):

            # 先准备当前表格段的原始行列表，供后续统一线性化处理。
            list_table_lines: list[str] = []  # 当前连续表格段的原始行列表

            # 持续收集连续表格行，直到遇到非表格行才结束。
            while int_index < len(list_lines) and list_lines[int_index].strip().startswith("|"):

                # 把当前表格原始行加入暂存列表，等待统一线性化。
                list_table_lines.append(list_lines[int_index])

                # 推进行游标到下一行，继续判断是否仍属于同一表格段。
                int_index += 1  # 表格段继续收集后的下一行游标位置

            # 逐条把当前表格段线性化成可阅读的普通段落文本。
            for str_table_line in list_table_lines:

                # 把当前表格行转换成线性文本，必要时过滤纯分隔行。
                str_table_text = normalize_table_row(str_table_line)  # 当前表格行转换后的线性正文

                # 只在当前表格行存在实际正文时才把它写入线性 block 列表。
                if str_table_text:

                    # 把当前表格行作为普通段落登记，保证回退导出也能阅读表格内容。
                    list_blocks.append({"kind": "paragraph", "text": str_table_text})

            # 直接进入下一轮扫描，避免表格段首行再次落入普通正文分支。
            continue

        # 把当前普通文本行登记成标题、列表或普通段落 block。
        append_text_block(list_blocks, str_line, str_stripped_line)

        # 推进行游标到下一行，继续扫描后续 Markdown 内容。
        int_index += 1  # 普通正文处理后的下一行游标位置

    # 在扫描结束时兜底处理未闭合代码块，避免最后一段源码完全丢失。
    if bool_in_code_block:

        # 把末尾未闭合代码块也转入正文占位和附件章节。
        flush_code_block(list_blocks, list_attachments, list_code_lines)

    # 把附件正文统一展开到文末，保持附件结构和顺序稳定。
    append_attachment_blocks(list_blocks, list_attachments)

    # 将完整线性 block 列表交给导出后端继续写入 DOCX。
    return list_blocks

# 去掉 Markdown 标题标记并返回清洗后的标题正文。
def normalize_markdown_heading(str_line: str) -> str:
    """归一化 Markdown 标题文本。

    参数：
    - `str_line`：原始 Markdown 行文本。

    返回：
    - `str`：去掉标题井号和行内标记后的标题正文。

    异常：
    - 无。
    """

    # 先尝试按 Markdown 标题规则识别当前行。
    obj_heading_match = RE_HEADING.match(str_line.strip())  # 当前行的 Markdown 标题匹配结果

    # 不是标题时返回空字符串，供调用方继续按普通正文处理。
    if obj_heading_match is None:

        # 空字符串表示当前行不参与章节切换。
        return ""

    # 清洗标题正文，避免行内 Markdown 标记影响模板章节匹配。
    str_heading_text = strip_markdown_inline_text(obj_heading_match.group(2))  # 当前标题正文文本

    # 返回清洗后的标题正文，供章节归一化逻辑继续判断。
    return str_heading_text

# 把 Markdown 标题映射到正式模板章节名。
def resolve_template_heading(str_heading_text: str) -> str:
    """把 Markdown 标题映射到模板章节名。

    参数：
    - `str_heading_text`：已经清洗过的 Markdown 标题正文。

    返回：
    - `str`：匹配到的模板章节名；无法匹配时返回空字符串。

    异常：
    - 无。
    """

    # 标题完全命中映射时直接返回模板章节名。
    if str_heading_text in TEMPLATE_HEADING_ALIASES:

        # 返回完全匹配得到的模板章节名。
        return TEMPLATE_HEADING_ALIASES[str_heading_text]

    # 处理旧草稿中带空格或半角冒号差异的标题文本。
    str_compact_heading = str_heading_text.replace(" ", "").replace(":", "：")  # 去空格后的标题文本

    # 逐项比较归一化后的标题，兼容旧标题里的轻微格式差异。
    for str_source_heading, str_target_heading in TEMPLATE_HEADING_ALIASES.items():

        # 准备映射源标题的紧凑形式，便于和当前标题比较。
        str_compact_source = str_source_heading.replace(" ", "").replace(":", "：")  # 映射源标题紧凑文本

        # 当前标题和映射源标题紧凑形式相同时返回模板标题。
        if str_compact_heading == str_compact_source:

            # 返回归一化后的模板章节名。
            return str_target_heading

    # 无法匹配模板章节时返回空字符串，让调用方决定是否进入 sidecar。
    return ""

# 判断当前标题是否属于内部审查章节。
def is_internal_heading(str_heading_text: str) -> bool:
    """判断当前标题是否属于内部审查章节。

    参数：
    - `str_heading_text`：已经清洗过的 Markdown 标题正文。

    返回：
    - `bool`：标题属于内部章节时返回 `True`。

    异常：
    - 无。
    """

    # 返回标题是否以受控内部章节前缀开头。
    return str_heading_text.startswith(INTERNAL_SECTION_PREFIXES)

# 把导出 Markdown 拆成模板正文分节和内部 sidecar 文本。
def collect_template_sections(str_markdown: str) -> dict[str, Any]:
    """收集模板正文分节和内部说明。

    参数：
    - `str_markdown`：待导出的 Markdown 全文。

    返回：
    - `dict[str, Any]`：包含模板分节和内部说明行的结构化结果。

    异常：
    - 无。
    """

    # 为每个模板章节准备正文列表，保持最终 DOCX 章节顺序稳定。
    dict_sections: dict[str, list[str]] = {str_heading: [] for str_heading in TEMPLATE_SECTION_ORDER}  # 模板章节正文映射

    # 准备用于 sidecar 的内部说明行列表，承接七至九节等内部内容。
    list_internal_lines: list[str] = []  # sidecar 暂存的内部正文行

    # 当前模板章节名为空时，普通正文会被忽略或转入 sidecar。
    str_current_section = ""  # 当前正在收集的模板章节名

    # 当前是否处于内部审查章节，用来决定后续正文是否进入 sidecar。
    bool_collecting_internal = False  # 是否正在收集内部审查章节

    # 顺序扫描 Markdown 行，按标题边界收集正文。
    for str_line in str_markdown.splitlines():

        # 提取当前行的 Markdown 标题正文；非标题时为空字符串。
        str_heading_text = normalize_markdown_heading(str_line)  # 当前行标题正文

        # 标题行会切换当前正文收集目标。
        if str_heading_text:

            # 内部标题会改写后续正文路由，确保其不进入主 DOCX。
            bool_collecting_internal = is_internal_heading(str_heading_text)  # 当前标题的 sidecar 路由标记

            # 在内部章节命中时记录标题并停止写入主 DOCX。
            if bool_collecting_internal:

                # 将内部章节标题写入 sidecar，便于人工回看被移出的内容。
                list_internal_lines.append(f"## {str_heading_text}")

                # 清空主 DOCX 章节目标，避免内部正文继续进入主交付件。
                str_current_section = ""  # 主 DOCX 章节目标清空

                # 当前标题已经处理完毕，继续扫描下一行。
                continue

            # 尝试把当前标题映射到模板章节名。
            str_mapped_heading = resolve_template_heading(str_heading_text)  # 当前标题对应的模板章节名

            # 在命中模板章节时切换主 DOCX 收集目标。
            if str_mapped_heading:

                # 记录当前主 DOCX 章节目标，后续正文写入该章节。
                str_current_section = str_mapped_heading  # 当前模板章节名

                # 当前标题只承担分节边界职责，不作为正文重复写入。
                continue

            # 未命中模板且非内部章节时清空当前章节，避免误写未知标题正文。
            str_current_section = ""  # 未识别标题后的主 DOCX 章节目标

            # 未识别标题不进入主文档，继续扫描后续行。
            continue

        # 清洗当前普通行，供正文和 sidecar 统一使用。
        str_clean_line = strip_markdown_inline_text(str_line.strip())  # 当前普通行清洗文本

        # 空行只承担段落分隔职责，不进入 DOCX 或 sidecar。
        if not str_clean_line:

            # 跳过空行，保持最终文档内容紧凑。
            continue

        # 内部章节正文进入提交说明 sidecar。
        if bool_collecting_internal:

            # 追加内部说明正文，避免其进入最终 DOCX 主体。
            list_internal_lines.append(str_clean_line)

            # 当前行已进入 sidecar，继续扫描下一行。
            continue

        # 已定位模板章节时，把当前正文写入对应章节。
        if str_current_section:

            # 把正文追加到当前模板章节，供后续按模板顺序输出。
            dict_sections[str_current_section].append(str_clean_line)

    # 返回模板分节和内部说明，供 DOCX 渲染与 sidecar 渲染共同使用。
    return {"sections": dict_sections, "internal_lines": list_internal_lines}

# 把 Markdown 拆成可提交模板章节的段落块和公式块，供代理交付版 DOCX 真正嵌入公式。
def collect_template_section_blocks(str_markdown: str) -> dict[str, Any]:
    """收集模板章节的结构化段落块和公式块。

    参数：
    - `str_markdown`：待导出的 Markdown 全文。

    返回：
    - `dict[str, Any]`：包含模板章节块列表和内部说明行的结构化结果。

    异常：
    - 无。
    """

    # 为每个模板章节准备结构化块列表，后续按顺序登记段落或公式块。
    dict_sections: dict[str, list[dict[str, str]]] = {str_heading: [] for str_heading in TEMPLATE_SECTION_ORDER}  # 模板章节块映射

    # 为内部审查 sidecar 暂存正文行，避免这些内容误入正式交付主稿。
    list_internal_lines: list[str] = []  # 内部审查说明行列表

    # 用空字符串表示扫描器尚未命中正式模板章节标题。
    str_current_section = ""  # 当前命中的正式模板章节标题

    # 单独跟踪是否进入内部审查段，供普通正文决定写入主稿还是 sidecar。
    bool_collecting_internal = False  # 内部审查区域状态位

    # 当前是否位于 display-math 公式块内部。
    bool_in_formula_block = False  # 是否正在收集公式块

    # 暂存当前公式块的原始正文行，待闭合后统一写入结构化结果。
    list_formula_lines: list[str] = []  # 当前公式块正文行列表

    # 顺序扫描 Markdown 各行，按章节边界和公式块边界收集结构化内容。
    for str_line in str_markdown.splitlines():

        # 读取当前去首尾空白后的文本，供标题、空行和公式块边界判断复用。
        str_stripped_line = str_line.strip()  # 当前去首尾空白后的文本行

        # 在 display-math 边界行上切换公式块状态。
        if str_stripped_line == "$$":

            # 命中公式块闭合边界时，把已收集公式正文写回当前模板章节。
            if bool_in_formula_block:

                # 先把公式块正文规整成非空纯文本行列表，避免空白行进入最终公式块。
                list_clean_formula_lines = collect_nonempty_stripped_lines(list_formula_lines)  # 当前公式块的规整正文行列表

                # 再把规整后的公式正文行按换行重新拼接，形成稳定的公式块文本。
                str_formula_text = "\n".join(list_clean_formula_lines)  # 当前公式块正文文本

                # 只在当前模板章节存在且公式正文非空时才把公式登记到主稿章节。
                if str_current_section and str_formula_text and not bool_collecting_internal:

                    # 把当前公式块登记到模板章节，供 DOCX 导出阶段渲染成嵌入式图片。
                    dict_sections[str_current_section].append({"kind": "formula", "text": str_formula_text})

                # 公式块正文已消费后清空暂存列表，避免污染后续公式块。
                list_formula_lines = []  # 已清空的公式块正文行列表

            # 切换公式块状态，让扫描器继续处理后续内容。
            bool_in_formula_block = not bool_in_formula_block  # 切换后的公式块状态

            # display-math 边界行本身不进入主稿正文，直接处理下一行。
            continue

        # 在公式块内部时只累计公式正文，不再执行标题或普通段落识别。
        if bool_in_formula_block:

            # 把当前公式正文行加入暂存列表，等待闭合边界统一提交。
            list_formula_lines.append(str_line)

            # 当前行已经作为公式正文处理，继续扫描下一行。
            continue

        # 先归一化当前行标题文本，后续再决定它指向主稿章节还是内部说明。
        str_heading_text = normalize_markdown_heading(str_line)  # 归一化后的标题候选文本

        # 在命中标题时重新决定后续普通行的路由目标。
        if str_heading_text:

            # 先判断当前标题是否属于内部章节。
            bool_collecting_internal = is_internal_heading(str_heading_text)  # 当前标题的内部章节标记

            # 内部标题只进入内部说明，不作为主稿章节继续扩展。
            if bool_collecting_internal:

                # 把内部标题写入内部说明，保留人工回看路径。
                list_internal_lines.append(f"## {str_heading_text}")

                # 清空当前模板章节，避免内部正文误写回主稿。
                str_current_section = ""  # 已清空的当前模板章节名

                # 当前标题已处理完成，继续扫描下一行。
                continue

            # 把当前标题映射到模板章节锚点，无法识别时保持后续普通行不进入主稿。
            str_current_section = resolve_template_heading(str_heading_text)  # 标题命中的模板章节锚点

            # 当前标题只承担章节切换职责，不作为正文再次写入。
            continue

        # 把当前普通行先清成纯文本，再决定它是否进入正式交底书主稿。
        str_clean_line = strip_markdown_inline_text(str_stripped_line)  # 去掉 Markdown 标记后的正文文本

        # 空行只承担段落分隔职责，不进入主稿或内部说明。
        if not str_clean_line:

            # 跳过空行，保持主稿段落列表紧凑。
            continue

        # 内部章节正文统一进入内部说明，不进入正式交付主稿。
        if bool_collecting_internal:

            # 把当前内部说明行追加到内部说明列表。
            list_internal_lines.append(str_clean_line)

            # 当前行已经写入内部说明，继续扫描下一行。
            continue

        # 只有在已经定位到模板章节时才把当前普通行写入主稿章节。
        if str_current_section:

            # 把当前普通段落写入主稿章节，供 DOCX 导出逐段写入。
            dict_sections[str_current_section].append({"kind": "paragraph", "text": str_clean_line})

    # 返回模板章节结构化块和内部说明，供严格模板导出复用。
    return {"sections": dict_sections, "internal_lines": list_internal_lines}

# 统计模板章节块中的公式数量，供最终 DOCX 媒体数量校验复用。
def count_formula_blocks(dict_sections: dict[str, list[dict[str, str]]]) -> int:
    """统计模板章节块中的公式数量。

    参数：
    - `dict_sections`：模板章节到结构化块列表的映射。

    返回：
    - `int`：所有章节中公式块的总数。

    异常：
    - 无。
    """

    # 逐章节汇总 `formula` 类型块数量，供 DOCX 媒体校验复用。
    return sum(
        1
        for list_blocks in dict_sections.values()
        for dict_block in list_blocks
        if dict_block.get("kind") == "formula"
    )

# 收集当前案件可用于 DOCX 正文嵌图的 PNG 附图路径。
def collect_delivery_figure_image_paths(path_case_dir: Path | None) -> list[Path]:
    """收集可用于 DOCX 嵌图的 PNG 附图路径。

    参数：
    - `path_case_dir`：当前案件目录路径；为空时返回空列表。

    返回：
    - `list[Path]`：按稳定顺序排列的 PNG 附图路径列表。

    异常：
    - 无。
    """

    # 缺少案件目录时无法定位正式附图目录，直接返回空列表。
    if path_case_dir is None:

        # 空列表表示当前导出不具备可用附图资产。
        return []

    # 固定正式附图目录路径，保持与案件目录合同一致。
    path_figures_dir = path_case_dir / "05_figures"  # 正式附图目录路径

    # 在附图目录缺失时直接返回空列表，让导出器只输出正文与公式。
    if not path_figures_dir.exists():

        # 空列表表示当前案件尚未生成可嵌入 DOCX 的附图资产。
        return []

    # 先按默认产品承诺的两张正式附图顺序收集 PNG 资产。
    list_default_paths = [  # 默认正式附图 PNG 路径列表
        path_figures_dir / "图1_方法流程图.png",  # 方法流程图的 PNG 交付路径
        path_figures_dir / "图2_系统模块图.png",  # 系统模块图的 PNG 交付路径
    ]

    # 仅保留已经真实落盘的默认 PNG 附图，避免把不存在的文件写进返回结果。
    list_existing_default_paths = [path_item for path_item in list_default_paths if path_item.exists()]  # 已存在的默认 PNG 附图路径列表

    # 在默认两张附图已经存在时直接按固定顺序返回。
    if list_existing_default_paths:

        # 返回默认附图路径列表，保证 DOCX 正文中的嵌图顺序稳定。
        return list_existing_default_paths

    # 默认文件名缺失时回退到扫描全部 `图*.png` 文件，兼容后续扩展附图场景。
    return sorted(path_figures_dir.glob("图*.png"))

# 根据章节层级选择 Word 标题样式。
def resolve_heading_style(str_heading: str) -> str:
    """根据章节标题选择 Word 样式。

    参数：
    - `str_heading`：模板章节标题。

    返回：
    - `str`：对应的 Word 段落样式 ID。

    异常：
    - 无。
    """

    # 数字编号小节使用更低层级标题样式，减少版式跳跃。
    if str_heading[0].isdigit():

        # 返回三级标题样式，供 3.x 和 4.x 小节使用。
        return "Heading3"

    # 中文大章节使用一级标题样式，突出六大模板部分。
    return "Heading1"

# 从模板 document.xml 中提取第一个信息表 XML。
def extract_first_table_xml(str_document_xml: str) -> str:
    """从模板主文档 XML 中提取首个表格。

    参数：
    - `str_document_xml`：模板 `word/document.xml` 文本。

    返回：
    - `str`：第一个 `w:tbl` 表格 XML；缺失时返回空字符串。

    异常：
    - 无。
    """

    # 用非贪婪匹配提取第一个 Word 表格，保留模板原始行政信息表结构。
    obj_table_match = re.search(r"<w:tbl\b.*?</w:tbl>", str_document_xml, flags=re.DOTALL)  # 模板首个表格匹配结果

    # 表格缺失时返回空字符串，让调用方继续生成正文但不伪造表格。
    if obj_table_match is None:

        # 空字符串表示模板里没有可复用的信息表。
        return ""

    # 返回模板首个表格 XML，供新主文档正文直接复用。
    return obj_table_match.group(0)

# 从模板 document.xml 中提取 section 属性。
def extract_section_xml(str_document_xml: str) -> str:
    """从模板主文档 XML 中提取 section 属性。

    参数：
    - `str_document_xml`：模板 `word/document.xml` 文本。

    返回：
    - `str`：模板 `w:sectPr` XML；缺失时返回最小兜底 section。

    异常：
    - 无。
    """

    # 匹配模板文档最后的 section 属性，尽量保留页面设置。
    obj_section_match = re.search(r"<w:sectPr\b.*?</w:sectPr>", str_document_xml, flags=re.DOTALL)  # 模板 section 匹配结果

    # 模板提供 section 时直接复用。
    if obj_section_match is not None:

        # 返回模板 section XML，保留页面尺寸和页边距。
        return obj_section_match.group(0)

    # 模板缺失 section 时回退到最小 A4 页面设置。
    return DEFAULT_SECTION_XML

# 从模板 document.xml 中提取根文档开始标签，保留全部命名空间声明。
def extract_document_open_xml(str_document_xml: str) -> str:
    """从模板主文档 XML 中提取 `w:document` 开始标签。

    参数：
    - `str_document_xml`：模板 `word/document.xml` 文本。

    返回：
    - `str`：包含命名空间声明的 `w:document` 开始标签。

    异常：
    - 模板主文档缺少根标签时抛出 `ValueError`。
    """

    # 从模板中保留完整根标签，避免复制表格后出现未绑定命名空间前缀。
    obj_document_match = re.search(r"<w:document\b[^>]*>", str_document_xml)  # 模板根文档开始标签匹配结果

    # 模板主文档缺少根标签时直接阻断，避免写出非法 DOCX。
    if obj_document_match is None:

        # 抛出明确错误，提示模板资产本身结构异常。
        raise ValueError("> ERR: [Python] 模板 DOCX 主文档缺少 w:document 根标签。")

    # 返回包含全部命名空间声明的根标签，供新 document.xml 复用。
    return obj_document_match.group(0)

# 把模板正文分节渲染成 Word body XML。
def render_template_body_xml(
    str_table_xml: str,
    dict_sections: dict[str, list[str]],
    str_section_xml: str,
) -> str:
    """把模板正文分节渲染成 Word body XML。

    参数：
    - `str_table_xml`：从模板复用的信息表 XML。
    - `dict_sections`：按模板章节归集的正文映射。
    - `str_section_xml`：模板或兜底 section XML。

    返回：
    - `str`：完整 `w:body` 内部 XML。

    异常：
    - 无。
    """

    # 先准备 body 片段列表，首段优先放入模板信息表。
    list_body_parts: list[str] = []  # Word body XML 片段列表

    # 在模板表格存在时保留到最终主文档。
    if str_table_xml:

        # 写入模板信息表，保证行政信息行和勾选项继续可见。
        list_body_parts.append(str_table_xml)

    # 按模板顺序写入章节标题和正文段落。
    for str_heading in TEMPLATE_SECTION_ORDER:

        # 选择当前章节对应的 Word 标题样式。
        str_style_id = resolve_heading_style(str_heading)  # 当前章节标题样式

        # 写入当前模板章节标题。
        list_body_parts.append(render_word_paragraph_xml(str_heading, str_style_id))

        # 读取当前章节正文行；空章节只保留标题，交由后续校验阻断交付。
        list_section_lines = dict_sections.get(str_heading, [])  # 当前章节正文行列表

        # 逐条写入当前章节正文段落。
        for str_section_line in list_section_lines:

            # 写入当前章节正文段落。
            list_body_parts.append(render_word_paragraph_xml(str_section_line))

    # 写入 section 属性，保持 Word 主文档结构完整。
    list_body_parts.append(str_section_xml)

    # 返回完整 body 内部 XML，供模板包替换 document.xml。
    return "".join(list_body_parts)

# 读取 DOCX 主文档 XML，供模板校验解析表格、标题和正文文本。
def read_docx_document_xml(path_docx: Path) -> str:
    """读取 DOCX 主文档 XML。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。

    返回：
    - `str`：`word/document.xml` 的 UTF-8 文本。

    异常：
    - DOCX 包缺失或主文档缺失时由底层异常继续上抛。
    """

    # 打开 DOCX ZIP 包并读取 Word 主文档 XML，避免依赖 python-docx。
    with zipfile.ZipFile(path_docx, "r") as obj_docx_zip:

        # 解码主文档 XML 文本，供后续 ElementTree 解析和字符串扫描复用。
        str_document_xml = obj_docx_zip.read("word/document.xml").decode("utf-8")  # 待校验的 document.xml 文本

    # 返回主文档 XML 文本，供模板校验继续处理。
    return str_document_xml

# 读取 DOCX ZIP 包中的媒体条目，供嵌图和嵌公式校验复用。
def read_docx_media_entries(path_docx: Path) -> list[str]:
    """读取 DOCX 媒体条目列表。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。

    返回：
    - `list[str]`：DOCX ZIP 包内 `word/media/` 条目路径列表。

    异常：
    - DOCX 包读取失败时由底层异常继续上抛。
    """

    # 打开 DOCX ZIP 包并枚举媒体条目，判断图像是否真实进入最终交付件。
    with zipfile.ZipFile(path_docx, "r") as obj_docx_zip:

        # 仅保留媒体目录条目，供最终媒体数量和存在性校验复用。
        return [str_name for str_name in obj_docx_zip.namelist() if str_name.startswith("word/media/")]

# 从主文档 XML 中提取可见段落文本。
def extract_docx_paragraph_texts(str_document_xml: str) -> list[str]:
    """从 DOCX 主文档 XML 中提取段落文本。

    参数：
    - `str_document_xml`：`word/document.xml` 文本。

    返回：
    - `list[str]`：按文档顺序提取的段落可见文本。

    异常：
    - XML 非法时由 `ElementTree` 异常继续上抛。
    """

    # 解析主文档 XML，供命名空间查询所有段落和文本节点。
    obj_document_root = ElementTree.fromstring(str_document_xml)  # DOCX 主文档 XML 根节点

    # 固定 WordprocessingML 命名空间，供段落和文本查询共享。
    dict_namespaces = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}  # Word XML 命名空间

    # 先准备段落文本列表，后续按 Word 段落顺序逐项填充。
    list_paragraph_texts: list[str] = []  # DOCX 可见段落文本列表

    # 遍历所有段落节点，并合并段落内可能被拆分的多个文本 run。
    for obj_paragraph in obj_document_root.findall(".//w:p", dict_namespaces):

        # 拼接当前段落内所有 w:t 文本节点，兼容 Word 按 run 拆分文字。
        str_paragraph_text = "".join(  # 当前段落可见文本
            obj_text_node.text or ""  # 当前文本节点内容
            for obj_text_node in obj_paragraph.findall(".//w:t", dict_namespaces)  # 当前段落的文本节点
        )

        # 把当前段落文本加入列表，后续用于章节顺序和正文非空校验。
        list_paragraph_texts.append(str_paragraph_text.strip())

    # 返回完整段落文本列表，供模板 DOCX 校验继续判断。
    return list_paragraph_texts

# 在段落序列中定位所有模板章节标题。
def collect_heading_indices(list_paragraph_texts: list[str]) -> dict[str, int]:
    """收集模板章节标题所在段落索引。

    参数：
    - `list_paragraph_texts`：按文档顺序提取的段落可见文本。

    返回：
    - `dict[str, int]`：模板标题到首个段落索引的映射。

    异常：
    - 无。
    """

    # 先准备标题索引映射，缺失标题不会进入该字典。
    dict_heading_indices: dict[str, int] = {}  # 模板章节标题索引映射

    # 顺序扫描段落文本，记录模板标题第一次出现的位置。
    for int_index, str_paragraph_text in enumerate(list_paragraph_texts):

        # 只记录模板章节顺序中的精确标题，避免误把普通正文当成章节边界。
        if str_paragraph_text in TEMPLATE_SECTION_ORDER and str_paragraph_text not in dict_heading_indices:

            # 记录当前标题所在段落索引，供正文非空检查使用。
            dict_heading_indices[str_paragraph_text] = int_index  # 当前模板标题的首个段落位置

    # 返回标题索引映射，供后续校验复用。
    return dict_heading_indices

# 提取某个模板章节标题到下一个标题之间的正文段落。
def collect_section_body_texts(
    str_heading: str,
    list_paragraph_texts: list[str],
    dict_heading_indices: dict[str, int],
) -> list[str]:
    """提取指定模板章节的正文段落。

    参数：
    - `str_heading`：待检查的模板章节标题。
    - `list_paragraph_texts`：DOCX 可见段落文本列表。
    - `dict_heading_indices`：模板标题到段落索引的映射。

    返回：
    - `list[str]`：当前章节标题之后、下一模板标题之前的非空正文段落。

    异常：
    - 无。
    """

    # 当前标题缺失时返回空列表，缺失问题由上层标题校验统一报告。
    if str_heading not in dict_heading_indices:

        # 空列表表示无法提取当前章节正文。
        return []

    # 读取当前标题在模板顺序中的位置，便于寻找下一个实际出现的章节标题。
    int_template_index = TEMPLATE_SECTION_ORDER.index(str_heading)  # 当前标题在模板顺序中的位置

    # 当前章节正文从标题后一段开始。
    int_start_index = dict_heading_indices[str_heading] + 1  # 当前章节正文起始段落索引

    # 默认把文档尾部作为当前章节结束位置，若后续有标题再收窄。
    int_end_index = len(list_paragraph_texts)  # 当前章节正文结束段落索引

    # 从模板顺序的后续标题中查找第一个真实出现的标题索引。
    for str_next_heading in TEMPLATE_SECTION_ORDER[int_template_index + 1:]:

        # 命中下一个已出现标题时，用它的位置作为当前章节结束边界。
        if str_next_heading in dict_heading_indices:

            # 记录下一标题索引并结束查找。
            int_end_index = dict_heading_indices[str_next_heading]  # 当前章节的右侧边界位置

            # 找到最近的后续模板标题后即可结束边界搜索。
            break

    # 返回当前章节边界内的非空、非标题正文段落。
    return [
        str_text  # 当前正文段落文本
        for str_text in list_paragraph_texts[int_start_index:int_end_index]  # 当前章节范围内的段落
        if str_text and str_text not in TEMPLATE_SECTION_ORDER  # 过滤空行和标题行
    ]

# 收集严格模板 DOCX 的结构与正文质量问题。
def collect_template_docx_findings(
    path_docx: Path,
    int_expected_media_count: int = 0,
) -> list[str]:
    """收集严格模板 DOCX 校验问题。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。
    - `int_expected_media_count`：按导出输入应至少嵌入的媒体数量。

    返回：
    - `list[str]`：校验问题文本列表；为空表示通过。

    异常：
    - DOCX 包读取或 XML 解析失败时由底层异常继续上抛。
    """

    # 读取最终 DOCX 的主文档 XML，供表格和正文检查共享。
    str_document_xml = read_docx_document_xml(path_docx)  # 待扫描的最终 document.xml 文本

    # 提取段落级可见文本，方便判断标题顺序与章节正文是否为空。
    list_paragraph_texts = extract_docx_paragraph_texts(str_document_xml)  # 用于章节校验的段落文本

    # 把所有可见文本串接起来，供信息表行名和禁用片段扫描使用。
    str_visible_text = "".join(list_paragraph_texts)  # DOCX 全部可见文本

    # 读取 DOCX 媒体条目列表，供附图和公式嵌入数量校验复用。
    list_media_entries = read_docx_media_entries(path_docx)  # DOCX 媒体条目列表

    # 先准备 finding 列表，后续校验规则逐项追加问题说明。
    list_findings: list[str] = []  # 模板 DOCX 校验问题列表

    # 检查模板信息表是否仍然存在。
    if str_document_xml.count("<w:tbl") < 1:

        # 记录缺少信息表的问题。
        list_findings.append("缺少模板信息表。")

    # 检查信息表行名是否仍然可见。
    for str_label in TEMPLATE_TABLE_LABELS:

        # 任一模板行名缺失都说明表格被破坏或错误替换。
        if str_label not in str_visible_text:

            # 记录缺失的表格行名，便于定位模板表格问题。
            list_findings.append(f"缺少模板信息表行名：{str_label}")

    # 收集模板标题索引，供标题存在性和章节正文非空检查复用。
    dict_heading_indices = collect_heading_indices(list_paragraph_texts)  # 模板标题段落索引映射

    # 检查全部模板标题是否存在。
    for str_heading in TEMPLATE_SECTION_ORDER:

        # 任一标题缺失都说明最终交底书不是严格模板结构。
        if str_heading not in dict_heading_indices:

            # 记录缺失的模板标题。
            list_findings.append(f"缺少模板章节标题：{str_heading}")

    # 检查关键技术章节是否有正文内容。
    for str_heading in TEMPLATE_REQUIRED_BODY_HEADINGS:

        # 提取当前章节的正文段落，父级总章不在该列表内。
        list_body_texts = collect_section_body_texts(  # 当前章节边界内的正文段落
            str_heading,  # 待检查的模板章节标题
            list_paragraph_texts,  # 最终 DOCX 的段落文本序列
            dict_heading_indices,  # 已定位的模板标题索引
        )

        # 关键章节没有任何正文时阻断最终交付。
        if not list_body_texts:

            # 记录空章节问题，提示上游补齐技术内容。
            list_findings.append(f"模板章节正文为空：{str_heading}")

    # 检查最终 DOCX 是否残留内部审查、模板提示或待确认占位。
    for str_forbidden_text in FORBIDDEN_DISCLOSURE_TEXTS:

        # 禁止片段一旦进入主 DOCX，就不应允许进入 completed 状态。
        if str_forbidden_text in str_visible_text:

            # 记录残留片段，供调用方定位正文清理问题。
            list_findings.append(f"最终 DOCX 残留禁止文本：{str_forbidden_text}")

    # 在当前导出按输入应嵌入媒体时，强制校验媒体条目和 drawing 节点数量。
    if int_expected_media_count > 0:

        # 媒体条目数量不足时阻断交付，说明图像并未完整写入最终 DOCX 包。
        if len(list_media_entries) < int_expected_media_count:

            # 记录媒体条目不足问题，帮助定位公式或附图嵌入链路缺口。
            list_findings.append(
                f"DOCX 媒体条目不足：expected>={int_expected_media_count}, actual={len(list_media_entries)}"
            )

        # drawing 节点数量不足时阻断交付，避免媒体文件进入 ZIP 但未挂接到正文。
        if str_document_xml.count("<w:drawing") < int_expected_media_count:

            # 记录正文未真实嵌入媒体的问题。
            list_findings.append(
                f"DOCX drawing 节点不足：expected>={int_expected_media_count}, actual={str_document_xml.count('<w:drawing')}"
            )

    # 返回所有发现的问题；空列表表示当前 DOCX 通过严格模板校验。
    return list_findings

# 执行严格模板 DOCX 校验，不通过时抛出明确错误。
def validate_template_docx_output(
    path_docx: Path,
    int_expected_media_count: int = 0,
) -> None:
    """校验严格模板 DOCX 输出。

    参数：
    - `path_docx`：待检查的 DOCX 文件路径。
    - `int_expected_media_count`：按导出输入应至少嵌入的媒体数量。

    返回：
    - `None`：通过时不返回额外数据。

    异常：
    - 校验不通过时抛出 `ValueError`。
    """

    # 收集当前 DOCX 的结构与正文质量问题。
    list_findings = collect_template_docx_findings(path_docx, int_expected_media_count)  # 阻断导出的模板问题列表

    # 没有发现问题时直接返回，表示最终 DOCX 可以作为主交付件。
    if not list_findings:

        # 校验通过，无需额外处理。
        return

    # 把问题合并为可读错误文本，便于 CLI stderr 直接定位失败原因。
    str_joined_findings = "\n".join(f"- {str_item}" for str_item in list_findings)  # 模板 DOCX 校验失败明细

    # 抛出明确错误，阻止 pipeline 把不合格 DOCX 标为 completed。
    raise ValueError("> ERR: [Python] 严格模板 DOCX 校验失败。\n" + str_joined_findings)

# 清空模板正文并保留首个信息表与最终 section，供代理交付版 DOCX 在模板壳上重建正文。
def clear_template_body_keep_first_table(obj_document: Any) -> None:
    """清空模板正文并保留首个信息表与最终 section。

    参数：
    - `obj_document`：基于模板打开的 python-docx `Document` 对象。

    返回：
    - `None`。

    异常：
    - 底层 XML 操作失败时由 python-docx 对象异常继续上抛。
    """

    # 读取模板 body XML 节点，供低层子节点筛选和删除复用。
    obj_body = obj_document._element.body  # 模板正文 body XML 节点

    # 复制当前所有 body 子节点，避免遍历时边删边改导致迭代错乱。
    list_children = list(obj_body)  # 模板 body 原始子节点列表

    # 记录需要保留的首个信息表节点；模板若无表则保持空值。
    obj_first_table = find_first_body_child_by_suffix(list_children, "}tbl")  # 模板首个信息表节点

    # 记录最终 section 节点，避免清正文时把页面设置一并删掉。
    obj_section = find_first_body_child_by_suffix(list_children, "}sectPr")  # 模板最终 section 节点

    # 逐项删除除首个信息表与最终 section 外的全部正文节点。
    for obj_child in list_children:

        # 命中需要保留的表格或 section 时跳过删除，保持模板外壳稳定。
        if obj_child is obj_first_table or obj_child is obj_section:

            # 当前子节点属于模板保留壳，不删除。
            continue

        # 删除当前正文节点，让后续导出按新的交底书内容重建主稿。
        obj_body.remove(obj_child)

# 根据模板章节标题选择 python-docx 的标题层级。
def resolve_docx_heading_level(str_heading: str) -> int:
    """根据模板章节标题选择 python-docx 标题层级。

    参数：
    - `str_heading`：模板章节标题。

    返回：
    - `int`：可直接传给 `add_heading` 的标题层级。

    异常：
    - 无。
    """

    # 数字编号小节使用较低层级标题，减少版式跳跃。
    if str_heading[:1].isdigit():

        # 返回适合 3.x 和 4.x 小节的三级标题层级。
        return 3

    # 中文大章节使用一级标题层级，突出正式模板主结构。
    return 1

# 把 Markdown 公式块渲染为 PNG 图片，供代理交付版 DOCX 以内嵌对象方式展示公式。
def render_formula_image(path_output_png: Path, str_formula: str) -> None:
    """把公式块渲染为 PNG 图片。

    参数：
    - `path_output_png`：公式图片输出路径。
    - `str_formula`：Markdown 公式块正文。

    返回：
    - `None`。

    异常：
    - 图片写入失败时由底层异常继续上抛。
    """

    # 先收集非空公式正文行，供空公式判定和公式渲染共用同一份规整结果。
    list_clean_formula_lines = collect_nonempty_stripped_lines(str_formula.splitlines())  # 渲染前保留的有效公式行

    # 把规整后的多行公式压成单行文本，供 matplotlib 与 Pillow 回退路径复用同一语义输入。
    str_formula_text = " ".join(list_clean_formula_lines)  # 当前公式块压缩后的单行文本

    # 预先准备纯文本 fallback 要使用的可读公式表达，避免回退时把 LaTeX 命令原样带进交付主稿。
    str_fallback_formula_text = normalize_formula_text_for_fallback(str_formula_text)  # 纯文本回退使用的可读公式正文

    # 在公式正文为空时直接写出空白兜底文本，避免后续图片插入阶段找不到文件。
    if not str_formula_text:

        # 为空公式写出最小兜底文本，保持导出链路不中断。
        str_formula_text = "formula unavailable"  # 空公式的最小兜底文本

        # 同步更新纯文本回退表达，避免空公式路径仍沿用旧内容。
        str_fallback_formula_text = "formula unavailable"  # 空公式对应的回退公式正文

    # 优先尝试用 matplotlib 渲染数学公式；环境不具备时回退到 Pillow 文本图。
    try:

        # 只在函数内部导入 matplotlib，避免模块导入期强依赖绘图库。
        from matplotlib import pyplot as plt

    # 缺少 matplotlib 时回退到 Pillow 文本图，至少保证代理侧能直接阅读公式内容。
    except Exception:

        # 直接进入 Pillow 渲染回退路径，保持模板导出在轻量运行时可继续执行。
        render_formula_image_with_pillow(path_output_png, str_fallback_formula_text)

        # Pillow 回退已经完成当前公式图片落盘，这里可以直接结束函数。
        return

    # 先准备白底 figure，后续按数学公式或纯文本两种模式二选一渲染。
    obj_figure = plt.figure(figsize=(10.0, 1.4), dpi=200)  # 当前公式图对象

    # 把 figure 背景固定为白色，保证插入 Word 后对代理阅读更稳定。
    obj_figure.patch.set_facecolor("white")

    # 准备一块无坐标轴的画布区域，只用于居中写公式图片正文。
    obj_axes = obj_figure.add_axes([0.0, 0.0, 1.0, 1.0])  # 当前公式图画布区域

    # 关闭坐标轴显示，避免公式图混入无关刻度和边框。
    obj_axes.axis("off")

    # 先假设当前公式可按数学公式渲染，失败时再回退到纯文本渲染。
    try:

        # 以数学公式模式把当前公式写入画布中央。
        obj_text = write_centered_formula_text(obj_axes, str_formula_text, 16, True)  # 当前公式绘制文本对象

        # 先触发一次画布排版，获取数学公式的真实边界。
        obj_figure.canvas.draw()

    # 数学公式渲染失败时回退到纯文本模式，至少保证代理可阅读公式正文。
    except Exception:

        # 清空当前画布，避免数学模式的残留对象影响纯文本回退。
        obj_axes.clear()

        # 继续保持当前回退画布无坐标轴显示。
        obj_axes.axis("off")

        # 以纯文本模式重新写入公式正文，保证数学含义至少可见可读。
        obj_text = write_centered_formula_text(obj_axes, str_fallback_formula_text, 14, False)  # 纯文本回退模式文本对象

        # 重新排版纯文本回退结果，供后续边界框估算复用。
        obj_figure.canvas.draw()

    # 根据当前文本对象的实际边界估算导出图片裁切范围。
    obj_bbox = obj_text.get_window_extent(renderer=obj_figure.canvas.get_renderer()).expanded(1.15, 1.5)  # 当前公式图片边界框

    # 把公式图裁切后保存为 PNG，供 DOCX 嵌图阶段直接消费。
    obj_figure.savefig(
        path_output_png,
        dpi=200,
        bbox_inches=obj_bbox.transformed(obj_figure.dpi_scale_trans.inverted()),
        facecolor="white",
        edgecolor="none",
    )

    # 关闭当前 figure，避免批量导出时累积绘图资源。
    plt.close(obj_figure)

# 使用 Pillow 把公式正文渲染为纯文本 PNG，作为缺少 matplotlib 时的受控回退。
def render_formula_image_with_pillow(path_output_png: Path, str_formula_text: str) -> None:
    """使用 Pillow 把公式正文渲染为 PNG。

    参数：
    - `path_output_png`：公式图片输出路径。
    - `str_formula_text`：已规整为单行或少量换行的公式正文。

    返回：
    - `None`。

    异常：
    - 图片写入失败时由底层异常继续上抛。
    """

    # 只在函数内部导入 Pillow，避免模块导入期强依赖图像库。
    from PIL import Image, ImageDraw, ImageFont

    # 公式文本为空时写出最小兜底内容，避免生成空白图片。
    if not str_formula_text.strip():

        # 当前回退文本为空时补默认文案，保持导出链路稳定。
        str_formula_text = "formula unavailable"  # 空文本回退时的最小公式文案

    # 把回退文本规整成非空文本行列表，供后续拼块和测量尺寸复用。
    list_formula_lines = collect_nonempty_stripped_lines(str_formula_text.splitlines())  # 待绘制的公式文本行列表

    # 没有可见文本行时补单行兜底文本，避免后续尺寸计算越界。
    if not list_formula_lines:

        # 兜底文本行用于支撑最小 PNG 输出。
        list_formula_lines = ["formula unavailable"]  # 回退路径的单行公式兜底文本

    # 把规整后的文本行拼成多行文本块，供 Pillow 的多行测量和绘制接口复用。
    str_formula_block = "\n".join(list_formula_lines)  # 当前回退路径的多行公式文本块

    # 优先加载支持乘号等运算符的 TrueType 字体，避免默认位图字体把公式符号渲染成方块。
    obj_font = build_formula_fallback_font(ImageFont)  # Pillow 公式回退字体对象

    # 先用探针画布估算文本尺寸，避免最终图片过窄导致裁切。
    obj_probe_image = Image.new("RGB", (FORMULA_FALLBACK_PROBE_SIZE, FORMULA_FALLBACK_PROBE_SIZE), "white")  # 文本尺寸探针画布

    # 基于探针画布创建绘图对象，后续统一复用它测量文本边界框。
    obj_probe_draw = ImageDraw.Draw(obj_probe_image)  # 探针画布绘制对象

    # 基于规整后的多行公式正文测量边界框，供最终 PNG 画布尺寸计算复用。
    obj_text_bbox = measure_formula_block_bbox(obj_probe_draw, str_formula_block, obj_font)  # 当前多行公式文本块边界框

    # 计算最大文本宽度，作为最终图片宽度的主要依据。
    int_text_width = max(obj_text_bbox[2] - obj_text_bbox[0], FORMULA_FALLBACK_TEXT_WIDTH_DEFAULT)  # 最大文本宽度

    # 计算多行文本块总高度，作为最终图片高度的主要依据。
    int_text_height = max(obj_text_bbox[3] - obj_text_bbox[1], FORMULA_FALLBACK_LINE_HEIGHT_DEFAULT)  # 多行文本块总高度

    # 固定图片左右留白，避免公式文本紧贴边界。
    int_horizontal_padding = FORMULA_FALLBACK_HORIZONTAL_PADDING  # 图片左右留白

    # 固定图片上下留白，保证公式内容和边缘有足够呼吸空间。
    int_vertical_padding = FORMULA_FALLBACK_VERTICAL_PADDING  # 图片上下留白

    # 固定多行文本间距，避免多行公式在回退图中粘连。
    int_line_gap = FORMULA_FALLBACK_LINE_GAP  # 多行文本之间的垂直间距

    # 计算最终图片宽度并设置最小阈值，避免短公式过窄。
    int_image_width = max(int_text_width + int_horizontal_padding * 2, FORMULA_FALLBACK_MIN_WIDTH)  # 最终图片宽度

    # 计算最终图片高度并设置最小阈值，兼容单行短公式场景。
    int_image_height = max(int_text_height + int_vertical_padding * 2, FORMULA_FALLBACK_MIN_HEIGHT)  # 最终图片高度

    # 创建白底目标图片，后续把公式文本居中绘制进去。
    obj_image = Image.new("RGB", (int_image_width, int_image_height), "white")  # 目标公式图片对象

    # 为目标图片创建绘图对象，统一负责文本写入。
    obj_draw = ImageDraw.Draw(obj_image)  # 目标图片绘图对象

    # 计算多行文本块的居中起点，避免回退公式图在短文本场景下偏向左上角。
    int_text_x = max((int_image_width - int_text_width) // 2, int_horizontal_padding)  # 多行文本块横向起点

    # 计算多行文本块的纵向起点，保证上下留白和整体视觉平衡。
    int_text_y = max((int_image_height - int_text_height) // 2, int_vertical_padding)  # 多行文本块纵向起点

    # 把多行公式文本块一次性写入白底图片，保证代理侧看到的是直接可读的公式表达。
    obj_draw.multiline_text(
        (int_text_x, int_text_y),
        str_formula_block,
        fill="black",
        font=obj_font,
        spacing=int_line_gap,
        align="center",
    )

    # 把最终公式图片写到目标路径，供 DOCX 主稿嵌图阶段直接消费。
    obj_image.save(path_output_png)

# 向 DOCX 主文档追加一段居中的图片段落，供附图和公式嵌入复用。
def add_centered_picture_paragraph(
    obj_document: Any,
    path_image: Path,
    float_width_inches: float,
) -> None:
    """向 DOCX 主文档追加居中的图片段落。

    参数：
    - `obj_document`：待写入的 python-docx `Document` 对象；shape=单个文档句柄，dtype=runtime object，unit=none。
    - `path_image`：待嵌入的图片路径；shape=单个文件路径，dtype=`Path`，unit=file path。
    - `float_width_inches`：图片宽度；shape=标量，dtype=`float`，unit=inch。

    返回：
    - `None`：shape=标量，dtype=`NoneType`，unit=none。

    异常：
    - python-docx 插图失败时由底层异常继续上抛。
    """

    # 只在函数内部导入 python-docx 的版式工具，避免模块导入期强依赖第三方包。
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    # 新建一个独立段落承载当前图片，避免图片和正文文字混排影响阅读。
    obj_paragraph = obj_document.add_paragraph()  # 当前图片段落对象

    # 把当前图片段落设置为居中，保证公式和附图在交底书正文中更符合阅读习惯。
    obj_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 当前图片段落使用居中对齐

    # 向当前段落追加一个 run，并把图片作为嵌入对象写入段落。
    obj_paragraph.add_run().add_picture(str(path_image), width=Inches(float_width_inches))

# 把模板章节块写入 python-docx 文档，并在附图章节自动嵌入正式 PNG 附图。
def append_template_section_blocks_to_document(
    obj_document: Any,
    str_heading: str,
    list_blocks: list[dict[str, str]],
    list_figure_paths: list[Path],
    path_temp_dir: Path, obj_runtime_module: Any,
    int_formula_index_start: int,
) -> int:
    """把模板章节块写入 python-docx 文档。

    参数：
    - `obj_document`：待写入的 python-docx `Document` 对象。
    - `str_heading`：当前模板章节标题。
    - `list_blocks`：当前章节的结构化块列表。
    - `list_figure_paths`：当前案件可嵌入 DOCX 的 PNG 附图路径列表。
    - `path_temp_dir`：本轮导出的临时渲染目录路径。
    - `obj_runtime_module`：共享运行时支持模块对象。
    - `int_formula_index_start`：当前公式图片编号起点。

    返回：
    - `int`：写完本章节后的下一个公式图片编号。

    异常：
    - 图片渲染、图片插入或正文写入失败时由底层异常继续上抛。
    """

    # 先写入当前模板章节标题，保证最终 DOCX 主结构与模板合同一致。
    obj_document.add_heading(str_heading, level=resolve_docx_heading_level(str_heading))

    # 记录当前章节可复用的公式图片编号，后续每写入一个公式自增一次。
    int_formula_index = int_formula_index_start  # 当前章节公式图片编号

    # 逐块写入当前章节正文；段落按纯文本写入，公式按 PNG 嵌图写入。
    for dict_block in list_blocks:

        # 当前普通段落直接写入 DOCX 正文。
        if dict_block["kind"] == "paragraph":

            # 把当前段落文本写入 DOCX，保持与 Markdown 主稿同序。
            obj_document.add_paragraph(str(dict_block["text"]))

            # 当前段落已写入，继续处理下一个结构化块。
            continue

        # 当前公式块需要先渲染成 PNG，再以嵌入对象方式写入 DOCX。
        if dict_block["kind"] == "formula":

            # 为当前公式图片构造稳定文件名，便于同一导出轮次多公式顺序复用。
            str_formula_file_stem = obj_runtime_module.sanitize_name(str_heading)  # 当前章节名清洗后的文件名片段

            # 拼出当前公式图路径，保持同一导出轮次内的命名稳定且可追踪。
            path_formula_png = (  # 当前公式图片路径
                path_temp_dir / f"formula_{int_formula_index:02d}_{str_formula_file_stem}.png"  # 当前公式图片完整路径
            )

            # 把当前公式块渲染成 PNG，供 DOCX 主稿以内嵌对象方式展示。
            render_formula_image(path_formula_png, dict_block["text"])

            # 把当前公式图片以居中图片段落方式嵌入主稿。
            add_centered_picture_paragraph(obj_document, path_formula_png, 5.8)

            # 当前公式图片已经占用一个编号，推进到下一张公式图编号。
            int_formula_index += 1  # 推进到下一张公式图片编号

    # 在附图说明章节后追加正式 PNG 附图，保证代理看到真实图而非只有图号说明。
    if str_heading == "五、附图及附图的简单说明":

        # 逐张写入正式 PNG 附图，保持方法流程图和系统模块图顺序稳定。
        for path_figure in list_figure_paths:

            # 只对真实存在的 PNG 文件执行嵌图，避免意外文件缺失时直接崩溃。
            if path_figure.exists():

                # 把当前正式附图写入 DOCX 主稿。
                add_centered_picture_paragraph(obj_document, path_figure, 6.0)

    # 返回推进后的公式图片编号，供后续章节继续顺序命名。
    return int_formula_index

# 使用模板 DOCX 生成严格交底书 DOCX，并把公式与附图真正嵌入主稿。
def export_with_template_docx(
    dict_paths: dict[str, Path | None],
    obj_runtime_module: Any,
) -> dict[str, Any]:
    """使用模板 DOCX 生成严格交底书 DOCX。

    参数：
    - `dict_paths`：已经解析完成的输入、输出和模板路径集合。
    - `obj_runtime_module`：共享运行时支持模块对象。

    返回：
    - `dict[str, Any]`：包含导出模式和内部说明行的结果字典。

    异常：
    - 模板缺失、Markdown 读取、图片渲染或 DOCX 写入失败时由底层异常继续上抛。
    """

    # 模板导出依赖 python-docx 打开模板并写入真实嵌图对象。
    from docx import Document

    # 读取模板路径，调用方已经在主流程中确认其存在。
    path_template = dict_paths["path_template"]  # 模板 DOCX 路径

    # 模板路径缺失时立即阻断，避免静默退回空白 DOCX 破坏交底书合同。
    if path_template is None or not path_template.exists():

        # 抛出明确错误，让上游知道严格模板导出缺少模板资产。
        raise FileNotFoundError("> ERR: [Python] 缺少专利技术交底书 DOCX 模板。")

    # 读取 Markdown 全文，准备拆分成模板章节块和内部说明内容。
    str_markdown = dict_paths["path_input"].read_text(encoding="utf-8")  # 输入 Markdown 全文

    # 解析模板章节块和内部说明，确保主稿只保留可提交代理的正式正文。
    dict_template_payload = collect_template_section_blocks(str_markdown)  # 模板导出结构化块载荷

    # 收集当前案件可嵌入 DOCX 的正式 PNG 附图路径列表。
    list_figure_paths = collect_delivery_figure_image_paths(dict_paths.get("path_case_dir"))  # 当前案件正式 PNG 附图路径列表

    # 统计当前 Markdown 中的公式块数量，供最终 DOCX 媒体数量校验复用。
    int_formula_count = count_formula_blocks(dict_template_payload["sections"])  # 当前 Markdown 公式块数量

    # 确保目标导出目录存在，避免 DOCX 保存阶段因目录缺失失败。
    obj_runtime_module.ensure_dir(dict_paths["path_output"].parent)

    # 打开模板文档对象，后续在保留信息表和版式的前提下重建正文主体。
    obj_document = Document(str(path_template))  # 基于模板打开的 Word 文档对象

    # 加载独立槽位渲染器，保留模板标题节点、分节和正文段落样式而非重建 Heading。
    obj_template_renderer = load_template_renderer_module()  # 模板槽位渲染器模块对象

    # 在独立临时目录中渲染公式图片，避免把中间 PNG 暴露到正式交付目录。
    with tempfile.TemporaryDirectory() as str_temp_dir:

        # 固定本轮导出的临时渲染目录路径，供公式 PNG 稳定落盘。
        path_temp_dir = Path(str_temp_dir)  # 本轮导出的临时渲染目录路径

        # 按原模板标题节点替换正文、公式和附图，保留其两个 section 与 Normal 段落样式。
        obj_template_renderer.replace_template_slots(
            obj_document,
            TEMPLATE_SECTION_ORDER,
            dict_template_payload["sections"],
            list_figure_paths,
            path_temp_dir,
            render_formula_image,
        )

        # 把当前模板文档保存到目标输出路径，形成正式交付 DOCX。
        obj_document.save(str(dict_paths["path_output"]))

    # 同内容附图会被 python-docx 合并为一个媒体部件，因此按真实字节内容去重估算媒体下限。
    set_unique_figure_contents = {
        path_figure.read_bytes()  # 以真实媒体字节作为去重键
        for path_figure in list_figure_paths  # 遍历本次准备嵌入的正式附图
        if path_figure.exists()  # 忽略已被上游移除的失效路径
    }  # 当前正式附图的去重字节内容集合

    # 公式各自独立渲染，附图按唯一内容计数，避免把关系引用数误当作 ZIP 媒体部件数。
    int_expected_media_count = int_formula_count + len(set_unique_figure_contents)  # 严格模板校验的最小媒体部件数量

    # 对最终 DOCX 执行严格模板校验，并要求媒体数量满足公式和附图嵌入预期。
    validate_template_docx_output(
        dict_paths["path_output"],
        int_expected_media_count=int_expected_media_count,
    )

    # 返回导出结果，供上游登记导出模式和可选内部说明。
    return {"mode": "template-docx", "internal_lines": dict_template_payload["internal_lines"]}

# 把单个线性 block 写入 python-docx 文档对象，供增强导出路径复用。
def append_block_to_document(obj_document: Any, dict_block: dict[str, Any]) -> None:
    """把单个线性 block 写入 python-docx 文档对象。

    参数：
    - `obj_document`：python-docx 的 `Document` 文档对象。
    - `dict_block`：当前待写入的线性正文 block 字典。

    返回：
    - `None`。

    异常：
    - python-docx 写入失败时由底层异常继续上抛。
    """

    # 在当前 block 是分页符时直接追加 Word 分页。
    if dict_block["kind"] == "page_break":

        # 把分页符插入文档流中，为附件章节预留清晰的版面切换点。
        obj_document.add_page_break()

        # 当前分页 block 已完成写入，不需要继续走正文和标题分支。
        return

    # 在当前 block 是标题时按受控层级写入 Word 标题样式。
    if dict_block["kind"] == "heading":

        # 用 Word 原生标题样式写入当前标题正文，提升阅读层次。
        obj_document.add_heading(str(dict_block["text"]), level=int(dict_block["level"]))

        # 当前标题 block 已完成写入，不需要继续落入普通段落分支。
        return

    # 将剩余普通段落 block 直接追加到 Word 正文中。
    obj_document.add_paragraph(str(dict_block["text"]))

# 在存在模板文件时复制其首页版式，避免增强导出路径完全丢失模板页边距设置。
def copy_template_layout(obj_document: Any, path_template: Path | None) -> None:
    """复制模板首页版式设置。

    参数：
    - `obj_document`：当前待写入的 python-docx 文档对象。
    - `path_template`：可选模板 DOCX 路径。

    返回：
    - `None`。

    异常：
    - 读取模板失败时由底层异常继续上抛。
    """

    # 在未提供模板路径或模板文件缺失时直接跳过版式复制。
    if path_template is None or not path_template.exists():

        # 在没有可读取模板时提前返回，让增强导出继续使用默认版式。
        return

    # 只在函数内部导入 python-docx，避免模块导入期强依赖第三方包。
    from docx import Document

    # 读取模板文档对象，准备复制其首页 section 版式。
    obj_template_document = Document(str(path_template))  # 模板 DOCX 文档对象

    # 在模板文档没有 section 时直接跳过版式复制，避免空模板触发越界访问。
    if not obj_template_document.sections:

        # 在模板不提供 section 信息时直接返回默认文档版式。
        return

    # 读取模板首页 section，作为当前导出文档的版式来源。
    obj_template_section = obj_template_document.sections[0]  # 模板首页 section 对象

    # 读取目标文档首页 section，后续把模板版式拷贝到这里。
    obj_target_section = obj_document.sections[0]  # 当前导出文档首页 section 对象

    # 复制模板顶部页边距，保持文档首部留白与模板一致。
    obj_target_section.top_margin = obj_template_section.top_margin  # 目标文档顶部页边距

    # 复制模板底部页边距，保持文档底部留白与模板一致。
    obj_target_section.bottom_margin = obj_template_section.bottom_margin  # 目标文档底部页边距

    # 复制模板左侧页边距，尽量贴近模板既有版式宽度。
    obj_target_section.left_margin = obj_template_section.left_margin  # 目标文档左侧页边距

    # 复制模板右侧页边距，避免正文宽度与模板差异过大。
    obj_target_section.right_margin = obj_template_section.right_margin  # 目标文档右侧页边距

    # 复制模板页眉边距，让页眉位置和模板保持一致。
    obj_target_section.header_distance = obj_template_section.header_distance  # 目标文档页眉边距

    # 复制模板页脚边距，让页脚位置和模板保持一致。
    obj_target_section.footer_distance = obj_template_section.footer_distance  # 目标文档页脚边距

# 构造最小 Word 段落 XML 片段，供标准库 DOCX 回退路径写入正文段落。
def render_word_paragraph_xml(
    str_text: str,
    str_style_id: str | None = None,
) -> str:
    """构造最小 Word 段落 XML 片段。

    参数：
    - `str_text`：待写入段落的纯文本内容。
    - `str_style_id`：可选的 Word 段落样式 ID。

    返回：
    - `str`：单个段落的 Word XML 片段。

    异常：
    - 无。
    """

    # 在传入样式 ID 时构造段落样式 XML，否则保持空字符串。
    str_style_xml = f'<w:pPr><w:pStyle w:val="{str_style_id}"/></w:pPr>' if str_style_id else ""  # 当前段落的样式 XML 片段

    # 返回带可选样式信息的最小 Word 段落 XML。
    return (
        "<w:p>"
        f"{str_style_xml}"
        '<w:r><w:t xml:space="preserve">'
        f"{escape(str_text)}"
        "</w:t></w:r>"
        "</w:p>"
    )

# 构造最小 Word 分页 XML 片段，供标准库 DOCX 回退路径插入分页符。
def render_word_page_break_xml() -> str:
    """构造最小 Word 分页 XML 片段。

    参数：
    - 无。

    返回：
    - `str`：带分页符的最小 Word XML 片段。

    异常：
    - 无。
    """

    # 直接返回带分页符的最小 Word XML 片段。
    return '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'

# 把线性 block 列表转换成最小 document.xml，供标准库 DOCX 回退路径打包写入。
def convert_blocks_to_document_xml(list_blocks: list[dict[str, Any]]) -> str:
    """把线性 block 列表转换成最小 document.xml。

    参数：
    - `list_blocks`：按正文顺序整理好的线性 block 列表。

    返回：
    - `str`：最小可用的 Word document.xml 文本。

    异常：
    - 无。
    """

    # 先准备 Word body XML 片段列表，后续逐项追加最小回退正文结构。
    list_body_xml: list[str] = []  # 最小 DOCX body 片段缓存

    # 顺序遍历线性 block，把标题、普通段落和分页转成 Word XML。
    for dict_block in list_blocks:

        # 在当前 block 是分页符时直接写入分页 XML。
        if dict_block["kind"] == "page_break":

            # 把分页 XML 片段加入 Word body，分隔正文主体与附件章节。
            list_body_xml.append(render_word_page_break_xml())

            # 当前分页 block 已完成转换，直接继续处理下一个 block。
            continue

        # 在当前 block 是标题时按受控层级写入 Heading 样式段落。
        if dict_block["kind"] == "heading":

            # 把当前标题转换成带 Heading 样式的段落 XML。
            list_body_xml.append(
                render_word_paragraph_xml(str(dict_block["text"]), f"Heading{int(dict_block['level'])}")
            )

            # 当前标题 block 已完成转换，直接继续处理下一个 block。
            continue

        # 将普通段落 block 直接转换成正文段落 XML。
        list_body_xml.append(render_word_paragraph_xml(str(dict_block["text"])))

    # 在 block 列表为空时补一个空段落，保证最小 document.xml 结构完整。
    if not list_body_xml:

        # 写入一个空段落兜底，避免最小 DOCX 缺少正文节点。
        list_body_xml.append(render_word_paragraph_xml(""))

    # 组装 section XML，统一声明页面尺寸和页边距。
    str_section_xml = (  # document.xml 结尾的 section XML 片段
        f'<w:sectPr><w:pgSz w:w="{WORD_PAGE_WIDTH}" w:h="{WORD_PAGE_HEIGHT}"/>'
        f'<w:pgMar w:top="{WORD_PAGE_MARGIN}" w:right="{WORD_PAGE_MARGIN}" '
        f'w:bottom="{WORD_PAGE_MARGIN}" w:left="{WORD_PAGE_MARGIN}" '
        f'w:header="{WORD_HEADER_FOOTER_MARGIN}" '
        f'w:footer="{WORD_HEADER_FOOTER_MARGIN}" '
        'w:gutter="0"/></w:sectPr>'
    )

    # 返回完整的最小 document.xml 文本，供 ZIP 打包步骤直接写入。
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'mc:Ignorable="w14 wp14">'
        "<w:body>"
        + "".join(list_body_xml)
        + str_section_xml
        + "</w:body></w:document>"
    )

# 返回最小 Word 样式 XML，供标准库 DOCX 回退路径提供标题样式。
def render_styles_xml() -> str:
    """返回最小 Word 样式 XML。

    参数：
    - 无。

    返回：
    - `str`：覆盖 Normal 到 Heading4 的最小 Word 样式 XML。

    异常：
    - 无。
    """

    # 用多行 XML 文本直接描述最小样式集，避免样式行列表带来过长行和多行赋值噪声。
    str_styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="32"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="heading 2"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading3">
    <w:name w:val="heading 3"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="26"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading4">
    <w:name w:val="heading 4"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="24"/></w:rPr>
  </w:style>
</w:styles>"""  # 最小 Word 样式 XML 文本

    # 把最小样式 XML 文本交回回退导出路径直接写入 ZIP。
    return str_styles_xml

# 使用 python-docx 执行 DOCX 导出，优先提供更自然的 Word 阅读体验。
def export_with_python_docx(
    dict_paths: dict[str, Path | None],
    obj_runtime_module: Any,
) -> str:
    """使用 python-docx 执行 DOCX 导出。

    参数：
    - `dict_paths`：已经解析完成的输入、输出和模板路径集合。
    - `obj_runtime_module`：共享运行时支持模块对象。

    返回：
    - `str`：本次导出模式标识，固定返回 `python-docx`。

    异常：
    - Markdown 读取、DOCX 写入或目录创建失败时由底层异常继续上抛。
    """

    # 只在函数体内部导入 python-docx，避免模块导入阶段要求第三方包必须存在。
    from docx import Document

    # 读取 Markdown 全文，供线性 block 解析逻辑统一处理正文结构。
    str_markdown = dict_paths["path_input"].read_text(encoding="utf-8")  # 回退导出读取到的 Markdown 原文

    # 把 Markdown 全文解析成线性 block 列表，供导出后端逐项写入 Word。
    list_blocks = collect_markdown_blocks(str_markdown)  # 回退导出准备写入 ZIP 的 block 序列

    # 创建新的回退导出文档对象，供无严格模板路径时承载正式正文内容。
    obj_document = Document()  # 回退导出的空白主稿文档

    # 在提供模板时复制其首页版式，尽量贴近模板设定的页面留白。
    copy_template_layout(obj_document, dict_paths["path_template"])

    # 逐项把线性 block 写入 Word 文档，保持解析顺序与正文顺序一致。
    for dict_block in list_blocks:

        # 把当前 block 追加到 Word 文档中，统一处理标题、段落和分页。
        append_block_to_document(obj_document, dict_block)

    # 确保 DOCX 输出目录存在，避免保存阶段因目录缺失而失败。
    obj_runtime_module.ensure_dir(dict_paths["path_output"].parent)

    # 把当前 Word 文档保存到目标输出路径，形成最终 DOCX 交付件。
    obj_document.save(str(dict_paths["path_output"]))

    # 用模式标识告知上游当前走的是 python-docx 增强导出路径。
    return "python-docx"

# 使用 Python 标准库直接打包最小 DOCX，作为 python-docx 缺失时的本地回退路径。
def export_with_stdlib_docx(
    dict_paths: dict[str, Path | None],
    obj_runtime_module: Any,
) -> str:
    """使用标准库回退导出 DOCX。

    参数：
    - `dict_paths`：已经解析完成的输入、输出和模板路径集合。
    - `obj_runtime_module`：共享运行时支持模块对象。

    返回：
    - `str`：本次导出模式标识，固定返回 `stdlib-docx`。

    异常：
    - Markdown 读取、ZIP 写入或目录创建失败时由底层异常继续上抛。
    """

    # 从导出输入读取 Markdown 原文，为 ZIP 回退路径准备源文本。
    str_markdown = dict_paths["path_input"].read_text(encoding="utf-8")  # 回退后端的 Markdown 源文本

    # 把源文本压平成 block 序列，供最小 DOCX 结构逐项写入正文。
    list_blocks = collect_markdown_blocks(str_markdown)  # 当前 Markdown 的线性 block 列表

    # 确保 DOCX 输出目录存在，避免 ZIP 打包阶段因目录缺失而失败。
    obj_runtime_module.ensure_dir(dict_paths["path_output"].parent)

    # 固定根内容类型清单文本，声明最小 DOCX 包内必需部件的 MIME 类型。
    str_content_types_xml = (  # 根内容类型清单 XML 文本
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        "</Types>"
    )

    # 固定根关系文件文本，把包入口指向主文档部件。
    str_root_relationships_xml = (  # 根关系 XML 文本
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )

    # 先生成标题样式部件文本，供 ZIP 包内 `word/styles.xml` 直接使用。
    str_styles_xml = render_styles_xml()  # 写入 ZIP 的 styles.xml 正文

    # 再渲染 document.xml 文本，供回退路径写入正文主体结构。
    str_document_xml = convert_blocks_to_document_xml(list_blocks)  # 最小 DOCX document.xml 文本

    # 打开目标 DOCX ZIP 文件，逐项写入最小 Office Open XML 结构。
    with zipfile.ZipFile(
        dict_paths["path_output"],
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as obj_zip_file:

        # 写入根内容类型清单，声明最小 DOCX 包的部件类型。
        obj_zip_file.writestr("[Content_Types].xml", str_content_types_xml)

        # 写入根关系文件，把包入口指向主文档部件。
        obj_zip_file.writestr("_rels/.rels", str_root_relationships_xml)

        # 写入最小样式 XML，保证标题层级在 Word 中可读。
        obj_zip_file.writestr("word/styles.xml", str_styles_xml)

        # 写入根据线性 block 生成的 document.xml 文本。
        obj_zip_file.writestr("word/document.xml", str_document_xml)

    # 用模式标识告知上游当前走的是标准库最小 DOCX 回退路径。
    return "stdlib-docx"

# 解析输入、输出和模板路径，统一处理按案件目录自动定位正文草稿的逻辑。
def resolve_paths(
    namespace_arguments: argparse.Namespace,
    obj_runtime_module: Any,
) -> dict[str, Path | None]:
    """解析输入、输出和模板路径。

    参数：
    - `namespace_arguments`：命令行解析后的参数对象。
    - `obj_runtime_module`：共享运行时支持模块对象。

    返回：
    - `dict[str, Path | None]`：输入 Markdown、输出 DOCX 和可选模板路径的封装字典。

    异常：
    - 缺少输入来源时抛出 `ValueError`。
    - 自动定位后的输入草稿缺失时抛出 `FileNotFoundError`。
    """

    # 在调用方显式给出案件目录时解析其绝对路径，否则保留空值。
    path_case_dir = Path(namespace_arguments.case_dir).resolve() if namespace_arguments.case_dir else None  # 案件目录绝对路径

    # 在调用方显式给出输入草稿时解析其绝对路径，否则保留空值供自动定位逻辑处理。
    path_input = Path(namespace_arguments.input).resolve() if namespace_arguments.input else None  # 输入 Markdown 绝对路径

    # 在调用方显式给出模板路径时解析其绝对路径，否则保留空值。
    path_template = Path(namespace_arguments.template).resolve() if namespace_arguments.template else None  # 模板 DOCX 绝对路径

    # 在既未提供输入文件也未提供案件目录时直接报错，避免自动定位无从开始。
    if path_input is None and path_case_dir is None:

        # 抛出明确参数错误，要求调用方至少提供一种正文来源。
        raise ValueError("> ERR: [Python] 请提供 --input 或 --case-dir。")

    # 在未显式提供输入草稿时按案件目录自动定位当前可用正文文件。
    if path_input is None:

        # 通过共享运行时支持模块查找当前案件下最合适的正文草稿。
        path_input = obj_runtime_module.find_disclosure_draft(path_case_dir)  # 自动定位到的正文草稿路径

    # 在最终仍无法获得有效正文草稿时立即报错。
    if path_input is None or not path_input.exists():

        # 抛出明确文件缺失错误，避免后续导出路径对空输入继续工作。
        raise FileNotFoundError("> ERR: [Python] 缺少 disclosure draft markdown。")

    # 在调用方显式给出输出路径时直接解析并使用该绝对路径。
    if namespace_arguments.output:

        # 解析调用方显式指定的 DOCX 输出路径，作为本次最终交付位置。
        path_output = Path(namespace_arguments.output).resolve()  # 显式指定的 DOCX 输出路径

    # 在未显式给出输出路径时按案件导出目录和时间戳自动构造文件名。
    else:

        # 在案件目录尚未明确时从输入 Markdown 的目录结构反推案件根目录。
        if path_case_dir is None:

            # 根据正式案件目录布局从输入文件位置回推出案件根目录。
            path_case_dir = path_input.parent.parent  # 由输入文件位置反推出的案件根目录

        # 确保案件导出目录存在，后续 DOCX 会稳定落到这里。
        path_export_dir = obj_runtime_module.ensure_dir(path_case_dir / "05_exports")  # 当前案件导出目录

        # 基于输入草稿名和当前时间戳自动构造 DOCX 文件名。
        str_output_name = (  # 自动生成的 DOCX 文件名
            f"{obj_runtime_module.sanitize_name(path_input.stem)}_"  # 清理后的草稿名称前缀
            f"{obj_runtime_module.now_timestamp()}.docx"  # 避免覆盖历史导出的时间戳后缀
        )

        # 拼出最终 DOCX 输出路径，保持正式导出目录结构一致。
        path_output = path_export_dir / str_output_name  # 自动构造的 DOCX 输出路径

    # 用字典封装解析结果，减少主流程中的多值拆包复杂度。
    dict_paths = {  # 当前导出流程使用的受控路径集合
        "path_case_dir": path_case_dir,  # 当前案件根目录
        "path_input": path_input,  # 当前要导出的 Markdown 主稿路径
        "path_output": path_output,  # 当前 DOCX 主交付件输出路径
        "path_template": path_template,  # 当前导出流程使用的模板路径
    }

    # 将已经解析完成的路径字典交回主流程继续导出。
    return dict_paths

# 生成导出说明 Markdown，记录源文件、导出模式和模板来源，便于回看导出上下文。
def render_export_note(
    path_input: Path,
    str_mode: str,
    path_template: Path | None,
) -> str:
    """渲染导出说明 Markdown 文本。

    参数：
    - `path_input`：输入 Markdown 路径。
    - `str_mode`：实际采用的导出模式标识。
    - `path_template`：可选模板 DOCX 路径。

    返回：
    - `str`：导出说明 Markdown 文本。

    异常：
    - 无。
    """

    # 在存在模板路径时提取模板文件名，否则回退到 `none`。
    str_template_name = path_template.name if path_template else "none"  # 导出说明中展示的模板名称

    # 先准备导出说明文本行列表，后续按固定顺序逐条登记说明内容。
    list_export_note_lines = [TEXT_EXPORT_NOTE_TITLE]  # 导出说明 Markdown 行列表

    # 为标题与正文条目之间补一个空行，保持 sidecar 可读性。
    list_export_note_lines.append("")

    # 登记本次导出的 Markdown 来源文件名，方便后续回看输入材料。
    list_export_note_lines.append(f"- source markdown: `{path_input.name}`")

    # 登记本次实际采用的导出模式，便于判断是否走了标准库回退。
    list_export_note_lines.append(f"- export mode: `{str_mode}`")

    # 登记本次使用的模板名称，便于追溯页面版式来源。
    list_export_note_lines.append(f"- template: `{str_template_name}`")

    # 在说明末尾补一个空行，保持 sidecar 文本结尾结构稳定。
    list_export_note_lines.append("")

    # 拼接导出说明 Markdown 文本，供 sidecar 文件直接写入。
    return "\n".join(list_export_note_lines)

# 生成提交说明 sidecar，承接行政空白和不进入主 DOCX 的内部审查内容。
def render_submission_note(
    path_input: Path,
    path_template: Path | None,
    list_internal_lines: list[str],
) -> str:
    """渲染提交说明 Markdown 文本。

    参数：
    - `path_input`：输入 Markdown 路径。
    - `path_template`：可选模板 DOCX 路径。
    - `list_internal_lines`：从主 DOCX 移出的内部审查说明行。

    返回：
    - `str`：提交说明 Markdown 文本。

    异常：
    - 无。
    """

    # 提交说明只展示文件名，避免把本地绝对路径写入 sidecar。
    str_template_name = path_template.name if path_template else "none"  # 提交说明中的模板名称

    # 先准备提交说明基础信息，避免内部材料混进最终 DOCX 主体。
    list_submission_lines = [  # 提交说明基础段落
        TEXT_SUBMISSION_NOTE_TITLE,  # 固定 sidecar 标题
        "",  # 标题后的 Markdown 空行
        f"- source markdown: `{path_input.name}`",  # 仅记录输入文件名
        f"- template: `{str_template_name}`",  # 仅记录模板文件名
        "",  # 基础信息与行政清单分隔
        "## 行政信息待确认",  # 行政空白集中列示
    ]

    # 逐项列出模板行政字段，避免空白信息在交付说明中不可见。
    for str_label in ADMIN_LABELS:

        # 写入当前待确认字段名，供人工按模板信息表逐项核对。
        list_submission_lines.append(f"- {str_label}")

    # 内部审查章节从主 DOCX 移出后在 sidecar 中保留可追溯文本。
    list_submission_lines.extend(["", "## 内部审查材料"])

    # 没有内部审查材料时给出明确说明，保持 sidecar 结构稳定。
    if not list_internal_lines:

        # 写入无内部材料的说明，避免空章节让读者误判写出失败。
        list_submission_lines.append("- 无")

    # 有内部审查材料时逐行展开，保留原先的标题和正文顺序。
    else:

        # 逐行写入内部材料，供后续审查或补充证据时回看。
        list_submission_lines.extend(list_internal_lines)

    # 末尾补空行，保证 Markdown 文件以换行结束。
    list_submission_lines.append("")

    # 拼接提交说明 Markdown 文本，交给统一文件写入器落盘。
    return "\n".join(list_submission_lines)

# 执行 DOCX 导出入口，按环境能力在 python-docx 与标准库回退之间选择后端。
def main() -> int:
    """执行 DOCX 导出入口。

    参数：
    - 无。

    返回：
    - `int`：导出成功时返回 `0`。

    异常：
    - 参数无效、输入草稿缺失或导出写入失败时由底层异常继续上抛。
    """

    # 加载共享运行时支持模块，复用统一路径、时间和正文草稿查找工具。
    obj_runtime_module = load_runtime_support_module()  # 共享运行时支持模块对象

    # 解析命令行参数，读取案件目录、输入、输出和模板配置。
    namespace_arguments = build_parser().parse_args()  # 导出入口命令行参数对象

    # 把命令行参数收束成统一路径字典，避免主流程手工分支拼路径。
    dict_paths = resolve_paths(namespace_arguments, obj_runtime_module)  # 主流程共享的输入输出路径字典

    # 在当前解释器缺少模板导出能力时，优先切到 Codex 文档运行时继续执行当前脚本。
    int_reexec_return_code = maybe_reexec_with_bundled_template_runtime(dict_paths["path_template"])  # 模板运行时重启退出码

    # 发生受控重启时直接复用子进程退出码，保持当前 CLI 契约不变。
    if int_reexec_return_code is not None:

        # 返回文档运行时子进程的退出码，避免当前进程继续重复执行导出逻辑。
        return int_reexec_return_code

    # 在模板文件存在时优先走严格模板导出，保留信息表和交底书章节合同。
    if dict_paths["path_template"] is not None and dict_paths["path_template"].exists():

        # 严格模板导出必须依赖可用的 python-docx；缺失时直接报出明确环境能力错误。
        if not is_python_docx_available():

            # 抛出清晰错误，阻止模板导出悄悄退化成不满足代理交付合同的回退模式。
            raise RuntimeError(
                "> ERR: [Python] 当前解释器缺少可用的 python-docx，且未找到可复用的 Codex 文档运行时。"
            )

        # 执行严格模板导出，把正式正文、附图和公式嵌入最终交底书主稿。
        export_with_template_docx(dict_paths, obj_runtime_module)

    # 在没有模板资产且 python-docx 可用时走增强导出路径，作为兼容回退。
    elif is_python_docx_available():

        # 执行 python-docx 增强导出，兼容没有严格模板资产的本地导出场景。
        export_with_python_docx(dict_paths, obj_runtime_module)

    # 在 python-docx 缺失时回退到标准库最小 DOCX 导出路径。
    else:

        # 执行标准库回退导出，保留最低可用 DOCX 写出能力。
        export_with_stdlib_docx(dict_paths, obj_runtime_module)

    # 把 DOCX 输出绝对路径作为机器可消费的单行结果写回上游流程。
    sys.stdout.write(str(dict_paths["path_output"].resolve()) + "\n")

    # 用零退出码告知调用方当前导出流程已经成功完成。
    return 0

# 在脚本被直接执行时进入命令行主流程，导入场景下不产生副作用。
if __name__ == "__main__":

    # 把主流程退出码交还给当前 shell 调用方。
    raise SystemExit(main())

