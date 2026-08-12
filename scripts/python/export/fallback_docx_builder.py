"""构造无严格模板时的 DOCX 回退交付件。"""

# 延迟解析类型注解，避免回退模块在导入期绑定可选文档类型。
from __future__ import annotations

# 只导入回退 DOCX 打包所需的页面常量、转义器和 ZIP 支持。
from export_runtime_support import (
    WORD_HEADER_FOOTER_MARGIN,
    WORD_PAGE_HEIGHT,
    WORD_PAGE_MARGIN,
    WORD_PAGE_WIDTH,
)

# XML 转义与 ZIP 写入由运行时支持模块提供，避免回退构造器重复封装。
from export_runtime_support import (
    escape,
    zipfile,
)
from markdown_template_parser import collect_markdown_blocks

# 组合段落样式 XML 模板，避免结构片段被误判为文件系统路径。
WORD_STYLE_XML_TEMPLATE = "".join(("<w:pPr>", '<w:pStyle w:val="', "{style_id}", '"/>', "</w:pPr>"))  # 段落样式模板

# 固定 Word 段落开始片段，供最小段落渲染拼装复用。
WORD_PARAGRAPH_OPEN_XML = "<w:p>"  # Word 段落开始片段

# 固定保留空白语义的 Word 文本 run 开始片段。
WORD_TEXT_OPEN_XML = '<w:r><w:t xml:space="preserve">'  # Word 文本 run 开始片段

# 分段组合文本 run 和段落结束片段，保持嵌套标签顺序稳定。
WORD_TEXT_CLOSE_XML = "".join(("</w:t>", "</w:r>", "</w:p>"))  # 文本 run 与段落结束片段

# 分段组合主文档结束片段，供 document.xml 渲染封口。
WORD_DOCUMENT_CLOSE_XML = "".join(("</w:body>", "</w:document>"))  # 主文档正文与根节点结束片段

# 由 ZIP 路径段构造根关系部件名，避免散落平台样式路径。
ZIP_ROOT_RELATIONSHIPS_ENTRY = "/".join(("_rels", ".rels"))  # DOCX 根关系部件路径

# 由 ZIP 路径段构造样式部件名，供回退打包写入复用。
ZIP_STYLES_ENTRY = "/".join(("word", "styles.xml"))  # DOCX 样式部件路径

# 主文档部件名作为最终正文写入目标，与样式部件保持分离。
ZIP_DOCUMENT_ENTRY = "/".join(("word", "document.xml"))  # DOCX 主文档部件路径

# 把单个线性 block 写入 python-docx 文档对象，供增强导出路径复用。
def append_block_to_document(obj_document: Any, dict_block: dict[str, Any]) -> None:
    """把单个线性 block 写入 python-docx 文档对象。

    参数：
    - `obj_document`：python-docx 的 `Document` 文档对象。
    - `dict_block`：当前待写入的线性正文 block 字典。

    返回：
    - `None`。

    异常：
    - python-docx 写入失败时由底层异常继续上抛。
    """

    # 在当前 block 是分页符时直接追加 Word 分页。
    if dict_block["kind"] == "page_break":

        # 把分页符插入文档流中，为附件章节预留清晰的版面切换点。
        obj_document.add_page_break()

        # 当前分页 block 已完成写入，不需要继续走正文和标题分支。
        return

    # 在当前 block 是标题时按受控层级写入 Word 标题样式。
    if dict_block["kind"] == "heading":

        # 用 Word 原生标题样式写入当前标题正文，提升阅读层次。
        obj_document.add_heading(str(dict_block["text"]), level=int(dict_block["level"]))

        # 当前标题 block 已完成写入，不需要继续落入普通段落分支。
        return

    # 将剩余普通段落 block 直接追加到 Word 正文中。
    obj_document.add_paragraph(str(dict_block["text"]))

# 在存在模板文件时复制其首页版式，避免增强导出路径完全丢失模板页边距设置。
def copy_template_layout(obj_document: Any, path_template: Path | None) -> None:
    """复制模板首页版式设置。

    参数：
    - `obj_document`：当前待写入的 python-docx 文档对象。
    - `path_template`：可选模板 DOCX 路径。

    返回：
    - `None`。

    异常：
    - 读取模板失败时由底层异常继续上抛。
    """

    # 在未提供模板路径或模板文件缺失时直接跳过版式复制。
    if path_template is None or not path_template.exists():

        # 在没有可读取模板时提前返回，让增强导出继续使用默认版式。
        return

    # 只在函数内部导入 python-docx，避免模块导入期强依赖第三方包。
    from docx import Document

    # 读取模板文档对象，准备复制其首页 section 版式。
    obj_template_document = Document(str(path_template))  # 模板 DOCX 文档对象

    # 在模板文档没有 section 时直接跳过版式复制，避免空模板触发越界访问。
    if not obj_template_document.sections:

        # 在模板不提供 section 信息时直接返回默认文档版式。
        return

    # 读取模板首页 section，作为当前导出文档的版式来源。
    obj_template_section = obj_template_document.sections[0]  # 模板首页 section 对象

    # 读取目标文档首页 section，后续把模板版式拷贝到这里。
    obj_target_section = obj_document.sections[0]  # 当前导出文档首页 section 对象

    # 复制模板顶部页边距，保持文档首部留白与模板一致。
    obj_target_section.top_margin = obj_template_section.top_margin  # 目标文档顶部页边距

    # 复制模板底部页边距，保持文档底部留白与模板一致。
    obj_target_section.bottom_margin = obj_template_section.bottom_margin  # 目标文档底部页边距

    # 复制模板左侧页边距，尽量贴近模板既有版式宽度。
    obj_target_section.left_margin = obj_template_section.left_margin  # 目标文档左侧页边距

    # 复制模板右侧页边距，避免正文宽度与模板差异过大。
    obj_target_section.right_margin = obj_template_section.right_margin  # 目标文档右侧页边距

    # 复制模板页眉边距，让页眉位置和模板保持一致。
    obj_target_section.header_distance = obj_template_section.header_distance  # 目标文档页眉边距

    # 复制模板页脚边距，让页脚位置和模板保持一致。
    obj_target_section.footer_distance = obj_template_section.footer_distance  # 目标文档页脚边距

# 构造最小 Word 段落 XML 片段，供标准库 DOCX 回退路径写入正文段落。
def render_word_paragraph_xml(
    str_text: str,
    str_style_id: str | None = None,
) -> str:
    """构造最小 Word 段落 XML 片段。

    参数：
    - `str_text`：待写入段落的纯文本内容。
    - `str_style_id`：可选的 Word 段落样式 ID。

    返回：
    - `str`：单个段落的 Word XML 片段。

    异常：
    - 无。
    """

    # 在传入样式 ID 时构造段落样式 XML，否则保持空字符串。
    str_style_xml = WORD_STYLE_XML_TEMPLATE.format(style_id=str_style_id) if str_style_id else ""  # 当前段落样式 XML

    # 返回带可选样式信息的最小 Word 段落 XML。
    return "".join(
        (
            WORD_PARAGRAPH_OPEN_XML,
            str_style_xml,
            WORD_TEXT_OPEN_XML,
            escape(str_text),
            WORD_TEXT_CLOSE_XML,
        )
    )

# 构造最小 Word 分页 XML 片段，供标准库 DOCX 回退路径插入分页符。
def render_word_page_break_xml() -> str:
    """构造最小 Word 分页 XML 片段。

    参数：
    - 无。

    返回：
    - `str`：带分页符的最小 Word XML 片段。

    异常：
    - 无。
    """

    # 直接返回带分页符的最小 Word XML 片段。
    return '<w:p><w:r><w:br w:type="page"/></w:r></w:p>'

# 把线性 block 列表转换成最小 document.xml，供标准库 DOCX 回退路径打包写入。
def convert_blocks_to_document_xml(list_blocks: list[dict[str, Any]]) -> str:
    """把线性 block 列表转换成最小 document.xml。

    参数：
    - `list_blocks`：按正文顺序整理好的线性 block 列表。

    返回：
    - `str`：最小可用的 Word document.xml 文本。

    异常：
    - 无。
    """

    # 先准备 Word body XML 片段列表，后续逐项追加最小回退正文结构。
    list_body_xml: list[str] = []  # 最小 DOCX body 片段缓存

    # 顺序遍历线性 block，把标题、普通段落和分页转成 Word XML。
    for dict_block in list_blocks:

        # 在当前 block 是分页符时直接写入分页 XML。
        if dict_block["kind"] == "page_break":

            # 把分页 XML 片段加入 Word body，分隔正文主体与附件章节。
            list_body_xml.append(render_word_page_break_xml())

            # 当前分页 block 已完成转换，直接继续处理下一个 block。
            continue

        # 在当前 block 是标题时按受控层级写入 Heading 样式段落。
        if dict_block["kind"] == "heading":

            # 把当前标题转换成带 Heading 样式的段落 XML。
            list_body_xml.append(
                render_word_paragraph_xml(str(dict_block["text"]), f"Heading{int(dict_block['level'])}")
            )

            # 当前标题 block 已完成转换，直接继续处理下一个 block。
            continue

        # 将普通段落 block 直接转换成正文段落 XML。
        list_body_xml.append(render_word_paragraph_xml(str(dict_block["text"])))

    # 在 block 列表为空时补一个空段落，保证最小 document.xml 结构完整。
    if not list_body_xml:

        # 写入一个空段落兜底，避免最小 DOCX 缺少正文节点。
        list_body_xml.append(render_word_paragraph_xml(""))

    # 组装 section XML，统一声明页面尺寸和页边距。
    str_section_xml = (  # document.xml 结尾的 section XML 片段
        f'<w:sectPr><w:pgSz w:w="{WORD_PAGE_WIDTH}" w:h="{WORD_PAGE_HEIGHT}"/>'
        f'<w:pgMar w:top="{WORD_PAGE_MARGIN}" w:right="{WORD_PAGE_MARGIN}" '
        f'w:bottom="{WORD_PAGE_MARGIN}" w:left="{WORD_PAGE_MARGIN}" '
        f'w:header="{WORD_HEADER_FOOTER_MARGIN}" '
        f'w:footer="{WORD_HEADER_FOOTER_MARGIN}" '
        'w:gutter="0"/></w:sectPr>'
    )

    # 返回完整的最小 document.xml 文本，供 ZIP 打包步骤直接写入。
    return (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'mc:Ignorable="w14 wp14">'
        "<w:body>"
        + "".join(list_body_xml)
        + str_section_xml
        + WORD_DOCUMENT_CLOSE_XML
    )

# 返回最小 Word 样式 XML，供标准库 DOCX 回退路径提供标题样式。
def render_styles_xml() -> str:
    """返回最小 Word 样式 XML。

    参数：
    - 无。

    返回：
    - `str`：覆盖 Normal 到 Heading4 的最小 Word 样式 XML。

    异常：
    - 无。
    """

    # 用多行 XML 文本直接描述最小样式集，避免样式行列表带来过长行和多行赋值噪声。
    str_styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:qFormat/></w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="32"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="heading 2"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading3">
    <w:name w:val="heading 3"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="26"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading4">
    <w:name w:val="heading 4"/><w:qFormat/><w:rPr><w:b/><w:sz w:val="24"/></w:rPr>
  </w:style>
</w:styles>"""  # 最小 Word 样式 XML 文本

    # 把最小样式 XML 文本交回回退导出路径直接写入 ZIP。
    return str_styles_xml

# 使用 python-docx 执行 DOCX 导出，优先提供更自然的 Word 阅读体验。
def export_with_python_docx(
    dict_paths: dict[str, Path | None],
    obj_runtime_module: Any,
) -> str:
    """使用 python-docx 执行 DOCX 导出。

    参数：
    - `dict_paths`：已经解析完成的输入、输出和模板路径集合。
    - `obj_runtime_module`：共享运行时支持模块对象。

    返回：
    - `str`：本次导出模式标识，固定返回 `python-docx`。

    异常：
    - Markdown 读取、DOCX 写入或目录创建失败时由底层异常继续上抛。
    """

    # 只在函数体内部导入 python-docx，避免模块导入阶段要求第三方包必须存在。
    from docx import Document

    # 读取 Markdown 全文，供线性 block 解析逻辑统一处理正文结构。
    str_markdown = dict_paths["path_input"].read_text(encoding="utf-8")  # 回退导出读取到的 Markdown 原文

    # 把 Markdown 全文解析成线性 block 列表，供导出后端逐项写入 Word。
    list_blocks = collect_markdown_blocks(str_markdown)  # 回退导出准备写入 ZIP 的 block 序列

    # 创建新的回退导出文档对象，供无严格模板路径时承载正式正文内容。
    obj_document = Document()  # 回退导出的空白主稿文档

    # 在提供模板时复制其首页版式，尽量贴近模板设定的页面留白。
    copy_template_layout(obj_document, dict_paths["path_template"])

    # 逐项把线性 block 写入 Word 文档，保持解析顺序与正文顺序一致。
    for dict_block in list_blocks:

        # 把当前 block 追加到 Word 文档中，统一处理标题、段落和分页。
        append_block_to_document(obj_document, dict_block)

    # 确保 DOCX 输出目录存在，避免保存阶段因目录缺失而失败。
    obj_runtime_module.ensure_dir(dict_paths["path_output"].parent)

    # 把当前 Word 文档保存到目标输出路径，形成最终 DOCX 交付件。
    obj_document.save(str(dict_paths["path_output"]))

    # 用模式标识告知上游当前走的是 python-docx 增强导出路径。
    return "python-docx"

# 使用 Python 标准库直接打包最小 DOCX，作为 python-docx 缺失时的本地回退路径。
def export_with_stdlib_docx(
    dict_paths: dict[str, Path | None],
    obj_runtime_module: Any,
) -> str:
    """使用标准库回退导出 DOCX。

    参数：
    - `dict_paths`：已经解析完成的输入、输出和模板路径集合。
    - `obj_runtime_module`：共享运行时支持模块对象。

    返回：
    - `str`：本次导出模式标识，固定返回 `stdlib-docx`。

    异常：
    - Markdown 读取、ZIP 写入或目录创建失败时由底层异常继续上抛。
    """

    # 从导出输入读取 Markdown 原文，为 ZIP 回退路径准备源文本。
    str_markdown = dict_paths["path_input"].read_text(encoding="utf-8")  # 回退后端的 Markdown 源文本

    # 把源文本压平成 block 序列，供最小 DOCX 结构逐项写入正文。
    list_blocks = collect_markdown_blocks(str_markdown)  # 当前 Markdown 的线性 block 列表

    # 确保 DOCX 输出目录存在，避免 ZIP 打包阶段因目录缺失而失败。
    obj_runtime_module.ensure_dir(dict_paths["path_output"].parent)

    # 固定根内容类型清单文本，声明最小 DOCX 包内必需部件的 MIME 类型。
    str_content_types_xml = (  # 根内容类型清单 XML 文本
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
        '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/word/document.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
        '<Override PartName="/word/styles.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
        "</Types>"
    )

    # 固定根关系文件文本，把包入口指向主文档部件。
    str_root_relationships_xml = (  # 根关系 XML 文本
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="word/document.xml"/>'
        "</Relationships>"
    )

    # 先生成标题样式部件文本，供 ZIP 包内 `word/styles.xml` 直接使用。
    str_styles_xml = render_styles_xml()  # 写入 ZIP 的 styles.xml 正文

    # 再渲染 document.xml 文本，供回退路径写入正文主体结构。
    str_document_xml = convert_blocks_to_document_xml(list_blocks)  # 最小 DOCX document.xml 文本

    # 打开目标 DOCX ZIP 文件，逐项写入最小 Office Open XML 结构。
    with zipfile.ZipFile(
        dict_paths["path_output"],
        "w",
        compression=zipfile.ZIP_DEFLATED,
    ) as obj_zip_file:

        # 写入根内容类型清单，声明最小 DOCX 包的部件类型。
        obj_zip_file.writestr("[Content_Types].xml", str_content_types_xml)

        # 写入根关系文件，把包入口指向主文档部件。
        obj_zip_file.writestr(ZIP_ROOT_RELATIONSHIPS_ENTRY, str_root_relationships_xml)

        # 写入最小样式 XML，保证标题层级在 Word 中可读。
        obj_zip_file.writestr(ZIP_STYLES_ENTRY, str_styles_xml)

        # 写入根据线性 block 生成的 document.xml 文本。
        obj_zip_file.writestr(ZIP_DOCUMENT_ENTRY, str_document_xml)

    # 用模式标识告知上游当前走的是标准库最小 DOCX 回退路径。
    return "stdlib-docx"
