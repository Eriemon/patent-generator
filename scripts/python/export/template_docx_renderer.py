"""按参考模板的原始节点替换专利交底书正文槽位。"""

# 启用延迟类型注解，避免运行时解析可选文档对象类型。
from __future__ import annotations

# 标准库负责复制 XML 属性、路径表达和回调类型约束。
from copy import deepcopy
from pathlib import Path
from typing import Any, Callable

# 图片引用关系属性由 OOXML 标准命名空间和本地名共同组成。
RELATIONSHIP_EMBED_ATTRIBUTE = "{" + "http://schemas.openxmlformats.org/officeDocument/2006/relationships" + "}embed"  # 图片引用关系属性

# 关系类型只比较协议末段，兼容标准前缀的不同序列化形式。
IMAGE_RELATIONSHIP_SUFFIX = "/" + "image"  # 图片关系类型末段

# 规整标题中的空白差异，使模板节点匹配不受软换行影响。
def normalize_slot_text(str_text: str) -> str:
    """返回去除全部空白的槽位标题文本。

    参数：
    - `str_text`：模板段落或合同槽位的原始标题文本。

    返回：
    - `str`：用于稳定比较的无空白标题键。
    """

    # 返回稳定的标题比较键，不改写模板中的真实显示文本。
    return "".join(str_text.split())

# 判断当前模板段落是否承载指定的正式章节标题。
def is_slot_heading(str_paragraph_text: str, str_slot_heading: str) -> bool:
    """判断模板段落是否命中指定槽位标题。

    参数：
    - `str_paragraph_text`：模板中的原始段落文本。
    - `str_slot_heading`：槽位合同中的标准标题。

    返回：
    - `bool`：原段落是纯标题或以标准标题开头时为真。
    """

    # 规整模板段落，兼容标题与提示语位于同一段落的情况。
    str_paragraph_key = normalize_slot_text(str_paragraph_text)  # 模板段落比较键

    # 规整合同标题，作为精确匹配和前缀匹配的共同基准。
    str_slot_key = normalize_slot_text(str_slot_heading)  # 正式槽位比较键

    # 同时接受纯标题和标题后附带模板提示语的原始节点。
    return str_paragraph_key == str_slot_key or str_paragraph_key.startswith(str_slot_key)

# 按模板出现顺序收集标题节点，防止相似标题被跨章节误匹配。
def collect_slot_paragraphs(obj_document: Any, list_slot_order: list[str]) -> dict[str, Any]:
    """收集全部强制槽位对应的原始模板段落。

    参数：
    - `obj_document`：从参考模板打开的 Word 文档对象。
    - `list_slot_order`：按模板正文顺序排列的强制槽位标题。

    返回：
    - `dict[str, Any]`：槽位标题到原始模板段落的映射。

    异常：
    - `ValueError`：参考模板缺少任一强制标题节点。
    """

    # 保存已定位的模板标题节点，供后续原位替换正文使用。
    dict_slot_paragraphs: dict[str, Any] = {}  # 槽位标题到模板段落的映射

    # 固化模板段落快照，避免后续 XML 修改影响本轮搜索顺序。
    list_paragraphs = list(obj_document.paragraphs)  # 模板正文段落快照

    # 从首段开始顺序查找，确保标题层级与参考模板一致。
    int_search_start = 0  # 下一个槽位的段落搜索起点

    # 依照合同顺序逐个定位强制标题节点。
    for str_slot_heading in list_slot_order:

        # 仅在前一个标题之后搜索，避免命中信息表或重复提示文字。
        for int_index in range(int_search_start, len(list_paragraphs)):

            # 读取当前候选段落，保留其原始样式和编号属性。
            obj_paragraph = list_paragraphs[int_index]  # 当前候选模板段落

            # 命中标题后记录节点，并推进下一轮搜索起点。
            if is_slot_heading(obj_paragraph.text, str_slot_heading):

                # 保存原始节点，正文装配阶段不会重新创建 Heading。
                dict_slot_paragraphs[str_slot_heading] = obj_paragraph  # 已定位的标题段落

                # 下一槽位必须位于当前标题之后。
                int_search_start = int_index + 1  # 后续标题搜索起点

                # 当前槽位已经定位，无需继续扫描后续段落。
                break

        # 缺少任一合同标题时阻断导出，避免生成结构不完整的交底书。
        if str_slot_heading not in dict_slot_paragraphs:

            # 报告具体缺失槽位，方便定位模板版本漂移。
            raise ValueError(f"> ERR: [Python] 模板缺少正文槽位：{str_slot_heading}")

    # 返回完整的标题节点映射供清理和正文插入复用。
    return dict_slot_paragraphs

# 清除标题之间的示例节点，同时保留首个正文段落的版式属性。
def remove_slot_example_content(obj_current_heading: Any, obj_next_heading: Any | None) -> Any | None:
    """删除当前槽位内的模板提示和示例正文。

    参数：
    - `obj_current_heading`：当前槽位的原始标题段落。
    - `obj_next_heading`：下一槽位标题；末尾槽位传入空值。

    返回：
    - `Any | None`：从首个正文段落复制的版式属性。
    """

    # 以当前标题 XML 节点作为遍历起点。
    obj_anchor_element = obj_current_heading._p  # 当前标题 XML 节点

    # 末尾槽位没有下一标题，遍历范围自然延伸到文档尾部。
    obj_next_element = obj_next_heading._p if obj_next_heading is not None else None  # 下一标题 XML 节点

    # 首个正文段落属性用于新正文继承模板的缩进、间距和字体语言。
    obj_template_properties: Any | None = None  # 可复用的正文段落属性

    # 从标题后的第一个兄弟节点开始清理示例内容。
    obj_candidate = obj_anchor_element.getnext()  # 当前待清理 XML 节点

    # 逐个处理到下一标题之前，避免越界删除正式标题节点。
    while obj_candidate is not None and obj_candidate is not obj_next_element:

        # 删除前保存后继节点，保证 XML 树修改后仍能继续遍历。
        obj_following = obj_candidate.getnext()  # 当前节点的后继节点

        # 文档主体末尾的独立 sectPr 承载最终页面设置，必须原样保留。
        if obj_candidate.tag.endswith("}sectPr"):

            # 跳过独立分节属性节点，只推进遍历游标。
            obj_candidate = obj_following  # 独立分节节点后的兄弟节点

            # 当前节点无需执行示例内容清理。
            continue

        # 首次遇到有版式的正文段落时保存属性，后续候选不再覆盖。
        obj_template_properties = capture_template_properties(  # 当前槽位正文版式副本
            obj_candidate,  # 当前示例内容节点
            obj_template_properties,  # 已经捕获的版式属性
        )

        # 清除当前示例节点，同时保护其中可能携带的分节属性。
        clear_template_candidate(obj_candidate)

        # 推进到删除前记录的后继节点继续清理。
        obj_candidate = obj_following  # 下一待清理 XML 节点

    # 返回当前槽位的正文版式副本供新增段落复用。
    return obj_template_properties

# 从首个合适的正文候选中复制段落属性。
def capture_template_properties(obj_candidate: Any, obj_existing_properties: Any | None) -> Any | None:
    """返回已经存在或刚从候选段落复制的正文属性。

    参数：
    - `obj_candidate`：当前待清理的模板示例节点。
    - `obj_existing_properties`：此前已经捕获的段落属性。

    返回：
    - `Any | None`：供新正文段落继承的独立属性节点。
    """

    # 已有版式样本或当前节点不是段落时直接保留原结果。
    if obj_existing_properties is not None or not obj_candidate.tag.endswith("}p"):

        # 返回既有样本，避免后续示例段落覆盖首段版式。
        return obj_existing_properties

    # 读取候选段落的属性节点，空段落可能没有该节点。
    obj_properties = obj_candidate.pPr  # 候选正文段落属性

    # 没有段落属性时继续等待下一个正文候选。
    if obj_properties is None:

        # 空值表示本槽位尚未找到可复用的版式样本。
        return None

    # 深拷贝属性节点，确保模板清理不会影响后续新增段落。
    return deepcopy(obj_properties)

# 清理单个模板示例节点，并保留分节段落的属性子树。
def clear_template_candidate(obj_candidate: Any) -> None:
    """删除示例节点或清空分节段落中的显示内容。

    参数：
    - `obj_candidate`：当前待清理的模板示例节点。

    返回：
    - `None`。
    """

    # 分节段落不能整体删除，只清理其中的文字和图片节点。
    if obj_candidate.tag.endswith("}p") and paragraph_contains_section_properties(obj_candidate):

        # 遍历分节段落的直接子节点，pPr 之外均属于显示内容。
        for obj_child in list(obj_candidate):

            # pPr 承载页面分节设置，必须保留在原文档位置。
            if not obj_child.tag.endswith("}pPr"):

                # 删除当前示例 run、图片或字段节点。
                obj_candidate.remove(obj_child)

        # 分节段落清理完成后不再执行整体删除。
        return

    # 普通示例段落、表格和图片节点可以从正文树整体移除。
    obj_candidate.getparent().remove(obj_candidate)

# 检查段落属性子树中是否包含分节设置，不依赖硬编码 XPath。
def paragraph_contains_section_properties(obj_paragraph_element: Any) -> bool:
    """判断段落节点是否承载 Word 分节属性。

    参数：
    - `obj_paragraph_element`：待检查的 WordprocessingML 段落节点。

    返回：
    - `bool`：任一段落属性子节点是分节属性时为真。
    """

    # 逐层读取 pPr 子树，避免不同 OOXML 库的命名空间前缀差异。
    for obj_properties in obj_paragraph_element:

        # 非段落属性节点不可能承载当前 section 设置。
        if not obj_properties.tag.endswith("}pPr"):

            # 继续检查当前段落的下一个直接子节点。
            continue

        # 检查段落属性中的直接子节点是否包含 sectPr。
        for obj_property_child in obj_properties:

            # 命中分节属性即可停止遍历。
            if obj_property_child.tag.endswith("}sectPr"):

                # 向清理逻辑确认当前段落必须保留。
                return True

    # 没有找到分节属性时允许按普通示例段落清理。
    return False

# 清除标题节点内的提示语，只写回合同规定的标准标题文本。
def set_heading_text(obj_paragraph: Any, str_slot_heading: str) -> None:
    """保留标题段落属性并替换其显示文本。

    参数：
    - `obj_paragraph`：需要原位更新的模板标题段落。
    - `str_slot_heading`：需要写回的标准槽位标题。

    返回：
    - `None`。
    """

    # 读取原始标题 XML 节点，编号和样式均位于其 pPr 中。
    obj_paragraph_element = obj_paragraph._p  # 原始标题 XML 节点

    # 遍历标题子节点，移除旧文字和提示语但保留段落属性。
    for obj_child in list(obj_paragraph_element):

        # pPr 是模板版式语言，不能随示例文本一起删除。
        if not obj_child.tag.endswith("}pPr"):

            # 移除当前标题中的旧 run 或其他显示节点。
            obj_paragraph_element.remove(obj_child)

    # 将标准章节标题写回原始节点，不调用 add_heading 重建结构。
    obj_paragraph.add_run(str_slot_heading)

# 在指定 XML 锚点后插入继承模板属性的新正文段落。
def insert_paragraph_after(
    obj_document: Any,
    obj_anchor_element: Any,
    obj_template_properties: Any | None,
) -> tuple[Any, Any]:
    """新增一个模板样式正文段落并返回段落及 XML 节点。

    参数：
    - `obj_document`：当前 Word 文档对象。
    - `obj_anchor_element`：新增段落需要紧随的 XML 节点。
    - `obj_template_properties`：可选的模板正文段落属性。

    返回：
    - `tuple[Any, Any]`：新增段落对象及其 XML 节点。
    """

    # 延迟导入 python-docx 类型，使模块静态检查不依赖可选文档运行时。
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    # 创建空段落 XML 节点，后续再按 block 类型写入文字或图片。
    obj_paragraph_element = OxmlElement("w:p")  # 新增正文段落 XML 节点

    # 当前槽位存在正文版式样本时复制其属性。
    if obj_template_properties is not None:

        # 深拷贝模板属性，避免多个正文段落共享可变 XML 节点。
        obj_paragraph_element.append(deepcopy(obj_template_properties))

    # 将新段落插入当前锚点之后，保持正文块顺序稳定。
    obj_anchor_element.addnext(obj_paragraph_element)

    # 返回 python-docx 包装对象和 XML 节点供调用方继续推进锚点。
    return Paragraph(obj_paragraph_element, obj_document._body), obj_paragraph_element

# 把普通正文中的文字和行内公式依次追加到同一 Word 段落。
def append_inline_content(obj_paragraph: Any, str_text: str, dict_render_context: dict[str, Any]) -> None:
    """写入普通正文及其中的 Office 原生行内公式。

    参数：
    - `obj_paragraph`：接收正文片段的 Word 段落对象。
    - `str_text`：可能包含行内公式分隔符的正文文本。
    - `dict_render_context`：公式拆分、转换和模式依赖。

    返回：
    - `None`。

    异常：
    - 行内公式分隔或转换失败时继续抛出底层异常。
    """

    # 顺序拆分普通文字与行内公式，保持公式位于原正文段落内。
    list_inline_segments = dict_render_context["inline_splitter"](str_text)  # 当前正文的行内片段列表

    # 逐片段写入文字 run 或原生行内 OMML。
    for dict_inline_segment in list_inline_segments:

        # 普通文字继续使用 Word 文本 run。
        if dict_inline_segment["kind"] == "text":

            # 保留当前片段的原始空格、标点与字符顺序。
            obj_paragraph.add_run(str(dict_inline_segment["text"]))

            # 当前文字已经写入，不进入公式转换分支。
            continue

        # MathType 模式先写入唯一定位标记，保存后由 Word COM 原位替换为 OLE。
        if dict_render_context["equation_mode"] == "mathtype":

            # 延迟导入书签 XML helper，只在原生 MathType 模式创建定位边界。
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            # 使用当前清单长度生成不会与技术正文冲突的稳定定位标记。
            str_marker = f"[[MATHTYPE_EQ_{len(dict_render_context['formula_records']) + 1:06d}]]"  # 当前行内公式标记

            # 生成符合 Word 书签命名限制的定位名称。
            str_bookmark = f"MT_EQ_{len(dict_render_context['formula_records']) + 1:06d}"  # 当前行内公式书签

            # 以递增编号创建书签开始节点，Word 会维护其真实字符位置。
            obj_bookmark_start = OxmlElement("w:bookmarkStart")  # 行内公式书签起点

            # 写入当前书签的数值标识，使起止节点可以配对。
            obj_bookmark_start.set(qn("w:id"), str(len(dict_render_context["formula_records"]) + 1000))

            # 写入供 Word COM 按名称读取的行内公式书签名。
            obj_bookmark_start.set(qn("w:name"), str_bookmark)

            # 把书签起点放在行内公式占位文本之前。
            obj_paragraph._p.append(obj_bookmark_start)

            # 把定位标记写入原公式位置，COM 完成后不会残留在最终文档。
            obj_paragraph.add_run(str_marker)

            # 在标记之后关闭书签范围，使 Word Range 只覆盖占位文本。
            obj_bookmark_end = OxmlElement("w:bookmarkEnd")  # 行内公式书签终点

            # 使用相同数值标识关闭当前行内书签。
            obj_bookmark_end.set(qn("w:id"), str(len(dict_render_context["formula_records"]) + 1000))

            # 把书签终点放在行内公式占位文本之后。
            obj_paragraph._p.append(obj_bookmark_end)

            # 登记源公式、布局和定位标记，供 MathType 写入器精确查找。
            dict_render_context["formula_records"].append(
                {
                    "latex": str(dict_inline_segment["text"]),
                    "display": False,
                    "marker": str_marker,
                    "bookmark": str_bookmark,
                }
            )

            # 当前公式已经登记为 MathType 占位符，不再生成中间 OMML。
            continue

        # 行内公式转换为 m:oMath，不创建独立段落或图片。
        obj_inline_formula = dict_render_context["formula_converter"](  # 当前可编辑行内公式节点
            str(dict_inline_segment["text"]),  # 不含分隔符的 LaTeX 正文
            False,  # 当前公式保持行内布局
            dict_render_context["equation_mode"],  # Office 或 MathType 中间模式
        )

        # 把当前数学节点原位追加到正文段落 XML。
        obj_paragraph._p.append(obj_inline_formula)

        # 记录公式源文本，供 MathType 模式按文档顺序替换中间 OMML。
        dict_render_context["formula_records"].append(
            {"latex": str(dict_inline_segment["text"]), "display": False}
        )

# 将单个槽位的正文、公式和附图依次插入原始标题节点之后。
def append_slot_blocks(
    obj_heading: Any,
    list_blocks: list[dict[str, str]],
    int_formula_index: int,
    dict_render_context: dict[str, Any],
) -> int:
    """写入一个模板槽位的正文块并返回下一公式编号。

    参数：
    - `obj_heading`：当前槽位的原始标题段落。
    - `list_blocks`：代理起草的正文和公式块列表，不涉及数组形状或数值 dtype。
    - `int_formula_index`：从一开始递增的无量纲公式编号。
    - `dict_render_context`：文档、版式、标题、OMML 公式回调和附图清单。

    返回：
    - `int`：供下一槽位继续使用的无量纲公式编号。

    异常：
    - 公式转换或 DOCX 嵌入失败时继续抛出底层异常。
    """

    # 延迟导入版式常量，只有正式导出时才要求 python-docx 可用。
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    # 首个正文块紧随原始标题节点插入。
    obj_anchor_element = obj_heading._p  # 当前正文插入锚点

    # 按代理起草结果中的稳定顺序写入正文或公式块。
    for dict_block in list_blocks:

        # 为当前 block 新建继承模板版式的正文段落。
        tuple_obj_paragraph, tuple_obj_anchor_element = insert_paragraph_after(  # 新段落及推进后的插入锚点
            dict_render_context["document"],  # 承载当前技术章节的模板文档
            obj_anchor_element,  # 前一正文块形成的 XML 顺序锚点
            dict_render_context["template_properties"],  # 从本槽位示例段落提取的版式
        )

        # 解包新增段落和后继锚点，分别承载内容写入和顺序推进职责。
        obj_paragraph = tuple_obj_paragraph  # 当前新增正文段落

        # 后续 block 必须紧随当前新增段落插入。
        obj_anchor_element = tuple_obj_anchor_element  # 推进后的正文插入锚点

        # 普通正文交给行内内容 helper 处理文字与可编辑公式。
        if dict_block["kind"] == "paragraph":

            # 写入正文文字并把其中的公式转换为原位 m:oMath 节点。
            append_inline_content(obj_paragraph, str(dict_block["text"]), dict_render_context)

            # 当前 block 已完成，继续处理下一个正文块。
            continue

        # MathType 模式在当前段落写入唯一标记，后续由 Word COM 原位创建 OLE。
        if dict_render_context["equation_mode"] == "mathtype":

            # 延迟导入书签 XML helper，只在原生 MathType 模式创建定位边界。
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            # 使用公式清单顺序生成当前行间公式的稳定定位标记。
            str_marker = f"[[MATHTYPE_EQ_{len(dict_render_context['formula_records']) + 1:06d}]]"  # 当前行间公式标记

            # 生成符合 Word 书签命名限制的行间公式定位名称。
            str_bookmark = f"MT_EQ_{len(dict_render_context['formula_records']) + 1:06d}"  # 当前行间公式书签

            # 公式段落保持居中，使替换后的 MathType 对象沿用预期版式。
            obj_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # MathType 行间公式对齐方式

            # 为独立公式段落创建书签起点，Word 将维护跨版式后的真实位置。
            obj_bookmark_start = OxmlElement("w:bookmarkStart")  # 行间公式书签起点

            # 写入当前行间书签的数值标识，使起止节点可以配对。
            obj_bookmark_start.set(qn("w:id"), str(len(dict_render_context["formula_records"]) + 1000))

            # 绑定独立公式段落与转换清单中的行间定位名称。
            obj_bookmark_start.set(qn("w:name"), str_bookmark)

            # 把书签起点放在居中公式占位文本之前。
            obj_paragraph._p.append(obj_bookmark_start)

            # 把标记写入目标段落，Word COM 将在同一 Range 内完成替换。
            obj_paragraph.add_run(str_marker)

            # 在居中标记之后关闭范围，确保书签不包含相邻技术正文。
            obj_bookmark_end = OxmlElement("w:bookmarkEnd")  # 行间公式书签终点

            # 使用相同数值标识关闭当前行间书签。
            obj_bookmark_end.set(qn("w:id"), str(len(dict_render_context["formula_records"]) + 1000))

            # 把书签终点放在行间公式占位文本之后。
            obj_paragraph._p.append(obj_bookmark_end)

            # 保存行间公式源文本、布局和精确定位标记。
            dict_render_context["formula_records"].append(
                {
                    "latex": str(dict_block["text"]),
                    "display": True,
                    "marker": str_marker,
                    "bookmark": str_bookmark,
                }
            )

            # 当前 MathType 公式已经登记，不进入 Office OMML 分支。
            continue

        # Office 模式将公式转换为原生 OMML，禁止创建任何公式图片。
        obj_formula = dict_render_context["formula_converter"](  # 当前可编辑行间公式节点
            str(dict_block["text"]),  # 当前公式的 LaTeX 正文
            True,  # 模板公式块按行间公式布局
            dict_render_context["equation_mode"],  # 行间公式沿用本次导出的兼容模式
        )

        # 公式段落保持居中，与参考模板的公式版式一致。
        obj_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 公式段落水平对齐方式

        # 将原生数学节点直接追加到段落 XML，Office 可双击进入公式编辑器。
        obj_paragraph._p.append(obj_formula)

        # 记录行间公式源文本，供原生 MathType OLE 写入器生成 MTEF 内容。
        dict_render_context["formula_records"].append(
            {"latex": str(dict_block["text"]), "display": True}
        )

        # 推进公式编号，供结构报告核对跨章节公式总量。
        int_formula_index += 1  # 下一条原生公式编号

    # 附图只写入模板规定的附图说明槽位，避免跨章节重复嵌入。
    if dict_render_context["slot_heading"] == "五、附图及附图的简单说明":

        # 按附图清单顺序写入方法流程图和系统模块图。
        for path_figure in dict_render_context["figure_paths"]:

            # 跳过不存在的可选附图，正式校验会检查实际媒体数量。
            if not path_figure.exists():

                # 当前路径无可嵌入内容，继续检查下一张附图。
                continue

            # 为当前附图创建独立居中段落。
            tuple_obj_paragraph, tuple_obj_anchor_element = insert_paragraph_after(  # 当前附图段落及新锚点
                dict_render_context["document"],  # 接收正式附图的模板文档
                obj_anchor_element,  # 前一附图或说明段落的 XML 锚点
                dict_render_context["template_properties"],  # 附图说明章节的段落版式
            )

            # 当前段落专门承载正式附图，不与说明文字混排。
            obj_paragraph = tuple_obj_paragraph  # 当前新增附图段落

            # 下一附图应排在当前图片之后。
            obj_anchor_element = tuple_obj_anchor_element  # 推进后的附图插入锚点

            # 附图在页面可用宽度内居中显示。
            obj_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 附图段落水平对齐方式

            # 嵌入正式 PNG，独立附图包仍保留原文件。
            obj_paragraph.add_run().add_picture(str(path_figure), width=Inches(6.0))

    # 返回推进后的公式编号供下一槽位继续使用。
    return int_formula_index

# 用已确认的发明名称和申请类型填写模板信息表，其余行政字段保持人工填写位。
def populate_information_table(obj_document: Any, dict_sections: dict[str, list[dict[str, str]]]) -> None:
    """填写模板信息表中的发明名称和发明专利类型。

    参数：
    - `obj_document`：当前 Word 文档对象。
    - `dict_sections`：槽位标题到正式正文块的映射。

    返回：
    - `None`。
    """

    # 模板没有信息表时不在正文中临时创建替代表格。
    if not obj_document.tables:

        # 无信息表的模板不满足后续严格校验，此处仅避免索引异常。
        return

    # 收集发明名称槽位中的正式正文块。
    list_name_blocks = dict_sections.get("一、发明名称", [])  # 发明名称正文块

    # 选择首个普通段落作为信息表中的发明名称。
    str_title = next(  # 可写入模板信息表的发明名称
        (str(dict_block["text"]) for dict_block in list_name_blocks if dict_block["kind"] == "paragraph"),  # 正文标题候选
        "",  # 发明名称缺失时保留空值供验证器阻断
    )

    # 只有真实标题存在时才覆盖模板目标单元格。
    if str_title:

        # 写入 11×2 信息表的发明名称值槽位。
        obj_document.tables[0].cell(0, 1).text = str_title  # 信息表发明名称值

    # 读取模板原有类型选项文本，仅把发明选项切换为已勾选状态。
    obj_type_cell = obj_document.tables[0].cell(1, 1)  # 发明创造类型值单元格

    # 保留实用新型、外观和 PCT 原有选项文本供精确替换。
    str_type_options = obj_type_cell.text  # 模板原始类型选项文本

    # 当前技能交付发明专利技术交底书，类型槽位必须明确勾选发明。
    if "□发明" in str_type_options:

        # 只替换首个发明选项，实用新型、外观和 PCT 选项保持模板原状。
        obj_type_cell.text = str_type_options.replace("□发明", "☒发明", 1)  # 勾选后的类型选项文本

# 清理示例图片删除后遗留的孤立图片关系。
def purge_unreferenced_image_relationships(obj_document: Any) -> None:
    """删除正文未引用的模板示例图片关系。

    参数：
    - `obj_document`：已经完成正文槽位替换的 Word 文档对象。

    返回：
    - `None`。
    """

    # DrawingML 使用该属性保存文档关系编号。
    str_embed_attribute = RELATIONSHIP_EMBED_ATTRIBUTE  # 图片关系属性名

    # 收集正文仍在使用的所有图片关系编号。
    set_referenced_ids = {  # 正文实际引用的图片关系编号
        str(obj_node.get(str_embed_attribute))  # 当前图片节点的关系编号
        for obj_node in obj_document._element.iter()  # 遍历正文 XML 子树
        if obj_node.tag.endswith("}blip") and obj_node.get(str_embed_attribute)  # 只收集有效 DrawingML 图片节点
    }

    # 对关系表创建快照后再删除，避免迭代期间改变原映射长度。
    for str_relationship_id, obj_relationship in list(obj_document.part.rels.items()):

        # 只删除未被正文引用的图片关系，其他外部关系保持不变。
        if (
            obj_relationship.reltype.endswith(IMAGE_RELATIONSHIP_SUFFIX)
            and str_relationship_id not in set_referenced_ids
        ):

            # 从文档部件移除孤立的示例图片关系。
            obj_document.part.drop_rel(str_relationship_id)

# 协调全部槽位清理、标题复用、正文插入和媒体关系收尾。
def replace_template_slots(
    obj_document: Any,
    # 槽位顺序和正文数据共同决定模板内容边界。
    list_slot_order: list[str],
    dict_sections: dict[str, list[dict[str, str]]],
    # 附图、公式转换与行内拆分均由调用方显式注入。
    list_figure_paths: list[Path],
    func_convert_formula: Callable[[str, bool, str], Any],
    func_split_inline_equations: Callable[[str], list[dict[str, str]]],
    # 渲染模式决定保留 OMML，还是在保存后替换为原生 MathType OLE。
    str_equation_mode: str,
) -> list[dict[str, Any]]:
    """在原模板节点上完成全部正式正文槽位替换。

    参数：
    - `obj_document`：从参考模板打开的 Word 文档对象。
    - `list_slot_order`：一维标题序列，不涉及数值 shape、dtype 或物理单位。
    - `dict_sections`：槽位标题到代理起草正文块的映射。
    - `list_figure_paths`：需要嵌入主稿的正式 PNG 路径。
    - `func_convert_formula`：将公式文本转换为 OMML 节点的回调。
    - `func_split_inline_equations`：把正文拆成文字与行内公式片段的回调。
    - `str_equation_mode`：Office OMML 或 MathType OLE 中间渲染模式。

    返回：
    - `list[dict[str, Any]]`：按文档顺序记录的公式源文本与布局类型。

    异常：
    - 模板槽位不完整、公式转换或附图嵌入失败时继续抛出底层异常。
    """

    # 先定位全部标题，任何槽位缺失时都不修改模板正文。
    dict_slot_paragraphs = collect_slot_paragraphs(obj_document, list_slot_order)  # 模板标题节点映射

    # 保存每个槽位从模板提取的正文版式属性。
    dict_slot_properties: dict[str, Any | None] = {}  # 槽位正文版式映射

    # 逐槽位清除示例内容，并保留原始标题节点。
    for int_index, str_slot_heading in enumerate(list_slot_order):

        # 获取当前标题节点作为清理起点。
        obj_current_heading = dict_slot_paragraphs[str_slot_heading]  # 当前槽位标题段落

        # 末尾槽位以文档尾部作为清理边界，其余槽位以下一标题为边界。
        obj_next_heading = (  # 下一槽位标题段落
            dict_slot_paragraphs[list_slot_order[int_index + 1]]  # 顺序中的下一标题节点
            if int_index + 1 < len(list_slot_order)  # 当前槽位不是末尾槽位
            else None  # 末尾槽位没有下一标题
        )

        # 提取当前章节的正文样式，并删除两个标题之间的示例材料。
        dict_slot_properties[str_slot_heading] = remove_slot_example_content(  # 本章节正文版式
            obj_current_heading,  # 本章节的原始标题锚点
            obj_next_heading,  # 本章节允许清理的结束边界
        )

        # 将标准标题文本写回原始标题节点。
        set_heading_text(obj_current_heading, str_slot_heading)

    # 将发明名称同步写入模板信息表。
    populate_information_table(obj_document, dict_sections)

    # 公式编号从一开始，便于报告公式转换数量和生成顺序。
    int_formula_index = 1  # 首个待分配的公式编号

    # 保存公式源文本及布局类型，供 MathType 模式完成 OLE 二次写入。
    list_formula_records: list[dict[str, Any]] = []  # 文档顺序公式记录

    # 按模板槽位顺序插入经过确认的正式正文。
    for str_slot_heading in list_slot_order:

        # 组装当前槽位稳定依赖，正文块和公式编号仍由调用参数显式表达。
        dict_render_context = {  # 当前槽位渲染上下文
            "document": obj_document,  # 本轮唯一的模板文档对象
            "template_properties": dict_slot_properties[str_slot_heading],  # 标题下正文继承的原版式
            "slot_heading": str_slot_heading,  # 当前正式槽位标题
            "formula_converter": func_convert_formula,  # Office 原生公式转换回调
            "inline_splitter": func_split_inline_equations,  # 行内公式片段识别回调
            "equation_mode": str_equation_mode,  # 当前公式兼容模式
            "formula_records": list_formula_records,  # 原生 MathType 写入清单
            "figure_paths": list_figure_paths,  # 当前案件正式附图路径
        }

        # 装配本章节正文，并延续跨章节递增的公式媒体编号。
        int_formula_index = append_slot_blocks(  # 后续章节使用的公式编号
            dict_slot_paragraphs[str_slot_heading],  # 本章节保留下来的标题节点
            dict_sections.get(str_slot_heading, []),  # 经确认后进入本章节的正文块
            int_formula_index,  # 本章节首个可用公式编号
            dict_render_context,  # 文档、样式和媒体渲染依赖
        )

    # 删除模板示例图片留下的孤立关系，避免无用媒体进入交付件。
    purge_unreferenced_image_relationships(obj_document)

    # 返回稳定公式顺序，使保存后的 Word COM 替换不会重新解析正文。
    return list_formula_records
